# inclined_master.py

import streamlit as st
import numpy as np
import pandas as pd
import io
import os
import re
import matplotlib as mpl
import matplotlib.pyplot as plt
from matplotlib.patches import Polygon
from docx import Document
from docx.shared import Cm, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# 💡 استدعاء قواعد بيانات القطاعات والشدات الموحدة من الملفات المرفقة
try:
    from config import SECTIONS_DB, STRUTS_DB, PROP_DB, CUPLOCK_DB, RINGLOCK_DB
except ImportError:
    st.error("⚠️ برجاء التأكد من وجود ملف config.py")
    SECTIONS_DB = {}
    STRUTS_DB = {}
    PROP_DB = {}
    CUPLOCK_DB = {}
    RINGLOCK_DB = {}

try:
    from math_solver import get_prop_allowable, get_scaffold_allowable
except ImportError:
    def get_prop_allowable(*args): return 20.0
    def get_scaffold_allowable(*args): return 40.0

# =========================================================
# 0. Helper Functions & Styles
# =========================================================
def get_shoring_capacity(t_nm, subtype, unb, req_ext):
    try:
        if t_nm == "Shorebrace Frame": return 54.00
        elif t_nm == "Cup-lock": return get_scaffold_allowable("Cup-lock", subtype, unb)
        elif t_nm == "Ringlock": return get_scaffold_allowable("Ringlock", subtype, unb)
        elif t_nm == "Acrow Prop": return get_prop_allowable(subtype, req_ext, True)
    except:
        pass
    return 20.0 

def get_valid_struts(req_len, struts_db):
    valid = []
    for name, props in struts_db.items():
        min_l, max_l = 0.0, 99.0
        if isinstance(props, dict) and 'min' in props and 'max' in props:
            min_l, max_l = props['min'], props['max']
        else:
            m = re.search(r"\((\d+\.?\d*):(\d+\.?\d*)m\)", name)
            if m:
                min_l, max_l = float(m.group(1)), float(m.group(2))
        
        if min_l <= req_len <= max_l:
            valid.append(name)
            
    if not valid:
        return list(struts_db.keys()) if struts_db else ["PPH (Fallback)"]
        
    def priority(name):
        n = name.upper()
        if "PPH" in n: return 1
        if "PPS" in n: return 2
        if "TILT" in n: return 3
        if "MMP" in n: return 4
        return 5
        
    return sorted(valid, key=priority)

def apply_plot_styles():
    mpl.rcParams['font.family'] = 'sans-serif'
    mpl.rcParams['font.sans-serif'] = ['Arial', 'Helvetica', 'DejaVu Sans']
    mpl.rcParams['axes.linewidth'] = 0.3
    mpl.rcParams['font.size'] = 7
    mpl.rcParams['font.weight'] = 'normal'

def get_short_name(sec_name):
    return re.sub(r'\s*\(.*?\)', '', sec_name).strip()

# =========================================================
# 1. Geometry & Mesh Generator
# =========================================================
def build_fea_mesh(L_segs, L_rem, X_segs, X_rem, angle_rad, applied_loads, inc_sec, base_sec, strut_types, corner_sup, base_sups):
    inc_key_pts = [0.0]
    curr = 0.0
    for seg in L_segs:
        curr += seg
        inc_key_pts.append(curr)
    if L_rem > 0:
        inc_key_pts.append(curr + L_rem)
    
    L_tot = curr + L_rem
    
    for ld in applied_loads:
        inc_key_pts.append(ld['start'])
        if ld['type'] != 'Point Load':
            inc_key_pts.append(ld['end'])
            
    inc_key_pts = sorted(list(set([round(p, 4) for p in inc_key_pts if 0 <= p <= L_tot + 1e-5])))
    
    inc_nodes_L = []
    for i in range(len(inc_key_pts)-1):
        A = inc_key_pts[i]
        B = inc_key_pts[i+1]
        num_sub = max(1, int(np.ceil((B - A) / 0.15))) 
        pts = np.linspace(A, B, num_sub+1)
        for p in pts[:-1]:
            inc_nodes_L.append(p)
    inc_nodes_L.append(inc_key_pts[-1])
    
    nodes = []
    inc_node_indices = []
    
    nodes.append([0.0, 0.0])
    inc_node_indices.append(0)
    
    for L_val in inc_nodes_L[1:]:
        nodes.append([L_val * np.cos(angle_rad), L_val * np.sin(angle_rad)])
        inc_node_indices.append(len(nodes)-1)
        
    X_cum = 0.0
    base_x_pts = [0.0]
    for X_seg in X_segs:
        X_cum += X_seg
        base_x_pts.append(X_cum)
    if X_rem > 0:
        base_x_pts.append(X_cum + X_rem)
        
    for sup in base_sups:
        base_x_pts.append(sup['x'])
        
    base_x_pts = sorted(list(set([round(x, 4) for x in base_x_pts])))
    
    base_node_indices = []
    for x in base_x_pts:
        if x == 0.0:
            base_node_indices.append(0)
        else:
            nodes.append([x, 0.0])
            base_node_indices.append(len(nodes)-1)
            
    supports_list = []
    supports_list.append({'node': 0, 'type': corner_sup['type'], 'angle': corner_sup.get('angle', 0.0)})
    for sup in base_sups:
        idx = base_x_pts.index(round(sup['x'], 4))
        n_idx = base_node_indices[idx]
        supports_list.append({'node': n_idx, 'type': sup['type'], 'angle': sup.get('angle', 0.0)})
        
    display_nodes = set([s['node'] for s in supports_list])
    display_nodes.add(inc_node_indices[-1])
    if len(base_node_indices) > 0:
        display_nodes.add(base_node_indices[-1])
    
    target_Ls = [sum(L_segs[:j+1]) for j in range(len(L_segs))]
    for j in range(len(L_segs)):
        target_L = round(target_Ls[j], 4)
        if target_L in inc_nodes_L:
            idx = inc_nodes_L.index(target_L)
            display_nodes.add(inc_node_indices[idx])
            
    for ld in applied_loads:
        if ld['type'] == 'Point Load':
            try:
                idx = inc_nodes_L.index(round(ld['start'], 4))
                display_nodes.add(inc_node_indices[idx])
            except ValueError: pass

    elements = []
    nodal_loads = []
    E_st = 210000000.0 
    inc_props = SECTIONS_DB.get(inc_sec, {'E': 2100.0, 'I': 412.0, 'Mall': 13.1, 'Qall': 100.8})
    
    for i in range(len(inc_node_indices)-1):
        n1 = inc_node_indices[i]
        n2 = inc_node_indices[i+1]
        L_mid = (inc_nodes_L[i] + inc_nodes_L[i+1]) / 2.0
        
        p_x1, p_y1 = 0.0, 0.0
        p_x2, p_y2 = 0.0, 0.0
        
        for ld in applied_loads:
            if ld['type'] == 'Point Load': continue
            if ld['start'] - 1e-4 <= L_mid <= ld['end'] + 1e-4:
                L_len = max(ld['end'] - ld['start'], 1e-5)
                w_a = ld['w1'] + (ld['w2'] - ld['w1']) * (inc_nodes_L[i] - ld['start']) / L_len
                w_b = ld['w1'] + (ld['w2'] - ld['w1']) * (inc_nodes_L[i+1] - ld['start']) / L_len
                
                if ld['dir'] == 'Gravity (Vertical ↓)':
                    c, s = np.cos(angle_rad), np.sin(angle_rad)
                    p_x1 += -w_a * s; p_y1 += -w_a * c
                    p_x2 += -w_b * s; p_y2 += -w_b * c
                else:
                    p_y1 += -w_a; p_y2 += -w_b
                    
        elements.append({
            'type': 'frame', 'group': 'inclined', 'sec': inc_sec,
            'n1': n1, 'n2': n2, 'px1': p_x1, 'py1': p_y1, 'px2': p_x2, 'py2': p_y2,
            'E': inc_props.get('E', 2100.0) * 10000.0, 'A': 0.00343, 'I': inc_props.get('I', 412.0) / 100000000.0
        })
        
    for ld in applied_loads:
        if ld['type'] == 'Point Load':
            try:
                idx = inc_nodes_L.index(round(ld['start'], 4))
                n_idx = inc_node_indices[idx]
                c, s = np.cos(angle_rad), np.sin(angle_rad)
                if ld['dir'] == 'Gravity (Vertical ↓)':
                    nodal_loads.append({'node': n_idx, 'Fx': 0.0, 'Fy': -ld['w1']})
                else:
                    nodal_loads.append({'node': n_idx, 'Fx': ld['w1']*s, 'Fy': -ld['w1']*c})
            except ValueError: pass
                
    if base_sec != "None (Direct to Ground)":
        base_props = SECTIONS_DB.get(base_sec, {'E': 2100.0, 'I': 412.0, 'Mall': 13.1, 'Qall': 100.8})
        for i in range(len(base_node_indices)-1):
            elements.append({
                'type': 'frame', 'group': 'base', 'sec': base_sec,
                'n1': base_node_indices[i], 'n2': base_node_indices[i+1],
                'px1': 0.0, 'py1': 0.0, 'px2': 0.0, 'py2': 0.0,
                'E': base_props.get('E', 2100.0) * 10000.0, 'A': 0.00343, 'I': base_props.get('I', 412.0) / 100000000.0
            })
        
    X_cum_strut = 0.0
    for j in range(len(L_segs)):
        target_L = round(target_Ls[j], 4)
        X_cum_strut += X_segs[j]
        if target_L in inc_nodes_L:
            idx_inc = inc_nodes_L.index(target_L)
            n_inc = inc_node_indices[idx_inc]
            
            idx_base = base_x_pts.index(round(X_cum_strut, 4))
            n_base = base_node_indices[idx_base]
            display_nodes.add(n_base)
            
            st_props = STRUTS_DB.get(strut_types[j], {'allow': 50.0}) 
            elements.append({
                'type': 'truss', 'group': 'strut', 'sec': strut_types[j],
                'n1': n_base, 'n2': n_inc,
                'E': 21000000.0, 'A': 0.001
            })
            
    X_tot = base_x_pts[-1]
    return nodes, elements, nodal_loads, L_tot, X_tot, display_nodes, supports_list

# =========================================================
# 2. Advanced 2D Frame FEA Solver
# =========================================================
def solve_fea_engine(nodes, elements, nodal_loads, supports_list):
    num_nodes = len(nodes)
    NDOF = num_nodes * 3
    K = np.zeros((NDOF, NDOF))
    F = np.zeros(NDOF)
    
    for el in elements:
        n1, n2 = el['n1'], el['n2']
        x1, y1 = nodes[n1][0], nodes[n1][1]
        x2, y2 = nodes[n2][0], nodes[n2][1]
        L = np.hypot(x2 - x1, y2 - y1)
        if L < 1e-5: continue
        c, s = (x2 - x1) / L, (y2 - y1) / L
        el['L'], el['c'], el['s'] = L, c, s
        
        E, A, I = el['E'], el['A'], el.get('I', 0.00005)
        
        T = np.array([
            [c, s, 0, 0, 0, 0], [-s, c, 0, 0, 0, 0], [0, 0, 1, 0, 0, 0],
            [0, 0, 0, c, s, 0], [0, 0, 0, -s, c, 0], [0, 0, 0, 0, 0, 1]
        ])
        
        if el['type'] == 'truss':
            k_loc = np.zeros((6, 6))
            k_loc[0, 0] = E * A / L
            k_loc[3, 3] = E * A / L
            k_loc[0, 3] = -E * A / L
            k_loc[3, 0] = -E * A / L
        else:
            k_loc = np.array([
                [E*A/L, 0, 0, -E*A/L, 0, 0],
                [0, 12*E*I/L**3, 6*E*I/L**2, 0, -12*E*I/L**3, 6*E*I/L**2],
                [0, 6*E*I/L**2, 4*E*I/L, 0, -6*E*I/L**2, 2*E*I/L],
                [-E*A/L, 0, 0, E*A/L, 0, 0],
                [0, -12*E*I/L**3, -6*E*I/L**2, 0, 12*E*I/L**3, -6*E*I/L**2],
                [0, 6*E*I/L**2, 2*E*I/L, 0, -6*E*I/L**2, 4*E*I/L]
            ])
            
            px1, py1, px2, py2 = el['px1'], el['py1'], el['px2'], el['py2']
            f_eq_loc = np.array([
                (2*px1 + px2)*L/6.0,
                (7*py1 + 3*py2)*L/20.0,
                (3*py1 + 2*py2)*L**2/60.0,
                (px1 + 2*px2)*L/6.0,
                (3*py1 + 7*py2)*L/20.0,
                -(2*py1 + 3*py2)*L**2/60.0
            ])
            f_eq_glob = T.T @ f_eq_loc
            dof_idx = [3*n1, 3*n1+1, 3*n1+2, 3*n2, 3*n2+1, 3*n2+2]
            for r in range(6): F[dof_idx[r]] += f_eq_glob[r]
            
        k_glob = T.T @ k_loc @ T
        dof_idx = [3*n1, 3*n1+1, 3*n1+2, 3*n2, 3*n2+1, 3*n2+2]
        for r in range(6):
            for col in range(6):
                K[dof_idx[r], dof_idx[col]] += k_glob[r, col]
                
    for nl in nodal_loads:
        F[3*nl['node'] + 0] += nl['Fx']
        F[3*nl['node'] + 1] += nl['Fy']
            
    K_orig = K.copy()
    fixed_dofs = []
    K_penalty = 1e12
    
    for sup in supports_list:
        n = sup['node']
        t = sup['type']
        a = sup['angle']
        if t == 'Fixed':
            fixed_dofs.extend([3*n, 3*n+1, 3*n+2])
        elif t == 'Hinged':
            fixed_dofs.extend([3*n, 3*n+1])
        elif t == 'Roller':
            if abs(a % 180) < 1e-5: 
                fixed_dofs.append(3*n+1)
            elif abs((a - 90) % 180) < 1e-5: 
                fixed_dofs.append(3*n)
            else:
                rad = np.radians(a)
                nx, ny = -np.sin(rad), np.cos(rad) 
                K[3*n, 3*n] += K_penalty * nx**2
                K[3*n+1, 3*n+1] += K_penalty * ny**2
                K[3*n, 3*n+1] += K_penalty * nx * ny
                K[3*n+1, 3*n] += K_penalty * nx * ny

    free_dof = [i for i in range(NDOF) if i not in fixed_dofs]
        
    K_ff = K[np.ix_(free_dof, free_dof)]
    F_f = F[free_dof]
    
    U = np.zeros(NDOF)
    try:
        U_f = np.linalg.solve(K_ff, F_f)
    except np.linalg.LinAlgError:
        U_f = np.linalg.lstsq(K_ff, F_f, rcond=None)[0]
    
    U[free_dof] = U_f
    R_reactions = K_orig @ U - F 
    
    for el in elements:
        n1, n2 = el['n1'], el['n2']
        c, s, L = el['c'], el['s'], el['L']
        E, A, I = el['E'], el['A'], el.get('I', 0.00005)
        
        dof_idx = [3*n1, 3*n1+1, 3*n1+2, 3*n2, 3*n2+1, 3*n2+2]
        u_glob = U[dof_idx]
        T = np.array([
            [c, s, 0, 0, 0, 0], [-s, c, 0, 0, 0, 0], [0, 0, 1, 0, 0, 0],
            [0, 0, 0, c, s, 0], [0, 0, 0, -s, c, 0], [0, 0, 0, 0, 0, 1]
        ])
        u_loc = T @ u_glob
        el['internal'] = {'u_loc': u_loc}
        
        if el['type'] == 'truss':
            N_val = (E * A / L) * (u_loc[3] - u_loc[0])
            el['internal'].update({'N': [N_val, N_val], 'V': [0,0], 'M': [0,0], 'x': [0, L], 'v_rel': [0,0]})
        else:
            k_loc = np.array([
                [E*A/L, 0, 0, -E*A/L, 0, 0],
                [0, 12*E*I/L**3, 6*E*I/L**2, 0, -12*E*I/L**3, 6*E*I/L**2],
                [0, 6*E*I/L**2, 4*E*I/L, 0, -6*E*I/L**2, 2*E*I/L],
                [-E*A/L, 0, 0, E*A/L, 0, 0],
                [0, -12*E*I/L**3, -6*E*I/L**2, 0, 12*E*I/L**3, -6*E*I/L**2],
                [0, 6*E*I/L**2, 2*E*I/L, 0, -6*E*I/el['L']**2, 4*E*I/el['L']]
            ])
            px1, py1, px2, py2 = el['px1'], el['py1'], el['px2'], el['py2']
            f_eq_loc = np.array([
                (2*px1 + px2)*L/6.0,
                (7*py1 + 3*py2)*L/20.0,
                (3*py1 + 2*py2)*L**2/60.0,
                (px1 + 2*px2)*L/6.0,
                (3*py1 + 7*py2)*L/20.0,
                -(2*py1 + 3*py2)*L**2/60.0
            ])
            f_end = k_loc @ u_loc - f_eq_loc
            
            xs = np.linspace(0, L, 51) 
            N_arr = np.zeros_like(xs)
            V_arr = np.zeros_like(xs)
            M_arr = np.zeros_like(xs)
            v_rel_arr = np.zeros_like(xs)
            
            v1, theta1 = u_loc[1], u_loc[2]
            v2, theta2 = u_loc[4], u_loc[5]
            w_avg = (py1 + py2) / 2.0 
            
            for i, x in enumerate(xs):
                N_arr[i] = -f_end[0] - (px1*x + (px2-px1)*x**2/(2*L))
                V_arr[i] = f_end[1] + (py1*x + (py2-py1)*x**2/(2*L))
                M_arr[i] = -f_end[2] + f_end[1]*x + py1*x**2/2.0 + (py2-py1)*x**3/(6*L)
                
                xi = x / L
                N1 = 1 - 3*xi**2 + 2*xi**3
                N2 = x * (1 - xi)**2
                N3 = 3*xi**2 - 2*xi**3
                N4 = x * (xi**2 - xi)
                
                v_shape = v1*N1 + theta1*N2 + v2*N3 + theta2*N4
                v_load = (w_avg * x**2 * (L - x)**2) / (24 * E * I) 
                v_tot = v_shape + v_load
                v_chord = v1 + xi * (v2 - v1) 
                v_rel_arr[i] = v_tot - v_chord 
                
            el['internal'].update({'N': N_arr, 'V': V_arr, 'M': M_arr, 'x': xs, 'v_rel': v_rel_arr})
            
    return U, R_reactions

# =========================================================
# 3. Plotting Engines (Independent Subplots)
# =========================================================
def get_img_buf(fig):
    plt.tight_layout()
    buf = io.BytesIO()
    fig.savefig(buf, format='png', dpi=150, bbox_inches='tight', transparent=True)
    plt.close(fig)
    buf.seek(0)
    return buf

def draw_base_geometry(ax, nodes, elements, supports_list):
    for el in elements:
        if el['group'] == 'base' and el['sec'] == "None (Direct to Ground)":
            continue
        n1, n2 = nodes[el['n1']], nodes[el['n2']]
        color = 'black' if el['type'] == 'frame' else 'gray'
        style = '-' if el['type'] == 'frame' else '--'
        ax.plot([n1[0], n2[0]], [n1[1], n2[1]], color=color, linestyle=style, linewidth=0.5, zorder=1)
        
    for sup in supports_list:
        n = sup['node']
        x, y = nodes[n][0], nodes[n][1]
        t = sup['type']
        a = sup['angle']
        
        rad = np.radians(a)
        nx, ny = np.sin(rad), -np.cos(rad) 
        tx, ty = -ny, nx 
        
        # 💡 تم تفرغ الركائز وجعلها بلون أخضر فاتح SAP2000-Style
        if t == 'Fixed':
            ax.plot(x, y, marker='s', markerfacecolor='none', markeredgecolor='limegreen', markersize=6, zorder=5)
            ax.plot([x - 0.2*tx, x + 0.2*tx], [y - 0.2*ty, y + 0.2*ty], color='limegreen', lw=2, zorder=4)
        elif t == 'Hinged':
            h, w = 0.3, 0.2
            p1 = (x, y)
            p2 = (x + h*nx + w*tx, y + h*ny + w*ty)
            p3 = (x + h*nx - w*tx, y + h*ny - w*ty)
            poly = Polygon([p1, p2, p3], facecolor='none', edgecolor='limegreen', lw=1.2, zorder=5)
            ax.add_patch(poly)
            ax.plot([p2[0]+0.1*tx, p3[0]-0.1*tx], [p2[1]+0.1*ty, p3[1]-0.1*ty], color='limegreen', lw=1.5, zorder=4)
        elif t == 'Roller':
            h, w, r = 0.25, 0.15, 0.08
            p1 = (x, y)
            p2 = (x + h*nx + w*tx, y + h*ny + w*ty)
            p3 = (x + h*nx - w*tx, y + h*ny - w*ty)
            poly = Polygon([p1, p2, p3], facecolor='none', edgecolor='limegreen', lw=1.2, zorder=5)
            ax.add_patch(poly)
            
            circ_x, circ_y = x + (h + r)*nx, y + (h + r)*ny
            circle = plt.Circle((circ_x, circ_y), r, facecolor='none', edgecolor='limegreen', lw=1.2, zorder=5)
            ax.add_patch(circle)
            
            # 💡 تصغير الخط ولصقه تماماً في الدائرة للـ Roller
            base_dist = h + 2*r
            line_w = 0.12
            lx1, ly1 = x + base_dist*nx - line_w*tx, y + base_dist*ny - line_w*ty
            lx2, ly2 = x + base_dist*nx + line_w*tx, y + base_dist*ny + line_w*ty
            ax.plot([lx1, lx2], [ly1, ly2], color='limegreen', lw=1.5, zorder=4)

def draw_section_names(ax, elements, nodes, L_tot, X_tot, angle_deg, inc_sec, base_sec, is_n_diagram=False):
    angle_rad = np.radians(angle_deg)
    c_ang, s_ang = np.cos(angle_rad), np.sin(angle_rad)
    
    inc_mid_x = (L_tot/2) * c_ang
    inc_mid_y = (L_tot/2) * s_ang
    ax.text(inc_mid_x - s_ang*0.15, inc_mid_y + c_ang*0.15, get_short_name(inc_sec), color='black', fontsize=7, alpha=0.9, ha='center', va='center', rotation=angle_deg, fontname='Arial')
    
    if base_sec != "None (Direct to Ground)":
        base_mid_x = X_tot/2
        ax.text(base_mid_x, -0.2, get_short_name(base_sec), color='black', fontsize=7, alpha=0.9, ha='center', va='center', fontname='Arial')
    
    drawn_struts = set()
    for el in elements:
        if el['group'] == 'strut':
            sig = f"{el['n1']}_{el['n2']}"
            if sig not in drawn_struts:
                n1, n2 = nodes[el['n1']], nodes[el['n2']]
                dx, dy = n2[0]-n1[0], n2[1]-n1[1]
                L_s = np.hypot(dx, dy)
                nx, ny = -dy/L_s, dx/L_s
                mid_x, mid_y = (n1[0]+n2[0])/2, (n1[1]+n2[1])/2
                rot = np.degrees(np.arctan2(dy, dx))
                
                if is_n_diagram:
                    ax.text(mid_x - nx*0.3, mid_y - ny*0.3, get_short_name(el['sec']), color='black', fontsize=6, alpha=0.9, ha='center', va='center', rotation=rot, fontname='Arial')
                else:
                    ax.text(mid_x + nx*0.15, mid_y + ny*0.15, get_short_name(el['sec']), color='black', fontsize=6, alpha=0.9, ha='center', va='center', rotation=rot, fontname='Arial')
                drawn_struts.add(sig)

def draw_reaction_arrow(ax, node_x, node_y, force_mag, axis_nx, axis_ny):
    if abs(force_mag) < 0.1: return
    arr_L = 0.8
    sgn = np.sign(force_mag)
    dx = sgn * axis_nx
    dy = sgn * axis_ny
    start_x = node_x - arr_L * dx
    start_y = node_y - arr_L * dy
    arr_c = 'blue' if force_mag >= 0 else 'red'
    ax.arrow(start_x, start_y, arr_L*dx, arr_L*dy, length_includes_head=True, 
             head_width=0.15, head_length=0.2, fc=arr_c, ec=arr_c, lw=1.0, zorder=5)
    ax.text(start_x - 0.25*dx, start_y - 0.25*dy, f"{abs(force_mag):.1f}", 
            color='black', fontsize=7, fontname='Arial', ha='center', va='center')

def plot_live_geometry(nodes, elements, applied_loads, L_segs, X_segs, angle_deg, inc_sec, base_sec, L_tot, X_tot, supports_list):
    apply_plot_styles()
    fig, ax = plt.subplots(figsize=(6, 5))
    ax.set_aspect('equal', adjustable='datalim')
    ax.axis('off')
    
    draw_base_geometry(ax, nodes, elements, supports_list)
    draw_section_names(ax, elements, nodes, L_tot, X_tot, angle_deg, inc_sec, base_sec)

    angle_rad = np.radians(angle_deg)
    c_ang, s_ang = np.cos(angle_rad), np.sin(angle_rad)
    
    curr_l = 0.0
    for i, seg in enumerate(L_segs):
        px = (curr_l + seg/2) * c_ang
        py = (curr_l + seg/2) * s_ang
        ax.text(px - s_ang*0.6, py + c_ang*0.6, f"L{i+1}={seg:.2f}m", color='black', fontsize=7, rotation=angle_deg, ha='center', va='center', fontname='Arial')
        curr_l += seg
        
    if base_sec != "None (Direct to Ground)":
        curr_x = 0.0
        for i, seg in enumerate(X_segs):
            px = curr_x + seg/2
            py = 0.0
            ax.text(px, py - 0.6, f"X{i+1}={seg:.2f}m", color='black', fontsize=7, ha='center', va='center', fontname='Arial')
            curr_x += seg

    if applied_loads:
        max_w = max([max(abs(ld['w1']), abs(ld['w2'])) for ld in applied_loads] + [1.0])
        scale_ld = 1.2 / max_w
        for ld in applied_loads:
            w1, w2 = ld['w1'], ld['w2']
            start_L, end_L = ld['start'], ld['end']
            dir_type = ld['dir']
            px1, py1 = start_L * c_ang, start_L * s_ang
            px2, py2 = end_L * c_ang, end_L * s_ang
            
            if ld['type'] == 'Point Load':
                arrow_len = 1.0
                pt_c = 'blue' if w1 >= 0 else 'red'
                if dir_type == 'Gravity (Vertical ↓)':
                    ax.arrow(px1, py1 + arrow_len + 0.1, 0, -arrow_len, head_width=0.15, head_length=0.2, length_includes_head=True, fc=pt_c, ec=pt_c, zorder=4, linewidth=0.5)
                    ax.text(px1, py1 + arrow_len + 0.3, f"{w1}", color='black', fontsize=8, ha='center', fontname='Arial')
                else:
                    start_x = px1 - s_ang*(arrow_len+0.1)
                    start_y = py1 + c_ang*(arrow_len+0.1)
                    ax.arrow(start_x, start_y, s_ang*arrow_len, -c_ang*arrow_len, head_width=0.15, head_length=0.2, length_includes_head=True, fc=pt_c, ec=pt_c, zorder=4, linewidth=0.5)
                    ax.text(start_x - s_ang*0.2, start_y + c_ang*0.2, f"{w1}", color='black', fontsize=8, ha='center', rotation=angle_deg, fontname='Arial')
            else:
                ld_c = 'blue' if (w1 + w2)/2.0 >= 0 else 'red'
                if dir_type == 'Gravity (Vertical ↓)':
                    hx1, hy1 = px1, py1 + w1 * scale_ld
                    hx2, hy2 = px2, py2 + w2 * scale_ld
                else:
                    hx1, hy1 = px1 - s_ang * w1 * scale_ld, py1 + c_ang * w1 * scale_ld
                    hx2, hy2 = px2 - s_ang * w2 * scale_ld, py2 + c_ang * w2 * scale_ld
                    
                poly = Polygon([(px1,py1), (hx1,hy1), (hx2,hy2), (px2,py2)], facecolor='none', edgecolor=ld_c, linewidth=0.8, zorder=3)
                ax.add_patch(poly)
                
                num_arrows = max(3, int((end_L - start_L) / 0.4))
                xs = np.linspace(start_L, end_L, num_arrows)
                for x_dist in xs:
                    w_curr = w1 + (w2 - w1) * (x_dist - start_L) / max(end_L - start_L, 1e-5)
                    if abs(w_curr) < 0.1: continue
                    px = x_dist * c_ang
                    py = x_dist * s_ang
                    hl = w_curr * scale_ld
                    arr_c = 'blue' if w_curr >= 0 else 'red'
                    if dir_type == 'Gravity (Vertical ↓)':
                        ax.arrow(px, py + hl, 0, -hl, head_width=0.08, head_length=0.1, length_includes_head=True, fc=arr_c, ec=arr_c, linewidth=0.3, zorder=2)
                    else:
                        ax.arrow(px - s_ang*hl, py + c_ang*hl, s_ang*hl, -c_ang*hl, head_width=0.08, head_length=0.1, length_includes_head=True, fc=arr_c, ec=arr_c, linewidth=0.3, zorder=2)
                        
                if dir_type == 'Gravity (Vertical ↓)':
                    ax.text(px1, py1 + w1*scale_ld + 0.15, f"{w1}", color='black', fontsize=7, ha='center', fontname='Arial')
                    ax.text(px2, py2 + w2*scale_ld + 0.15, f"{w2}", color='black', fontsize=7, ha='center', fontname='Arial')
                else:
                    ax.text(hx1 - s_ang*0.15, hy1 + c_ang*0.15, f"{w1}", color='black', fontsize=7, ha='center', rotation=angle_deg, fontname='Arial')
                    ax.text(hx2 - s_ang*0.15, hy2 + c_ang*0.15, f"{w2}", color='black', fontsize=7, ha='center', rotation=angle_deg, fontname='Arial')

    return get_img_buf(fig)

def plot_sap2000_diagrams(nodes, elements, R_reactions, scales, display_nodes, applied_loads, angle_deg, L_tot, X_tot, inc_sec, base_sec, supports_list):
    apply_plot_styles()
    angle_rad = np.radians(angle_deg)
    c_ang, s_ang = np.cos(angle_rad), np.sin(angle_rad)
    
    figs_dict = {}
    
    # --- 1. Load Diagram ---
    fig_ld, ax_ld = plt.subplots(figsize=(6, 5))
    ax_ld.set_aspect('equal', adjustable='datalim')
    ax_ld.axis('off')
    draw_base_geometry(ax_ld, nodes, elements, supports_list)
    draw_section_names(ax_ld, elements, nodes, L_tot, X_tot, angle_deg, inc_sec, base_sec)
    
    if applied_loads:
        max_w = max([max(abs(ld['w1']), abs(ld['w2'])) for ld in applied_loads] + [1.0])
        scale_ld = 1.2 / max_w
        for ld in applied_loads:
            w1, w2 = ld['w1'], ld['w2']
            start_L, end_L = ld['start'], ld['end']
            dir_type = ld['dir']
            px1, py1 = start_L * c_ang, start_L * s_ang
            px2, py2 = end_L * c_ang, end_L * s_ang
            
            if ld['type'] == 'Point Load':
                arrow_len = 1.0
                pt_c = 'blue' if w1 >= 0 else 'red'
                if dir_type == 'Gravity (Vertical ↓)':
                    ax_ld.arrow(px1, py1 + arrow_len + 0.1, 0, -arrow_len, head_width=0.15, head_length=0.2, length_includes_head=True, fc=pt_c, ec=pt_c, zorder=4, linewidth=0.5)
                    ax_ld.text(px1, py1 + arrow_len + 0.3, f"{w1}", color='black', fontsize=8, ha='center', fontname='Arial')
                else:
                    start_x = px1 - s_ang*(arrow_len+0.1)
                    start_y = py1 + c_ang*(arrow_len+0.1)
                    ax_ld.arrow(start_x, start_y, s_ang*arrow_len, -c_ang*arrow_len, head_width=0.15, head_length=0.2, length_includes_head=True, fc=pt_c, ec=pt_c, zorder=4, linewidth=0.5)
                    ax_ld.text(start_x - s_ang*0.2, start_y + c_ang*0.2, f"{w1}", color='black', fontsize=8, ha='center', rotation=angle_deg, fontname='Arial')
            else:
                ld_c = 'blue' if (w1 + w2)/2.0 >= 0 else 'red'
                if dir_type == 'Gravity (Vertical ↓)':
                    hx1, hy1 = px1, py1 + w1 * scale_ld
                    hx2, hy2 = px2, py2 + w2 * scale_ld
                else:
                    hx1, hy1 = px1 - s_ang * w1 * scale_ld, py1 + c_ang * w1 * scale_ld
                    hx2, hy2 = px2 - s_ang * w2 * scale_ld, py2 + c_ang * w2 * scale_ld
                    
                poly = Polygon([(px1,py1), (hx1,hy1), (hx2,hy2), (px2,py2)], facecolor='none', edgecolor=ld_c, linewidth=0.8, zorder=3)
                ax_ld.add_patch(poly)
                
                num_arrows = max(3, int((end_L - start_L) / 0.4))
                xs = np.linspace(start_L, end_L, num_arrows)
                for x_dist in xs:
                    w_curr = w1 + (w2 - w1) * (x_dist - start_L) / max(end_L - start_L, 1e-5)
                    if abs(w_curr) < 0.1: continue
                    px = x_dist * c_ang
                    py = x_dist * s_ang
                    hl = w_curr * scale_ld
                    arr_c = 'blue' if w_curr >= 0 else 'red'
                    if dir_type == 'Gravity (Vertical ↓)':
                        ax_ld.arrow(px, py + hl, 0, -hl, head_width=0.08, head_length=0.1, length_includes_head=True, fc=arr_c, ec=arr_c, linewidth=0.3, zorder=2)
                    else:
                        ax_ld.arrow(px - s_ang*hl, py + c_ang*hl, s_ang*hl, -c_ang*hl, head_width=0.08, head_length=0.1, length_includes_head=True, fc=arr_c, ec=arr_c, linewidth=0.3, zorder=2)
                        
                if dir_type == 'Gravity (Vertical ↓)':
                    ax_ld.text(px1, py1 + w1*scale_ld + 0.15, f"{w1}", color='black', fontsize=7, ha='center', fontname='Arial')
                    ax_ld.text(px2, py2 + w2*scale_ld + 0.15, f"{w2}", color='black', fontsize=7, ha='center', fontname='Arial')
                else:
                    ax_ld.text(hx1 - s_ang*0.15, hy1 + c_ang*0.15, f"{w1}", color='black', fontsize=7, ha='center', rotation=angle_deg, fontname='Arial')
                    ax_ld.text(hx2 - s_ang*0.15, hy2 + c_ang*0.15, f"{w2}", color='black', fontsize=7, ha='center', rotation=angle_deg, fontname='Arial')
    figs_dict['Load'] = get_img_buf(fig_ld)

    # --- 2. Reactions Diagram ---
    fig_react, ax_react = plt.subplots(figsize=(6, 5))
    ax_react.set_aspect('equal', adjustable='datalim')
    ax_react.axis('off')
    draw_base_geometry(ax_react, nodes, elements, supports_list)
    draw_section_names(ax_react, elements, nodes, L_tot, X_tot, angle_deg, inc_sec, base_sec)
    
    for sup in supports_list:
        n = sup['node']
        t = sup['type']
        a = sup['angle']
        
        Rx, Ry = R_reactions[3*n], R_reactions[3*n+1]
        rad = np.radians(a)
        c_a, s_a = np.cos(rad), np.sin(rad)
        
        R_u = Rx * c_a + Ry * s_a
        R_v = -Rx * s_a + Ry * c_a
        x, y = nodes[n][0], nodes[n][1]
        
        if t == 'Roller':
            draw_reaction_arrow(ax_react, x, y, R_v, -s_a, c_a)
        else:
            draw_reaction_arrow(ax_react, x, y, R_u, c_a, s_a)
            draw_reaction_arrow(ax_react, x, y, R_v, -s_a, c_a)
                
    figs_dict['React'] = get_img_buf(fig_react)

    # --- Force Diagrams Helper ---
    def create_force_diagram(val_key, scale, color_pos, color_neg):
        fig_f, ax_f = plt.subplots(figsize=(6, 5))
        ax_f.set_aspect('equal', adjustable='datalim')
        ax_f.axis('off')
        
        draw_base_geometry(ax_f, nodes, elements, supports_list)
        draw_section_names(ax_f, elements, nodes, L_tot, X_tot, angle_deg, inc_sec, base_sec, is_n_diagram=(val_key=='N'))
        
        plotted_texts = set()
        def write_val(txt_x, txt_y, v, rot=0):
            if abs(v) >= 0.1:
                lbl = f"{abs(v):.1f}"
                sig = f"{round(txt_x,1)}_{round(txt_y,1)}"
                if sig not in plotted_texts:
                    ax_f.text(txt_x, txt_y, lbl, color='black', fontsize=7, fontname='Arial', ha='center', va='center', rotation=rot)
                    plotted_texts.add(sig)

        for el in elements:
            if el['group'] == 'base' and el['sec'] == "None (Direct to Ground)":
                continue
            
            n1, n2 = nodes[el['n1']], nodes[el['n2']]
            x1, y1 = n1[0], n1[1]
            x2, y2 = n2[0], n2[1]
            dx, dy = x2 - x1, y2 - y1
            L_s = np.hypot(dx, dy)
            if L_s < 1e-5: continue
            
            c, s = dx/L_s, dy/L_s
            rot_ang = np.degrees(np.arctan2(dy, dx))
            
            if el['type'] == 'truss' and val_key == 'N':
                val = el['internal']['N'][0]
                if abs(val) < 0.1: continue
                nx, ny = dy/L_s, -dx/L_s 
                h = max(0.4, abs(val) * scale) 
                color = color_pos if val >= 0 else color_neg
                
                p1, p2 = (x1, y1), (x2, y2)
                p3, p4 = (x2 + nx * h, y2 + ny * h), (x1 + nx * h, y1 + ny * h)
                
                ax_f.add_patch(Polygon([p1, p2, p3, p4], facecolor='none', edgecolor=color, linewidth=0.8, zorder=2))
                
                num_lines = max(5, int(L_s / 0.3))
                for i in range(1, num_lines):
                    frac = i / num_lines
                    lx, ly = x1 + frac * dx, y1 + frac * dy
                    ax_f.plot([lx, lx + nx * h], [ly, ly + ny * h], color=color, linewidth=0.3, alpha=0.6)
                    
                mid_h_x, mid_h_y = x1 + dx/2 + nx*h/2, y1 + dy/2 + ny*h/2
                write_val(mid_h_x, mid_h_y, val, rot_ang)
                continue
                
            if el['type'] == 'frame':
                xs_arr = el['internal']['x']
                vals_orig = el['internal'][val_key]
                plot_vals = -vals_orig 
                
                px_arr = x1 + c * xs_arr - s * plot_vals * scale
                py_arr = y1 + s * xs_arr + c * plot_vals * scale
                
                # رسم خط التغطية الخارجي مقطعاً حسب الإشارة (أزرق/أحمر)
                ax_f.plot([x1, px_arr[0]], [y1, py_arr[0]], color=color_pos if vals_orig[0] >= 0 else color_neg, linewidth=0.8)
                for k in range(len(px_arr)-1):
                    avg_v = (vals_orig[k] + vals_orig[k+1]) / 2.0
                    seg_color = color_pos if avg_v >= 0 else color_neg
                    ax_f.plot([px_arr[k], px_arr[k+1]], [py_arr[k], py_arr[k+1]], color=seg_color, linewidth=0.8)
                ax_f.plot([px_arr[-1], x2], [py_arr[-1], y2], color=color_pos if vals_orig[-1] >= 0 else color_neg, linewidth=0.8)
                
                num_lines = max(2, int(L_s / 0.4))
                for i in range(1, num_lines):
                    frac = i / num_lines
                    lx, ly = x1 + frac * dx, y1 + frac * dy
                    idx_val = int(frac * (len(plot_vals)-1))
                    lv = plot_vals[idx_val]
                    hx, hy = lx - s * lv * scale, ly + c * lv * scale
                    line_color = color_pos if vals_orig[idx_val] >= 0 else color_neg
                    ax_f.plot([lx, hx], [ly, hy], color=line_color, linewidth=0.3, alpha=0.6)
                    
                offset = 0.25
                v_start = plot_vals[0]
                if el['n1'] in display_nodes:
                    txt_x = x1 - s * v_start * scale - s * np.sign(v_start) * offset
                    txt_y = y1 + c * v_start * scale + c * np.sign(v_start) * offset
                    write_val(txt_x, txt_y, vals_orig[0])
                    
                v_end = plot_vals[-1]
                if el['n2'] in display_nodes:
                    txt_x = x2 - s * v_end * scale - s * np.sign(v_end) * offset
                    txt_y = y2 + c * v_end * scale + c * np.sign(v_end) * offset
                    write_val(txt_x, txt_y, vals_orig[-1])
                    
                if val_key == 'M':
                    for i in range(1, len(plot_vals)-1):
                        v_prev, v_curr, v_next = plot_vals[i-1], plot_vals[i], plot_vals[i+1]
                        if (v_curr > v_prev and v_curr > v_next) or (v_curr < v_prev and v_curr < v_next):
                            if abs(v_curr) > 0.1 and abs(v_curr) > 0.05 * max(abs(plot_vals)):
                                txt_x = x1 + c * xs_arr[i] - s * v_curr * scale - s * np.sign(v_curr) * offset
                                txt_y = y1 + s * xs_arr[i] + c * v_curr * scale + c * np.sign(v_curr) * offset
                                write_val(txt_x, txt_y, vals_orig[i])
                                
        return get_img_buf(fig_f)

    # توحيد ألوان الدياجرامات لتكون (أزرق للموجب، أحمر للسالب)
    figs_dict['N'] = create_force_diagram('N', scales['N'], 'blue', 'red')
    figs_dict['V'] = create_force_diagram('V', scales['V'], 'blue', 'red')
    figs_dict['M'] = create_force_diagram('M', scales['M'], 'blue', 'red')

    return figs_dict

# =========================================================
# 4. Report Generator for Inclined Systems
# =========================================================
def generate_inclined_report(sys_data):
    if os.path.exists("Acrow_Template.docx"):
        doc = Document("Acrow_Template.docx")
        doc.add_page_break()
    else:
        doc = Document()
        
    def force_ltr_left(p):
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        pPr = p._element.get_or_add_pPr()
        bidi = OxmlElement('w:bidi')
        bidi.set(qn('w:val'), '0')
        pPr.append(bidi)
        
    def add_line(text, bold=False):
        p = doc.add_paragraph()
        force_ltr_left(p)
        p.paragraph_format.line_spacing = 1.5
        r = p.add_run(text)
        r.font.name = 'Arial'
        r.font.size = Pt(12)
        r.font.bold = bold
        r.font.rtl = False
        
    def add_check(component, param, act, allw, unit):
        p = doc.add_paragraph()
        force_ltr_left(p)
        p.paragraph_format.line_spacing = 1.5
        r_title = p.add_run(f"• Check {component} ({param}):\n")
        r_title.bold = True
        r_title.font.rtl = False
        if allw > 9000:
            r_act = p.add_run(f"  Actual = {act:.2f} {unit}  (No Limit Required)")
            r_res = p.add_run("  SAFE")
            r_res.font.color.rgb = RGBColor(0, 128, 0)
        else:
            r_act = p.add_run(f"  Actual = {act:.2f} {unit}  <  Allowable = {allw:.2f} {unit}  ")
            res = "SAFE" if act <= allw else "UNSAFE"
            r_res = p.add_run(res)
            r_res.font.color.rgb = RGBColor(0, 128, 0) if res == "SAFE" else RGBColor(255, 0, 0)
        r_act.font.rtl = False
        r_res.font.bold = True
        r_res.font.rtl = False
    
    p_title = doc.add_paragraph()
    force_ltr_left(p_title)
    run_title = p_title.add_run("CALCULATION SHEET FOR INCLINED FORMWORK SYSTEM")
    run_title.font.name = 'Arial'
    run_title.font.size = Pt(16)
    run_title.font.bold = True
    run_title.font.rtl = False
    
    add_line("="*50, bold=True)
    
    add_line(f"1. Geometry & Inputs:", bold=True)
    add_line(f"- Total Inclined Soldier Length = {sys_data['L_tot']:.2f} m")
    add_line(f"- Inclination Angle = {sys_data['angle']:.1f} degrees")
    add_line(f"- Applied Loads = Variable (Refer to Diagram)")
    
    doc.add_paragraph()
    add_line(f"2. Safety Checks:", bold=True)
    
    inc_sec = sys_data['inc_sec']
    add_line(f"A. Inclined Soldier ({inc_sec})", bold=True)
    inc_db = SECTIONS_DB.get(inc_sec, {'Mall': 13.1, 'Qall': 100.8})
    add_check("Moment", "M_max", sys_data.get('max_M_inc', 0), inc_db.get('Mall', 999), "kN.m")
    add_check("Shear", "V_max", sys_data.get('max_V_inc', 0), inc_db.get('Qall', 999), "kN")
    add_check("Deflection", "Local Relative", sys_data.get('max_def_inc', 0), sys_data.get('allw_def_inc', 10), "mm")
    
    base_sec = sys_data['base_sec']
    if base_sec != "None (Direct to Ground)":
        add_line(f"B. Horizontal Base Soldier ({base_sec})", bold=True)
        base_db = SECTIONS_DB.get(base_sec, {'Mall': 13.1, 'Qall': 100.8})
        add_check("Moment", "M_max", sys_data.get('max_M_base', 0), base_db.get('Mall', 999), "kN.m")
        add_check("Shear", "V_max", sys_data.get('max_V_base', 0), base_db.get('Qall', 999), "kN")
        add_check("Deflection", "Local Relative", sys_data.get('max_def_base', 0), sys_data.get('allw_def_base', 10), "mm")
    
    add_line("C. Push-Pull Struts (Axial Force)", bold=True)
    for idx, st_val in enumerate(sys_data.get('struts_res', [])):
        st_data = STRUTS_DB.get(st_val['type'], {})
        allow = st_data.get('allow', st_data.get('pts', {0: 50.0}).get(list(st_data.get('pts', {0:50.0}).keys())[0], 50.0))
        add_check(f"Strut {idx+1} ({st_val['type']})", "N_max", st_val['N'], allow, "kN")
        
    doc.add_page_break()
    add_line("3. Analysis Diagrams:", bold=True)
    
    titles = {
        'Load': "Assigned Load Diagram",
        'React': "Reactions Diagram (kN)",
        'N': "Axial Force Diagram (kN)",
        'V': "Shear Force Diagram (kN)",
        'M': "Bending Moment Diagram (kN.m)"
    }
    
    for key in ['Load', 'React', 'N', 'V', 'M']:
        buf = sys_data['img_bufs'][key]
        buf.seek(0)
        
        p_img = doc.add_paragraph()
        force_ltr_left(p_img)
        p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_img.add_run().add_picture(io.BytesIO(buf.read()), width=Cm(14.0))
        
        p_txt = doc.add_paragraph()
        force_ltr_left(p_txt)
        p_txt.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r_txt = p_txt.add_run(titles[key])
        r_txt.font.name = 'Arial'
        r_txt.font.size = Pt(11)
        r_txt.underline = True
        r_txt.font.rtl = False
        
        doc.add_page_break()
    
    out = io.BytesIO()
    doc.save(out)
    return out

# =========================================================
# 5. Main UI Module for Inclined Systems
# =========================================================
def render_inclined_module():
    st.markdown("## 📐 Inclined Formwork System (Advanced FEA)")
    
    if 'inclined_solved' not in st.session_state:
        st.session_state.inclined_solved = False
        
    st.markdown("#### ⚙️ 1. Geometry & Profiles")
    
    c_tot1, c_tot2 = st.columns(2)
    L_tot_val = c_tot1.number_input("Total Inclined Length (m)", value=5.0, step=0.5, on_change=lambda: st.session_state.update(inclined_solved=False))
    X_tot_val = c_tot2.number_input("Total Base Length (m)", value=3.5, step=0.5, on_change=lambda: st.session_state.update(inclined_solved=False))
    
    c_p1, c_p2 = st.columns(2)
    sec_list = list(SECTIONS_DB.keys()) if SECTIONS_DB else ["Soldier U100"]
    base_sec_list = ["Soldier U100", "None (Direct to Ground)"]
    default_idx = next((i for i, sec in enumerate(sec_list) if 'Soldier' in sec), 0)
    inc_sec = c_p1.selectbox("Profile (Inclined Soldier)", sec_list, index=default_idx, on_change=lambda: st.session_state.update(inclined_solved=False))
    base_sec = c_p2.selectbox("Profile (Base Soldier)", base_sec_list, index=0, on_change=lambda: st.session_state.update(inclined_solved=False))
    
    st.markdown("#### 🔗 2. Segments & Connections")
    c_g1, c_g2, c_g3, c_g4 = st.columns(4)
    angle_deg = c_g1.number_input("Inclination Angle (°)", value=60.0, step=5.0, on_change=lambda: st.session_state.update(inclined_solved=False))
    angle_rad = np.radians(angle_deg)
    num_struts = c_g2.number_input("Push-Pulls Count", min_value=1, max_value=5, value=2, step=1, on_change=lambda: st.session_state.update(inclined_solved=False))
    
    L_segs, X_segs, strut_types = [], [], []
    L_cum, X_cum = 0.0, 0.0
    
    for j in range(int(num_struts)):
        cl1, cl2, cl3 = st.columns([1, 1, 1.5])
        l_val = cl1.number_input(f"L{j+1} on Inclined (m)", value=2.0, step=0.5, key=f"L_{j}", on_change=lambda: st.session_state.update(inclined_solved=False))
        x_val = cl2.number_input(f"X{j+1} on Base (m)", value=1.5, step=0.5, key=f"X_{j}", on_change=lambda: st.session_state.update(inclined_solved=False))
        
        L_segs.append(l_val)
        X_segs.append(x_val)
        L_cum += l_val
        X_cum += x_val
        
        req_len = np.hypot(X_cum - L_cum * np.cos(angle_rad), 0 - L_cum * np.sin(angle_rad))
        valid_struts = get_valid_struts(req_len, STRUTS_DB)
        st_type = cl3.selectbox(f"Strut {j+1} (Req: {req_len:.2f}m)", valid_struts, key=f"st_{j}", on_change=lambda: st.session_state.update(inclined_solved=False))
        strut_types.append(st_type)
        
    L_rem = max(0.0, L_tot_val - sum(L_segs))
    X_rem = max(0.0, X_tot_val - sum(X_segs))
    st.info(f"📏 **Calculated Cantilevers:** Top Cantilever = {L_rem:.2f} m | Right Cantilever = {X_rem:.2f} m")
    
    st.markdown("#### ⚓ 3. Ground Supports Configuration")
    c_s1, c_s2, c_s3, c_s4 = st.columns(4)
    corner_check_opt = c_s1.selectbox("Corner Securing", ["2 Tie Rods", "Shoring System", "None (On Ground)"], key="corn_chk_opt", on_change=lambda: st.session_state.update(inclined_solved=False))
    sys_opts = ["None (On Ground directly)", "Acrow Prop", "Cup-lock", "Ringlock", "Shorebrace Frame"]
    shoring_sys = c_s2.selectbox("Base Shoring", sys_opts, on_change=lambda: st.session_state.update(inclined_solved=False))
    corner_type = c_s3.selectbox("Corner Support Type", ["Hinged", "Roller", "Fixed"], key="corn_type", on_change=lambda: st.session_state.update(inclined_solved=False))
    corner_ang = c_s4.number_input("Corner Angle (°)", value=0.0, step=15.0, key="corn_ang", on_change=lambda: st.session_state.update(inclined_solved=False))
    
    corner_sup = {'type': corner_type, 'angle': corner_ang}
    corner_tr_cap = 180.0
    
    allw_sh = 9999.0
    shoring_desc = "On Ground directly"
    
    if shoring_sys == "Acrow Prop":
        cp1, cp2, cp3, cp4 = st.columns(4)
        prop_keys = list(PROP_DB.keys()) if PROP_DB else ["AEP E-450"]
        prop_type = cp1.selectbox("Prop Type", prop_keys, key="sh_pt", on_change=lambda: st.session_state.update(inclined_solved=False))
        prop_ext = cp2.number_input("Extension (m)", value=3.0, step=0.1, key="sh_pe", on_change=lambda: st.session_state.update(inclined_solved=False))
        allw_sh = get_shoring_capacity("Acrow Prop", prop_type, 1.5, prop_ext)
        shoring_desc = f"{shoring_sys} ({prop_type}, Ext: {prop_ext}m)"
    elif shoring_sys == "Cup-lock":
        cp1, cp2, cp3, cp4 = st.columns(4)
        cu_grade = cp1.selectbox("Grade", list(CUPLOCK_DB.keys()) if CUPLOCK_DB else ["S355 (st.52)"], key="sh_cg", on_change=lambda: st.session_state.update(inclined_solved=False))
        
        cu_opts = list(CUPLOCK_DB[cu_grade].keys()) if CUPLOCK_DB and cu_grade in CUPLOCK_DB else [1.0, 1.5, 2.0]
        def_idx = cu_opts.index(1.5) if 1.5 in cu_opts else 0
        cu_unb = cp2.selectbox("Unbraced (m)", cu_opts, index=def_idx, key="sh_cu", format_func=lambda x: f"{float(x):.2f}", on_change=lambda: st.session_state.update(inclined_solved=False))
        
        allw_sh = get_shoring_capacity("Cup-lock", cu_grade, cu_unb, 3.0)
        shoring_desc = f"{shoring_sys} ({cu_grade}, Unb: {cu_unb}m)"
    elif shoring_sys == "Ringlock":
        cp1, cp2, cp3, cp4 = st.columns(4)
        ri_size = cp1.selectbox("Size", list(RINGLOCK_DB.keys()) if RINGLOCK_DB else ["Ringlock 1.5\""], key="sh_rs", on_change=lambda: st.session_state.update(inclined_solved=False))
        
        ri_opts = list(RINGLOCK_DB[ri_size].keys()) if RINGLOCK_DB and ri_size in RINGLOCK_DB else [1.0, 1.5, 2.0]
        def_idx = ri_opts.index(1.5) if 1.5 in ri_opts else 0
        ri_unb = cp2.selectbox("Unbraced (m)", ri_opts, index=def_idx, key="sh_ru", format_func=lambda x: f"{float(x):.2f}", on_change=lambda: st.session_state.update(inclined_solved=False))
        
        allw_sh = get_shoring_capacity("Ringlock", ri_size, ri_unb, 3.0)
        shoring_desc = f"{shoring_sys} ({ri_size}, Unb: {ri_unb}m)"
    elif shoring_sys == "Shorebrace Frame":
        allw_sh = 54.0
        shoring_desc = "Shorebrace Frame"

    num_base_sups = st.number_input("Additional Base Supports Count", 0, 10, int(num_struts), on_change=lambda: st.session_state.update(inclined_solved=False))
    base_sups = []
    default_xs = [sum(X_segs[:i+1]) for i in range(len(X_segs))]
    
    for i in range(int(num_base_sups)):
        cs1, cs2, cs3 = st.columns(3)
        def_x = default_xs[i] if i < len(default_xs) else float((i+1)*1.5)
        sx = cs1.number_input(f"Sup {i+1} X (m)", value=def_x, step=0.5, key=f"sx_{i}", on_change=lambda: st.session_state.update(inclined_solved=False))
        stype = cs2.selectbox(f"Sup {i+1} Type", ["Hinged", "Roller", "Fixed"], key=f"stype_{i}", on_change=lambda: st.session_state.update(inclined_solved=False))
        base_sups.append({'x': sx, 'type': stype, 'angle': 0.0})

    c_in, c_plot = st.columns([1.3, 1])
    with c_in:
        st.markdown("#### 🎯 4. Applied Loads")
        c_ld1, c_ld2, c_ld3 = st.columns(3)
        l_type = c_ld1.selectbox("Load Type", ["Uniform", "Trapezoidal/Triangular", "Point Load"], on_change=lambda: st.session_state.update(inclined_solved=False))
        num_items = c_ld2.number_input(f"Count of Loads", 1, 20, 1, on_change=lambda: st.session_state.update(inclined_solved=False))
        ldir = c_ld3.selectbox("Direction", ["Gravity (Vertical ↓)", "Perpendicular (Local ↘)"], on_change=lambda: st.session_state.update(inclined_solved=False))
        
        applied_loads = []
        for item in range(int(num_items)):
            if l_type == "Point Load":
                c1, c2 = st.columns(2)
                start_l = c1.number_input(f"P{item+1} Distance from Corner (m)", value=0.0, step=0.5, key=f"ls_{item}", on_change=lambda: st.session_state.update(inclined_solved=False))
                w1 = c2.number_input(f"P{item+1} Value (kN)", value=15.0, step=1.0, key=f"w1_{item}", on_change=lambda: st.session_state.update(inclined_solved=False))
                applied_loads.append({'type': l_type, 'start': start_l, 'end': start_l, 'w1': w1, 'w2': w1, 'dir': ldir})
            elif l_type == "Uniform":
                c1, c2, c3 = st.columns(3)
                start_l = c1.number_input(f"L{item+1} Start from Corner (m)", value=0.0, step=0.5, key=f"ls_{item}", on_change=lambda: st.session_state.update(inclined_solved=False))
                end_l = c2.number_input(f"L{item+1} End from Corner (m)", value=L_tot_val, step=0.5, key=f"le_{item}", on_change=lambda: st.session_state.update(inclined_solved=False))
                w1 = c3.number_input(f"L{item+1} W (kN/m)", value=15.0, step=1.0, key=f"w1_{item}", on_change=lambda: st.session_state.update(inclined_solved=False))
                applied_loads.append({'type': l_type, 'start': min(start_l, end_l), 'end': max(start_l, end_l), 'w1': w1, 'w2': w1, 'dir': ldir})
            else:
                c1, c2, c3, c4 = st.columns(4)
                start_l = c1.number_input(f"L{item+1} Start from Corner (m)", value=0.0, step=0.5, key=f"ls_{item}", on_change=lambda: st.session_state.update(inclined_solved=False))
                end_l = c2.number_input(f"L{item+1} End from Corner (m)", value=L_tot_val, step=0.5, key=f"le_{item}", on_change=lambda: st.session_state.update(inclined_solved=False))
                w1 = c3.number_input(f"L{item+1} W1 (kN/m)", value=15.0, step=1.0, key=f"w1_{item}", on_change=lambda: st.session_state.update(inclined_solved=False))
                w2 = c4.number_input(f"L{item+1} W2 (kN/m)", value=0.0, step=1.0, key=f"w2_{item}", on_change=lambda: st.session_state.update(inclined_solved=False))
                applied_loads.append({'type': l_type, 'start': min(start_l, end_l), 'end': max(start_l, end_l), 'w1': w1, 'w2': w2, 'dir': ldir})

    nodes, elements, nodal_loads, L_tot, X_tot, display_nodes, supports_list = build_fea_mesh(L_segs, L_rem, X_segs, X_rem, angle_rad, applied_loads, inc_sec, base_sec, strut_types, corner_sup, base_sups)

    with c_plot:
        st.markdown("<h4 style='text-align: center; font-family: Arial; font-weight: normal; border-bottom: 1px solid gray; padding-bottom: 5px;'>Live Assigned Loads</h4>", unsafe_allow_html=True)
        live_img_buf = plot_live_geometry(nodes, elements, applied_loads, L_segs, X_segs, angle_deg, inc_sec, base_sec, L_tot, X_tot, supports_list)
        st.image(live_img_buf, use_container_width=True)

    st.markdown("---")
    
    col_btn, col_chk, col_blank = st.columns([1.5, 1.5, 1])
    with col_btn:
        if st.button("🚀 Run Advanced FEA & Generate Report", type="primary", use_container_width=True):
            with st.spinner("Building Matrix & Solving FEA..."):
                U, R = solve_fea_engine(nodes, elements, nodal_loads, supports_list)
                
                max_M_inc, max_V_inc, max_M_base, max_V_base = 0, 0, 0, 0
                max_def_inc, max_def_base = 0.0, 0.0
                max_span_inc, max_span_base = 0.0, 0.0
                struts_results = []
                
                for el in elements:
                    if el['type'] == 'frame':
                        max_M = max(abs(el['internal']['M'][0]), abs(el['internal']['M'][-1]))
                        if len(el['internal']['M']) > 2: max_M = max(max_M, np.max(np.abs(el['internal']['M'])))
                        max_V = max(abs(el['internal']['V'][0]), abs(el['internal']['V'][-1]))
                        if len(el['internal']['V']) > 2: max_V = max(max_V, np.max(np.abs(el['internal']['V'])))
                        
                        max_def_el = np.max(np.abs(el['internal']['v_rel'])) * 1000.0
                        
                        if el['group'] == 'inclined': 
                            max_M_inc = max(max_M_inc, max_M)
                            max_V_inc = max(max_V_inc, max_V)
                            max_def_inc = max(max_def_inc, max_def_el)
                            max_span_inc = max(max_span_inc, el['L'])
                        elif el['group'] == 'base': 
                            max_M_base = max(max_M_base, max_M)
                            max_V_base = max(max_V_base, max_V)
                            max_def_base = max(max_def_base, max_def_el)
                            max_span_base = max(max_span_base, el['L'])
                    elif el['type'] == 'truss':
                        struts_results.append({'type': el['sec'], 'N': abs(el['internal']['N'][0])})

                allw_def_inc = (max_span_inc * 1000) / 400.0 if max_span_inc > 0 else 10.0
                allw_def_base = (max_span_base * 1000) / 400.0 if max_span_base > 0 else 10.0

                st.session_state.inclined_fea_data = {
                    'U': U, 'R': R, 'nodes': nodes, 'elements': elements, 'display_nodes': display_nodes, 'supports_list': supports_list,
                    'sys_data': {
                        'L_tot': L_tot, 'X_tot': X_tot, 'angle': angle_deg, 'W': "Variable", 'ld_dir': "Variable",
                        'inc_sec': inc_sec, 'base_sec': base_sec, 
                        'corner_check_opt': corner_check_opt, 'corner_tr_cap': corner_tr_cap, 
                        'shoring_type': shoring_desc, 'allw_sh': allw_sh,
                        'max_M_inc': max_M_inc, 'max_V_inc': max_V_inc, 'max_def_inc': max_def_inc, 'allw_def_inc': allw_def_inc,
                        'max_M_base': max_M_base, 'max_V_base': max_V_base, 'max_def_base': max_def_base, 'allw_def_base': allw_def_base,
                        'struts_res': struts_results
                    }
                }
                st.session_state.inclined_solved = True
                st.session_state.show_safety_table = False
                
    with col_chk:
        if st.button("📊 Quick Safety Check Table", use_container_width=True):
            if st.session_state.inclined_solved:
                st.session_state.show_safety_table = not st.session_state.get('show_safety_table', False)
            else:
                st.warning("⚠️ Please run the FEA analysis first!")
    
    if getattr(st.session_state, 'show_safety_table', False) and st.session_state.inclined_solved:
        fea_data = st.session_state.inclined_fea_data
        sd = fea_data['sys_data']
        struts_results = sd['struts_res']
        
        html = """
        <style>
        .safety-table { width: 100%; border-collapse: collapse; margin: 15px 0; font-size: 14px; font-family: 'Arial', sans-serif; box-shadow: 0 0 15px rgba(0, 0, 0, 0.05); border-radius: 8px 8px 0 0; overflow: hidden; }
        .safety-table thead tr { background-color: #2c3e50; color: #ffffff; text-align: center; font-weight: bold; }
        .safety-table th, .safety-table td { padding: 12px 15px; text-align: center; border-bottom: 1px solid #dddddd; }
        .safety-table tbody tr:hover { background-color: #f5f5f5; }
        .badge-safe { background-color: #d4edda; color: #155724; padding: 4px 10px; border-radius: 12px; font-weight: bold; border: 1px solid #c3e6cb; }
        .badge-unsafe { background-color: #f8d7da; color: #721c24; padding: 4px 10px; border-radius: 12px; font-weight: bold; border: 1px solid #f5c6cb; }
        </style>
        <table class="safety-table">
        <thead><tr><th>Component</th><th>Description</th><th>Actual (Applied)</th><th>Allowable (Capacity)</th><th>Status</th></tr></thead>
        <tbody>
        """
        
        def add_row(comp, desc, act, allw, unit):
            if allw > 9000: 
                status_html = "<span style='color:gray;'>Not Required</span>"
                allw_str = "No Limit"
            else:
                is_safe = act <= allw
                status_html = "<span class='badge-safe'>✅ SAFE</span>" if is_safe else "<span class='badge-unsafe'>❌ UNSAFE</span>"
                allw_str = f"{allw:.2f} {unit}"
            return f"<tr><td><b>{comp}</b></td><td>{desc}</td><td>{act:.2f} {unit}</td><td>{allw_str}</td><td>{status_html}</td></tr>"

        inc_db = SECTIONS_DB.get(sd['inc_sec'], {'Mall': 13.1, 'Qall': 100.8})
        base_db = SECTIONS_DB.get(sd['base_sec'], {'Mall': 13.1, 'Qall': 100.8})

        html += add_row("Inclined Soldier", f"{sd['inc_sec']} - Moment", sd.get('max_M_inc', 0), inc_db.get('Mall', 999), "kN.m")
        html += add_row("Inclined Soldier", f"{sd['inc_sec']} - Shear", sd.get('max_V_inc', 0), inc_db.get('Qall', 999), "kN")
        html += add_row("Inclined Soldier", f"{sd['inc_sec']} - Deflection", sd.get('max_def_inc', 0), sd.get('allw_def_inc', 10), "mm")
        
        if sd['base_sec'] != "None (Direct to Ground)":
            html += add_row("Base Soldier", f"{sd['base_sec']} - Moment", sd.get('max_M_base', 0), base_db.get('Mall', 999), "kN.m")
            html += add_row("Base Soldier", f"{sd['base_sec']} - Shear", sd.get('max_V_base', 0), base_db.get('Qall', 999), "kN")
            html += add_row("Base Soldier", f"{sd['base_sec']} - Deflection", sd.get('max_def_base', 0), sd.get('allw_def_base', 10), "mm")
        
        for i, st_res in enumerate(struts_results):
            st_data = STRUTS_DB.get(st_res['type'], {})
            allow = st_data.get('allow', st_data.get('pts', {0: 50.0}).get(list(st_data.get('pts', {0:50.0}).keys())[0], 50.0))
            html += add_row(f"Push-Pull Strut {i+1}", st_res['type'], st_res['N'], allow, "kN")
            
        R = fea_data['R']
        for i, sup in enumerate(fea_data['supports_list']):
            n = sup['node']
            if i == 0:
                chk_opt = sd.get('corner_check_opt', "2 Tie Rods")
                if chk_opt == "2 Tie Rods":
                    R_res = np.hypot(R[3*n], R[3*n+1])
                    html += add_row("Corner Support", "2 x Tie Rods (Resultant Force)", R_res, sd['corner_tr_cap'], "kN")
                elif chk_opt == "Shoring System":
                    Ry_abs = abs(R[3*n+1])
                    html += add_row("Corner Support", f"{sd['shoring_type']} (Vertical)", Ry_abs, sd['allw_sh'], "kN")
                else:
                    R_res = np.hypot(R[3*n], R[3*n+1])
                    html += add_row("Corner Support", "None (On Ground)", R_res, 9999.0, "kN")
            else:
                Ry_abs = abs(R[3*n+1])
                html += add_row(f"Base Support {i}", sd['shoring_type'], Ry_abs, sd['allw_sh'], "kN")

        html += "</tbody></table>"
        st.markdown("### 📊 System Components Safety Summary")
        st.markdown(html, unsafe_allow_html=True)
    
    if st.session_state.inclined_solved:
        fea_data = st.session_state.inclined_fea_data
        
        st.markdown("### 🎛️ Analysis Results & Diagrams")
        with st.expander("⚙️ Diagram Scale Controls", expanded=True):
            c_s1, c_s2, c_s3 = st.columns(3)
            sc_n = c_s1.slider("Axial Force Scale", 0.001, 0.100, 0.02, step=0.005)
            sc_v = c_s2.slider("Shear Force Scale", 0.001, 0.100, 0.02, step=0.005)
            sc_m = c_s3.slider("Moment Scale", 0.01, 0.50, 0.10, step=0.01)
            scales = {'N': sc_n, 'V': sc_v, 'M': sc_m}
            
        img_bufs = plot_sap2000_diagrams(fea_data['nodes'], fea_data['elements'], fea_data['R'], scales, fea_data['display_nodes'], applied_loads, angle_deg, fea_data['sys_data']['L_tot'], fea_data['sys_data']['X_tot'], fea_data['sys_data']['inc_sec'], fea_data['sys_data']['base_sec'], fea_data['supports_list'])
        
        titles = {
            'Load': "Assigned Load Diagram",
            'React': "Reactions Diagram (kN)",
            'N': "Axial Force Diagram (kN)",
            'V': "Shear Force Diagram (kN)",
            'M': "Bending Moment Diagram (kN.m)"
        }
        
        c_p1, c_p2, c_p3 = st.columns(3)
        cols = [c_p1, c_p2, c_p3, c_p1, c_p2]
        
        for idx, key in enumerate(['Load', 'React', 'N', 'V', 'M']):
            with cols[idx]:
                st.image(img_bufs[key], use_container_width=True)
                st.markdown(f"<p style='text-align: center; border-bottom: 1px solid gray; padding-bottom: 5px; font-family: Arial; font-size: 14px;'>{titles[key]}</p>", unsafe_allow_html=True)
        
        fea_data['sys_data'].update({'img_bufs': img_bufs})
        
        docx_out = generate_inclined_report(fea_data['sys_data'])
        
        st.success("✅ SAP2000-Style Analysis Complete!")
        st.download_button("⬇️ Download Inclined System Calculation Sheet", 
                           data=docx_out.getvalue(), 
                           file_name="Inclined_System_Report.docx", 
                           mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
