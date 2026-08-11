# app.py

import os
import io
import gc
import re
from datetime import date
import pandas as pd
import numpy as np
import streamlit as st

# ==========================================
# 1. Page Configuration & Setup
# ==========================================
st.set_page_config(layout="wide", page_title="Acrow - Pro 3-Moment Solver")

import matplotlib
matplotlib.use('Agg') # 💡 الحل الجذري لإغلاق واجهات الرسم وحماية السيرفر من الانهيار

from docx import Document
from docx.shared import Cm, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_COLOR_INDEX
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import fitz
from PIL import Image

st.markdown(
    """
    <style>
    .block-container { max-width: 95% !important; padding-top: 2rem; padding-bottom: 2rem; } 
    header { visibility: hidden; } 
    footer { visibility: hidden; }
    </style>
    """, 
    unsafe_allow_html=True
)

import auth
from config import SECTIONS_DB, STRUTS_DB
from helpers import convert_transparent_to_pdf_stream, crop_white_margins
from math_solver import get_element_safety_details, solve_beam_advanced
from plot_core import generate_acrow_diagrams
from ui_project import render_project_details
from ui_standard import render_slab_element, render_vertical_element
from report_builder import *

# ==========================================
# 2. Authentication & Admin Dashboard
# ==========================================
if not auth.check_access(): 
    st.stop()

if not st.session_state.get("is_admin", False) and "welcome_msg_shown" not in st.session_state:
    st.info("💡 **IMPORTANT:** Bookmark this page now. Next time, open it from your bookmarks to login automatically!")
    st.session_state["welcome_msg_shown"] = True

auth.render_admin_dashboard()

with st.sidebar:
    st.markdown("<h1 style='color: #E10000; font-weight: bold; letter-spacing: 2px;'>ACROW</h1>", unsafe_allow_html=True)
    st.markdown("---")
    if st.session_state.get("is_admin", False): 
        st.markdown(f"### 🟢 Active Users: **{auth.track_active_users()}**")
    else: 
        st.markdown(f"#### Welcome {st.session_state.get('user_name', 'User')} 👋")

# ==========================================
# 3. Project Details UI
# ==========================================
proj_name, contractor, calc_sub, sys_name, proj_no, calc_by, date_val, chk_by, ref_code, cover_img, data_sheets, def_sec, def_main = render_project_details()

# ==========================================
# SMART INITIALS LOGIC 
# ==========================================
user_email = st.query_params.get("user", "")

if not user_email:
    if st.query_params.get("admin") == "acrow_master":
        user_email = "islam.morsy@acrow.co"
    elif hasattr(st, "experimental_user") and st.experimental_user.email:
        user_email = st.experimental_user.email

if user_email and '@' in user_email:
    name_part = user_email.split('@')[0]
    parts = name_part.split('.')
    if len(parts) >= 2:
        calc_by = f"Eng. {parts[0][0].upper()}.{parts[1][0].upper()}"
    elif len(parts) == 1:
        calc_by = f"Eng. {parts[0][0].upper()}"
elif not calc_by or calc_by == "Eng. ":
    calc_by = "Eng."
# ==========================================

# الكود البريطاني = 1.50، والأمريكي = 2.40
def_live_load = 1.50 if "BS" in ref_code else 2.40

# ==========================================
# 4. Structural System Configurator
# ==========================================
st.divider()
st.subheader("2. Structural System Configurator")

# 💡 تمت إضافة خيار الكباري والأشكال المتقدمة للقائمة
sys_cat = st.radio("Select Structural Category:", ["Slab Elements", "Vertical Elements (Walls, Columns)", "Inclined Elements (Frames)", "Slab Back-propping", "Bridges", "Advanced Shapes (Curved/Multi)"], horizontal=True)

if "Inclined Elements" in sys_cat:
    import inclined_master
    inclined_master.render_inclined_module()
    st.stop()

if "Back-propping" in sys_cat:
    import backprop_master
    backprop_master.render_backprop_module(ref_code)
    st.stop()

# 💡 تمرير بيانات المشروع لملف الكباري لتوليد النوتة بشكل متطابق مع باقي البرنامج
if "Bridges" in sys_cat:
    import bridge_master
    proj_info = {
        "proj_name": proj_name, "contractor": contractor, "calc_sub": calc_sub,
        "sys_name": sys_name, "proj_no": proj_no, "calc_by": calc_by,
        "date_val": date_val, "chk_by": chk_by, "ref_code": ref_code,
        "cover_img": cover_img, "data_sheets": data_sheets
    }
    bridge_master.render_bridge_module(proj_info)
    st.stop()

# 💡 استدعاء ملف الأشكال المتقدمة
if "Advanced Shapes" in sys_cat:
    import advanced_shape_master
    advanced_shape_master.render_advanced_shape_module()
    st.stop()

configs = []

if "Slab Elements" in sys_cat:
    cg1, cg2, cg3, cg4 = st.columns(4)
    gamma_c = cg1.number_input("Conc. Density (kN/m³)", value=25.0)
    live_load = cg2.number_input("Live Load (kN/m²)", value=float(def_live_load), step=0.05)
    fw_load = cg3.number_input("Formwork Load (kN/m²)", value=0.50, step=0.05)
    num_elements = cg4.number_input("Number of Elements", min_value=1, max_value=5, value=1)
    
    tabs = st.tabs([f"Element {i+1}" for i in range(int(num_elements))])
    for i, tab in enumerate(tabs):
        with tab:
            conf = render_slab_element(i, gamma_c, live_load, fw_load, def_sec, def_main)
            conf['gamma_c'] = gamma_c
            conf['live_load'] = live_load
            conf['fw_load'] = fw_load
            configs.append(conf)
else:
    element_subtype = st.radio("Element Type:", ["Wall", "Column"], horizontal=True)
    num_elements = st.number_input(f"Number of {element_subtype}s", min_value=1, max_value=5, value=1)
    
    tabs = st.tabs([f"{element_subtype} {i+1}" for i in range(int(num_elements))])
    for i, tab in enumerate(tabs):
        with tab:
            conf = render_vertical_element(i, element_subtype, def_sec, def_main)
            configs.append(conf)

# ==========================================
# 5. Global Buttons & Action Handlers
# ==========================================
st.divider()
st.warning("⚠️ **تنبيه:** يرجى مراجعة وتأكيد تطابق البيانات الأساسية مع العناصر المدخلة قبل استخراج النوتة الحسابية.")

col_btn1, col_btn2 = st.columns(2)
with col_btn1: 
    check_safety_btn = st.button("🔍 Pre-Check Safety", use_container_width=True)
with col_btn2: 
    generate_doc_btn = st.button("🚀 Generate Automated Calculation Sheet", type="primary", use_container_width=True)

if check_safety_btn:
    errors = [c for c in configs if not c.get('is_panel_system') and c.get('cat') == 'horizontal' and (c.get("s_cr", 0) < -0.01 or c.get("m_cr", 0) < -0.01)]
    if errors: 
        st.error("❌ Error: Spans exceed total length. Please fix geometry first.")
    else:
        st.markdown("### 🛡️ Pre-Check Safety Results")
        all_safe = perform_global_safety_check(configs) 
        has_valid_configs = False
        
        for idx, conf in enumerate(configs):
            has_valid_configs = True
            if conf['cat'] == 'horizontal':
                header_val = f"Beam {conf['beam_b']}x{conf['ts']}m" if conf['sub_cat'] == 'Beam' else f"{conf['sub_cat']} {conf['ts']}m"
            else:
                header_val = f"{conf['sub_cat']} {conf.get('height')}m"
                
            st.markdown(f"**Section {idx+1} Summary ({header_val}):**")
            df_data = []
            
            if conf.get('is_panel_system'):
                panel_safe = conf['w'] <= conf['panel_allowable']
                df_data.append({"Component": f"{conf['sys_name']} Panel", "Moment/Z": "-", "Shear": "-", "Deflection": "-", "Support/Reaction": f"{'SAFE' if panel_safe else 'UNSAFE'} {conf['w']:.2f} < {conf['panel_allowable']:.2f}", "Status": "SAFE" if panel_safe else "UNSAFE"})
                if conf.get('tie_h') and conf['sys_name'] != "Circular Steel Panel System" and not conf.get('strongback', {}).get('active'):
                    tie_load = conf['w'] * conf.get('tie_h',0) * conf.get('tie_v',0)
                    tie_safe = tie_load <= 90.0
                    df_data.append({"Component": "Tie Rod 15mm", "Moment/Z": "-", "Shear": "-", "Deflection": "-", "Support/Reaction": f"{'SAFE' if tie_safe else 'UNSAFE'} {tie_load:.2f} < 90.00", "Status": "SAFE" if tie_safe else "UNSAFE"})
                if conf.get('bolt_h'):
                    bolt_load = conf['w'] * conf['bolt_h'] * conf['bolt_v']
                    bolt_safe = bolt_load <= 50.0
                    df_data.append({"Component": "Acrow Bolt", "Moment/Z": "-", "Shear": "-", "Deflection": "-", "Support/Reaction": f"{'SAFE' if bolt_safe else 'UNSAFE'} {bolt_load:.2f} < 50.00", "Status": "SAFE" if bolt_safe else "UNSAFE"})
            else:
                M_ply = (conf['w'] * (conf['s_spc']**2)) / 10
                Z_req = (M_ply * 100) / 3.41
                E_ply, I_ply = (74.52, 48.60) if "18" in conf['ply_thick'] else (74.52, 77.0)
                D_ply = (0.0068 * conf['w'] * (conf['s_spc']*100)**4) / (100 * E_ply * I_ply)
                all_ply_d = (conf['s_spc']*1000)/300
                ply_m_safe = Z_req <= conf['ply_mall']
                ply_d_safe = D_ply <= all_ply_d
                
                df_data.append({"Component": "Plywood", "Moment/Z": f"{'SAFE' if ply_m_safe else 'UNSAFE'} {Z_req:.2f} < {conf['ply_mall']:.2f}", "Shear": "-", "Deflection": f"{'SAFE' if ply_d_safe else 'UNSAFE'} {D_ply:.2f} < {all_ply_d:.2f}", "Support/Reaction": "-", "Status": "SAFE" if (ply_m_safe and ply_d_safe) else "UNSAFE"})
                
                s_act_M, s_act_V, s_act_D, s_all_M, s_all_V, s_all_D = get_element_safety_details(conf, True)
                df_data.append({"Component": f"Secondary ({conf['s_sec']})", "Moment/Z": f"{'SAFE' if s_act_M<=s_all_M else 'UNSAFE'} {s_act_M:.2f} < {s_all_M:.2f}", "Shear": f"{'SAFE' if s_act_V<=s_all_V else 'UNSAFE'} {s_act_V:.2f} < {s_all_V:.2f}", "Deflection": f"{'SAFE' if s_act_D<=s_all_D else 'UNSAFE'} {s_act_D:.2f} < {s_all_D:.2f}", "Support/Reaction": "-", "Status": "SAFE" if (s_act_M<=s_all_M and s_act_V<=s_all_V and s_act_D<=s_all_D) else "UNSAFE"})
                
                m_act_M, m_act_V, m_act_D, m_all_M, m_all_V, m_all_D = get_element_safety_details(conf, False)
                df_data.append({"Component": f"Main ({conf['m_sec']})", "Moment/Z": f"{'SAFE' if m_act_M<=m_all_M else 'UNSAFE'} {m_act_M:.2f} < {m_all_M:.2f}", "Shear": f"{'SAFE' if m_act_V<=m_all_V else 'UNSAFE'} {m_act_V:.2f} < {m_all_V:.2f}", "Deflection": f"{'SAFE' if m_act_D<=m_all_D else 'UNSAFE'} {m_act_D:.2f} < {m_all_D:.2f}", "Support/Reaction": "-", "Status": "SAFE" if (m_act_M<=m_all_M and m_act_V<=m_all_V and m_act_D<=m_all_D) else "UNSAFE"})
                
                if conf.get('t_allow') is not None and conf['t_allow'] < 900:
                    prop_m = SECTIONS_DB[conf['m_sec']]
                    _, _, _, _, R_m = solve_beam_advanced(conf['m_L'], conf['m_sup'], conf['m_ld'], prop_m['E'], prop_m['I'])
                    max_R = np.max(R_m) if len(R_m) > 0 else 0
                    t_safe = max_R <= conf['t_allow']
                    
                    support_comp_name = conf['t_name']
                    if conf['t_name'] == 'Acrow Prop' and conf.get('t_sub'):
                        support_comp_name = f"Support ({conf['t_sub']})"
                    elif conf['t_name'] in ['Cup-lock', 'Ring-lock'] and conf.get('t_sub'):
                        support_comp_name = f"Support ({conf['t_sub']}) [Unbraced={conf['t_unb']:.2f}m]"
                    else:
                        support_comp_name = f"Support ({conf['t_name']})"
                        
                    df_data.append({"Component": support_comp_name, "Moment/Z": "-", "Shear": "-", "Deflection": "-", "Support/Reaction": f"{'SAFE' if t_safe else 'UNSAFE'} {max_R:.2f} < {conf['t_allow']:.2f}", "Status": "SAFE" if t_safe else "UNSAFE"})

            if conf.get('strongback', {}).get('active'):
                sb = conf['strongback']
                sb_m_v_safe, sb_v_v_safe = sb['M_v'] <= SECTIONS_DB[sb['sv']]['Mall'], sb['V_v'] <= SECTIONS_DB[sb['sv']]['Qall']
                sb_m_h_safe, sb_v_h_safe = sb['M_h'] <= SECTIONS_DB[sb['sh']]['Mall'], sb['V_h'] <= SECTIONS_DB[sb['sh']]['Qall']
                sb_tie_safe, sb_waler_safe = sb.get('tie_T_single', 0.0) <= 90.0, sb['waler_M'] <= SECTIONS_DB[sb['waler_sec']]['Mall']
                sb_pin_safe = sb['max_diag_force'] <= 80.0
                
                df_data.append({"Component": f"Vert Soldier ({sb['sv']})", "Moment/Z": f"{'SAFE' if sb_m_v_safe else 'UNSAFE'}", "Shear": f"{'SAFE' if sb_v_v_safe else 'UNSAFE'}", "Deflection": "-", "Support/Reaction": "-", "Status": "SAFE" if (sb_m_v_safe and sb_v_v_safe) else "UNSAFE"})
                df_data.append({"Component": f"Horz Soldier ({sb['sh']})", "Moment/Z": f"{'SAFE' if sb_m_h_safe else 'UNSAFE'}", "Shear": f"{'SAFE' if sb_v_h_safe else 'UNSAFE'}", "Deflection": "-", "Support/Reaction": "-", "Status": "SAFE" if (sb_m_h_safe and sb_v_h_safe) else "UNSAFE"})
                df_data.append({"Component": "Tie Rod & Revit Pin", "Moment/Z": "-", "Shear": f"Pin: {'SAFE' if sb_pin_safe else 'UNSAFE'}", "Deflection": "-", "Support/Reaction": f"Tie: {'SAFE' if sb_tie_safe else 'UNSAFE'}", "Status": "SAFE" if (sb_pin_safe and sb_tie_safe) else "UNSAFE"})
                df_data.append({"Component": f"Waler ({sb['waler_sec']})", "Moment/Z": f"{'SAFE' if sb_waler_safe else 'UNSAFE'}", "Shear": "-", "Deflection": "-", "Support/Reaction": "-", "Status": "SAFE" if sb_waler_safe else "UNSAFE"})
                
                for d in sb['diags']:
                    d_force = next((abs(e['N_ax']) for e in sb['elements'] if e['type'] == 'truss' and e['sec'] == d['type'].split()[0]), 0)
                    df_data.append({"Component": f"Diagonal ({d['type'].split()[0]})", "Moment/Z": "-", "Shear": "-", "Deflection": "-", "Support/Reaction": f"{'SAFE' if d_force <= d['allow'] else 'UNSAFE'}", "Status": "SAFE" if d_force <= d['allow'] else "UNSAFE"})

            if conf.get('cat') == 'vertical' and conf.get('tilting', {}).get('active'):
                td = conf['tilting']
                tilt_safe = True
                if not td.get('length_safe', True): tilt_safe = False
                if 'struts' in td:
                    for st_c in td['struts']:
                        if abs(st_c['N']) > st_c['allow']: tilt_safe = False
                if td.get('max_n',0) > 80.0 or (max(td.get('ry1',0), td.get('ry2',0))/2) > 15.1 or (max(td.get('rx1',0), td.get('rx2',0))/2) > 29.5: tilt_safe = False
                
                df_data.append({"Component": "Tilting System", "Moment/Z": "-", "Shear": "-", "Deflection": "-", "Support/Reaction": "-", "Status": "SAFE" if tilt_safe else "UNSAFE"})
                if td.get('block_data') and td['block_data'].get('active'):
                    df_data.append({"Component": "Concrete Block", "Moment/Z": "-", "Shear": "-", "Deflection": "-", "Support/Reaction": "-", "Status": "SAFE" if td['block_data'].get('safe', False) else "UNSAFE"})
                    
            st.table(pd.DataFrame(df_data))
                
        if has_valid_configs and not all_safe: st.warning("⚠️ Please fix UNSAFE elements before generating the report.")
        elif not has_valid_configs: st.info("No calculable sections found. Please configure a valid system first.")

# ==========================================
# 6. Generate Document (The Output Engine)
# ==========================================
if generate_doc_btn:
    errors = [c for c in configs if not c.get('is_panel_system') and c.get('cat') == 'horizontal' and (c.get("s_cr", 0) < -0.01 or c.get("m_cr", 0) < -0.01)]
    tilt_len_errors = [c for c in configs if c.get('cat') == 'vertical' and c.get('tilting', {}).get('active') and not c.get('tilting', {}).get('length_safe', True)]
    
    if errors: 
        st.error("❌ Error: Cannot generate report! Spans exceed total length.")
    elif tilt_len_errors: 
        st.error("❌ Error: Cannot generate report! One or more Struts exceed their valid length range.")
    elif not os.path.exists("Acrow_Template.docx"): 
        st.error("❌ Template file 'Acrow_Template.docx' not found.")
    else:
        if not perform_global_safety_check(configs): 
            st.error("❌ **Report Generation Rejected:** One or more structural components are UNSAFE. Please review the 'Pre-Check Safety' results and fix the failing elements before generating the calculation sheet.")
            st.stop()
            
        with st.spinner("🔄 Running Advanced 3-Moment Equation Solver & Building Document..."):
            doc = Document("Acrow_Template.docx")
            
            file_sub_cat = configs[0]['sub_cat'] if (configs and 'sub_cat' in configs[0]) else ("Slab" if "Slab" in sys_cat else "Vertical")
            sys_name_clean = sys_name.replace("/", "-").replace("\\", "-")
            file_name = f"Calculation Sheet for {file_sub_cat} - Using {sys_name_clean}.docx"

            # ==========================================
            # Smart Find & Replace logic for Cover Page
            # ==========================================
            def remove_hardcoded_prefix(p):
                if p.text and "CALCULATION SHEET FOR" in p.text.upper():
                    for r in p.runs:
                        if "CALCULATION SHEET FOR" in r.text.upper():
                            r.text = re.sub(r'(?i)CALCULATION SHEET FOR\s*', '', r.text)
                            
                    if "CALCULATION SHEET FOR" in p.text.upper():
                        clean_text = re.sub(r'(?i)CALCULATION SHEET FOR\s*', '', p.text)
                        if p.runs:
                            font_name = p.runs[0].font.name
                            font_size = p.runs[0].font.size
                            font_bold = p.runs[0].font.bold
                            font_color = p.runs[0].font.color.rgb if p.runs[0].font.color else None
                            
                            for r in p.runs:
                                r.text = ""
                                
                            p.runs[0].text = clean_text
                            p.runs[0].font.name = font_name
                            p.runs[0].font.size = font_size
                            p.runs[0].font.bold = font_bold
                            if font_color:
                                p.runs[0].font.color.rgb = font_color
                        else:
                            p.text = clean_text

            for p in doc.paragraphs:
                remove_hardcoded_prefix(p)
                
            for tbl in doc.tables:
                for row in tbl.rows:
                    for cell in row.cells:
                        for p in cell.paragraphs:
                            remove_hardcoded_prefix(p)

            # 2. INJECT NEW VARIABLES
            replacements = {
                "[PROJECT_NAME]": proj_name,
                "[CONTRACTOR]": contractor,
                "[CALC_SUBJECT]": calc_sub,
                "[SYSTEM_NAME]": sys_name
            }
            
            for p in doc.paragraphs:
                if "[COVER_IMAGE]" in p.text:
                    for r in p.runs:
                        r.text = r.text.replace("[COVER_IMAGE]", "")
                    if cover_img and cover_img != "No images found." and os.path.exists(cover_img):
                        p.add_run().add_picture(cover_img, width=Cm(15.0))
                        
                for k, v in replacements.items():
                    if k in p.text:
                        for r in p.runs:
                            if k in r.text:
                                r.text = r.text.replace(k, str(v))
                        if k in p.text:  
                            p.text = p.text.replace(k, str(v))

            for tbl in doc.tables:
                for row in tbl.rows:
                    for cell in row.cells:
                        for p in cell.paragraphs:
                            for k, v in replacements.items():
                                if k in p.text:
                                    for r in p.runs:
                                        if k in r.text:
                                            r.text = r.text.replace(k, str(v))
                                    if k in p.text:
                                        p.text = p.text.replace(k, str(v))
                                        
            doc.add_page_break()
            
            # --- Headers and Footers ---
            for sec in doc.sections:
                for hf in [sec.header, sec.first_page_header, sec.footer, sec.first_page_footer]:
                    if hf:
                        for tbl in hf.tables:
                            for row in tbl.rows:
                                for cell in row.cells:
                                    for p in cell.paragraphs:
                                        for k, v in {"[PROJECT_NAME]": proj_name, "[CONTRACTOR]": contractor, "[PROJ_NO]": proj_no, "[DATE]": date_val, "[CALC_BY]": calc_by, "[CHK_BY]": chk_by, "[REV]": "00"}.items():
                                            if k in p.text:
                                                for r in p.runs:
                                                    if k in r.text:
                                                        r.text = r.text.replace(k, str(v))
                                                if k in p.text:
                                                    p.text = p.text.replace(k, str(v))
                                                for r in p.runs: 
                                                    r.font._element.set(qn('w:ascii'), 'Arial')
            
            # --- Regulations ---
            insert_blue_banner(doc, "REGULATIONS AND STANDARDS", font_size=16)
            doc.add_paragraph()
            
            if "BS" in ref_code: 
                add_eq(doc, "1- BS 5975-1996: FORMWORK FOR CONCRETE")
                add_eq(doc, "2- BS 5975-2008: FORMWORK FOR CONCRETE")
                add_eq(doc, "3- FORMWORK AGUIDE TO AGOOD PRATICE")
                add_eq(doc, "4- WISA®-FORM PLYWOOD.")
                add_eq(doc, "5- THE SAUDI BUILDING CODE (SBC) 2018")
            else: 
                add_eq(doc, "1- ACI 347R-14 ....... GUIDE TO FORMWORK FOR CONCRETE.")
                add_eq(doc, "2- ACI SP-4 ......... FORMWORK FOR CONCRETE.")
                add_eq(doc, "3- WISA®-FORM PLYWOOD.")
                add_eq(doc, "4- THE SAUDI BUILDING CODE (SBC) 2018")
            
            # --- Data Sheets ---
            if data_sheets:
                doc.add_page_break()
                insert_blue_banner(doc, "FORMWORK MATERIALS TECHNICAL DATA", font_size=14)
                for f in data_sheets:
                    if os.path.exists(f): 
                        append_pdf_stream_to_word(f, doc, is_path=True, max_width_cm=17.5, max_height_cm=24.0, add_border=True, reduce_first_page=True)
            
            # --- Slab Design Loads ---
            if "Slab Elements" in sys_cat:
                design_pdf = "Design_Loads_BS.pdf" if "BS" in ref_code and os.path.exists("Design_Loads_BS.pdf") else ("Design_Loads_ACI.pdf" if "ACI" in ref_code and os.path.exists("Design_Loads_ACI.pdf") else None)
                if design_pdf: 
                    doc.add_page_break()
                    insert_blue_banner(doc, "DESIGN LOADS FOR SLAB", font_size=14)
                    append_pdf_stream_to_word(design_pdf, doc, is_path=True, max_width_cm=17.5, max_height_cm=24.0, add_border=True, reduce_first_page=True)
                
                doc.add_page_break()
                insert_blue_banner(doc, "FORMWORK SKETCHES", font_size=16)

            # ==========================================
            # Word Generation per Config
            # ==========================================
            for idx, conf in enumerate(configs):
                if conf['cat'] == 'vertical':
                    if conf.get("wall_pdf_curr"):
                        doc.add_page_break()
                        insert_blue_banner(doc, f"CONCRETE PRESSURE CALCULATION ({conf['sub_cat'].upper()} {idx+1})", font_size=14)
                        conf["wall_pdf_curr"].seek(0)
                        append_pdf_stream_to_word(conf["wall_pdf_curr"].read(), doc, is_path=False, max_width_cm=17.5, max_height_cm=24.0, add_border=True, reduce_first_page=True)
                    
                    doc.add_page_break()
                    insert_blue_banner(doc, "FORMWORK SKETCHES", font_size=16)

                doc.add_page_break()
                
                # --- Panel Systems ---
                if conf.get('is_panel_system'):
                    insert_blue_banner(doc, f"CHECK FORMWORK ELEMENTS FOR {conf['sub_cat'].upper()} HEIGHT {conf.get('height', 0):.2f} m", font_size=14)
                    add_eq(doc, f"{conf['sub_cat']} Load = Concrete Pressure", italic=True)
                    add_eq(doc, f"W = {conf['w']:.2f} KN/m²\n")
                    chk_counter = 1
                    
                    add_heading_14(doc, f"{chk_counter}. Check of {conf['sys_name']} Panel:")
                    add_eq(doc, f"- Actual Pressure = {conf['w']:.2f} KN/m²")
                    add_eq(doc, f"- Allowable Pressure = {conf['panel_allowable']:.2f} KN/m²")
                    add_word_check(doc, "Check for Panel Safety", conf['w'], conf['panel_allowable'], "KN/m²")
                    chk_counter += 1
                    
                    if conf.get('tie_h') and conf['sys_name'] != "Circular Steel Panel System" and not conf.get('strongback', {}).get('active'):
                        add_heading_14(doc, f"{chk_counter}. Check of Tie Rod:")
                        tie_a = conf.get('tie_h', 0) * conf.get('tie_v', 0)
                        tie_f = conf.get('w', 0) * tie_a
                        add_eq(doc, f"- Area covered by 1 Tie rod = {conf.get('tie_h', 0)} x {conf.get('tie_v', 0)} = {tie_a:.2f} m²")
                        add_eq(doc, f"- Tension force on Tie rod = P x Area = {conf.get('w', 0):.2f} x {tie_a:.2f} = {tie_f:.2f} KN")
                        add_word_check(doc, "Check for Tie rod Safety", tie_f, 90.00, "KN")
                        chk_counter += 1
                        
                    if conf.get('bolt_h'):
                        add_heading_14(doc, f"{chk_counter}. Check of Bolts:")
                        bolt_a = conf.get('bolt_h', 0) * conf.get('bolt_v', 0)
                        bolt_f = conf.get('w', 0) * bolt_a
                        add_eq(doc, f"- Area covered by 1 Bolt = {conf.get('bolt_h', 0)} x {conf.get('bolt_v', 0)} = {bolt_a:.2f} m²")
                        add_eq(doc, f"- Tension/Shear force on Bolt = P x Area = {conf.get('w', 0):.2f} x {bolt_a:.2f} = {bolt_f:.2f} KN")
                        add_word_check(doc, "Check for Bolt Safety", bolt_f, 50.00, "KN")
                        chk_counter += 1
                    
                # --- Non-Panel Systems OR Single Sided Wall ---
                else:
                    if conf['cat'] == 'horizontal':
                        if conf['sub_cat'] == 'Beam': 
                            insert_blue_banner(doc, f"CHECK FORMWORK ELEMENTS FOR {conf['sub_cat'].upper()} ({conf['beam_b']}x{conf['ts']} m)", font_size=14)
                            add_eq(doc, f"Beam load = γc * depth + live load + formwork load", italic=True)
                        else: 
                            insert_blue_banner(doc, f"CHECK FORMWORK ELEMENTS FOR {conf['sub_cat'].upper()} THICKNESS {conf['ts']} m", font_size=14)
                            add_eq(doc, f"{conf['sub_cat']} load = γc * ts + live load + formwork load", italic=True)
                        
                        add_eq(doc, f"W = {conf['gamma_c']:.1f} * {conf['ts']:.2f} + {conf['live_load']:.1f} + {conf['fw_load']:.1f} = {conf['w']:.2f} KN/m²\n")
                    else: 
                        insert_blue_banner(doc, f"CHECK FORMWORK ELEMENTS FOR {conf['sub_cat'].upper()} HEIGHT {conf['height']:.2f} m", font_size=14)
                        add_eq(doc, f"{conf['sub_cat']} Load = Concrete Pressure", italic=True)
                        add_eq(doc, f"W = {conf['w']:.2f} KN/m²\n")
                    
                    chk_counter = 1
                    
                    # 1. Plywood
                    add_heading_14(doc, f"{chk_counter}. Plywood {conf['ply_thick']}:")
                    add_eq(doc, f"W_plywood = {conf['w']:.2f} KN/m²")
                    add_eq(doc, f"Max Spacing = {conf['s_spc']} m\n")
                    
                    add_eq(doc, "Check for moment:", bold=True)
                    M_ply = (conf['w'] * (conf['s_spc']**2)) / 10
                    Z_req = (M_ply * 100) / 3.41
                    add_eq(doc, f"M = W * L² / 10 = {conf['w']:.2f} * ({conf['s_spc']})² / 10 = {M_ply:.2f} KN.m")
                    add_eq(doc, f"Z_req = M * 100 / 3.41 = {M_ply:.2f} * 100 / 3.41 = {Z_req:.2f} cm³")
                    add_word_check(doc, None, Z_req, conf['ply_mall'], "cm³")
                    
                    E_ply, I_ply = (74.52, 48.60) if "18" in conf['ply_thick'] else (74.52, 77.0)
                    D_ply = (0.0068 * conf['w'] * (conf['s_spc']*100)**4) / (100 * E_ply * I_ply)
                    all_ply_d = (conf['s_spc']*1000)/300
                    
                    doc.add_paragraph()
                    add_eq(doc, "Check for deflection:", bold=True)
                    add_eq(doc, f"D = 0.0068 * W * L⁴ / (E * I) = 0.0068 * {conf['w']:.2f} * ({conf['s_spc']*100:.1f})⁴ / (100 * {E_ply:.2f} * {I_ply:.1f}) = {D_ply:.2f} mm")
                    add_word_check(doc, None, D_ply, all_ply_d, "mm", f"Allowable = L/300 = {all_ply_d:.2f} mm")
                    
                    chk_counter += 1
                    doc.add_page_break() 
                    
                    # 2. Secondary
                    max_s_span = max([float(x.strip()) for x in conf['s_sp'].split(',')])
                    add_heading_14(doc, f"{chk_counter}. Secondary Decking {conf['s_sec']}:")
                    add_eq(doc, f"- Secondary Beam length = {conf['s_L']:.2f} m")
                    add_eq(doc, f"- Max. spacing between main decking = {max_s_span:.2f} m")
                    add_eq(doc, f"- Max. spacing between Secondary decking = {conf['s_spc']:.2f} m")
                    s_w_calc = conf['w'] * conf['s_spc']
                    add_eq_highlight(doc, f"- W_sec = {conf['w']:.2f} x {conf['s_spc']:.2f} = ", f"{s_w_calc:.2f} KN/m'")
                    
                    prop_s = SECTIONS_DB[conf['s_sec']]
                    p_sk = doc.add_paragraph()
                    p_sk.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    p_sk.add_run().add_picture(io.BytesIO(conf['s_ld_img']), width=Cm(15.0))
                    add_centered_text(doc, "Load Assignment & Spans", size=12, color=RGBColor(100,100,100))
                    
                    add_eq(doc, "\nMaximum loads & deflections from attached Program Results:", underline=True)
                    s_img_bytes, s_M, s_V, s_D, _, _, s_Dtxt = generate_acrow_diagrams(
                        conf['s_sec'], conf['s_L'], conf['s_sup'], conf['s_ld'], 
                        prop_s['E'], prop_s['I'], prop_s['Mall'], prop_s['Qall'], Rall=None, transparent_bg=False
                    )
                    
                    add_word_check(doc, "Check for Moment", s_M, prop_s['Mall'], "KN.m")
                    add_word_check(doc, "Check for Shear", s_V, prop_s['Qall'], "KN")
                    add_word_check(doc, "Check for deflection", s_D, float(s_Dtxt.split('=')[-1].replace('mm','')), "mm", f"{s_Dtxt}")
                    
                    chk_counter += 1
                    doc.add_page_break()
                    
                    p_s = doc.add_paragraph()
                    p_s.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    p_s.add_run().add_picture(io.BytesIO(s_img_bytes), width=Cm(16.5))
                    add_centered_text(doc, f"Analysis Diagrams for Secondary Beam ({conf['s_sec']})", size=12)
                    
                    doc.add_page_break() 
                    
                    # 3. Main
                    max_m_span = max([float(x.strip()) for x in conf['m_sp'].split(',')])
                    add_heading_14(doc, f"{chk_counter}. Main Decking {conf['m_sec']}:")
                    add_eq(doc, f"- Main Beam length = {conf['m_L']:.2f} m")
                    add_eq(doc, f"- Max. Span of main {conf['m_sec'].split()[0]} = {max_m_span:.2f} m")
                    
                    if "Secondary Reaction" in conf.get("m_load_method", ""):
                        max_s = conf.get("max_s_rxn", 0.0)
                        m_w_calc = max_s / conf['s_spc'] if conf['s_spc'] > 0 else 0.0
                        add_eq(doc, f"- Spacing between Secondary decking = {conf['s_spc']:.2f} m")
                        add_eq(doc, f"- Max Reaction from Secondary Beam = {max_s:.2f} KN")
                        add_eq_highlight(doc, f"- W_main = Max Reaction / Sec. Spacing = {max_s:.2f} / {conf['s_spc']:.2f} = ", f"{m_w_calc:.2f} KN/m'")
                    else:
                        add_eq(doc, f"- Loaded width by one row of main {conf['m_sec'].split()[0]} = {conf['m_spc']:.2f} m")
                        m_w_calc = conf['w'] * conf['m_spc']
                        add_eq_highlight(doc, f"- W_main = {conf['w']:.2f} x {conf['m_spc']:.2f} = ", f"{m_w_calc:.2f} KN/m'")
                    
                    prop_m = SECTIONS_DB[conf['m_sec']]
                    p_msk = doc.add_paragraph()
                    p_msk.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    p_msk.add_run().add_picture(io.BytesIO(conf['m_ld_img']), width=Cm(15.0))
                    add_centered_text(doc, "Load Assignment & Spans", size=12, color=RGBColor(100,100,100))
                    
                    add_eq(doc, "\nMaximum loads & deflections from attached Program Results:", underline=True)
                    m_img_bytes, m_M, m_V, m_D, m_R, _, m_Dtxt = generate_acrow_diagrams(
                        conf['m_sec'], conf['m_L'], conf['m_sup'], conf['m_ld'], 
                        prop_m['E'], prop_m['I'], prop_m['Mall'], prop_m['Qall'], 
                        Rall=conf['t_allow'] if not conf.get('strongback',{}).get('active') else None, transparent_bg=False
                    )
                    
                    add_word_check(doc, "Check for Moment", m_M, prop_m['Mall'], "KN.m")
                    add_word_check(doc, "Check for Shear", m_V, prop_m['Qall'], "KN")
                    add_word_check(doc, "Check for deflection", m_D, float(m_Dtxt.split('=')[-1].replace('mm','')), "mm", f"{m_Dtxt}")
                    
                    chk_counter += 1
                    doc.add_page_break()
                    
                    p_m = doc.add_paragraph()
                    p_m.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    p_m.add_run().add_picture(io.BytesIO(m_img_bytes), width=Cm(16.5))
                    add_centered_text(doc, f"Analysis Diagrams for Main Beam ({conf['m_sec']})", size=12)
                    
                    doc.add_page_break() 

                    # 4. Support
                    if conf.get('t_allow') is not None and conf['t_allow'] < 900:
                        support_title = conf['t_name']
                        display_name = conf['t_name']
                        if conf['t_name'] == 'Acrow Prop' and conf.get('t_sub'):
                            support_title = f"{conf['t_name']} ({conf['t_sub']})"
                            display_name = conf['t_sub']
                        elif conf['t_name'] in ['Cup-lock', 'Ring-lock'] and conf.get('t_sub'):
                            support_title = f"{conf['t_name']} ({conf['t_sub']}) - Unbraced Length = {conf['t_unb']:.2f} m"
                            
                        add_heading_14(doc, f"{chk_counter}. {support_title}:")
                        add_eq(doc, f"- Load on {display_name} = Max. Reaction from Main Beam = {m_R:.2f} KN < {conf['t_allow']:.2f} KN")
                        add_word_check(doc, "Check for Support", m_R, conf['t_allow'], "KN")
                        chk_counter += 1

                # ----------------- STRONGBACK (Side-by-Side Diagrams) -----------------
                if conf.get('strongback', {}).get('active'):
                    sb = conf['strongback']
                    doc.add_page_break()
                    insert_blue_banner(doc, "STRONGBACK SYSTEM CHECK", font_size=14)
                    doc.add_paragraph()
                    
                    add_eq(doc, f"- Max Pressure (Pmax) = {conf.get('w', 0):.2f} kN/m²")
                    add_eq(doc, f"- Strongback Spacing = {sb['spc']:.3f} m")
                    add_eq(doc, f"- Base Load (W) = Pmax x Spacing = {sb['w']:.2f} kN/m'\n")
                    
                    add_heading_14(doc, f"- Check for Vertical Soldier:")
                    add_eq(doc, "FROM SAP CALCULATIONS .......", color=RGBColor(100,100,100))
                    add_word_check(doc, "Max. Bending Moment", sb['M_v'], SECTIONS_DB[sb['sv']]['Mall'], "KN.m")
                    add_word_check(doc, "Max. Shear force", sb['V_v'], SECTIONS_DB[sb['sv']]['Qall'], "KN")
                    add_word_check(doc, "Max deflection", sb['D_v'], sb['allw_D_v'], "mm", f"Allowable deflection = L/400 = {sb['allw_D_v']:.2f} mm")
                    
                    doc.add_paragraph()
                    add_heading_14(doc, f"- Check for Horizontal Soldier at ground:")
                    add_eq(doc, "FROM SAP CALCULATIONS .......", color=RGBColor(100,100,100))
                    add_word_check(doc, "Max. Bending Moment", sb['M_h'], SECTIONS_DB[sb['sh']]['Mall'], "KN.m")
                    add_word_check(doc, "Max. Shear force", sb['V_h'], SECTIONS_DB[sb['sh']]['Qall'], "KN")

                    doc.add_paragraph()
                    add_heading_14(doc, "- Check for Tilting Push Pull:")
                    add_eq(doc, "FROM SAP CALCULATIONS .......", color=RGBColor(100,100,100))
                    for st_d in sb['diags']:
                        val_t = next((abs(e['N_ax']) for e in sb.get('elements', []) if e['type'] == 'truss' and e['sec'] == st_d['type'].split()[0]), 0)
                        add_eq(doc, f"For Push Pull {st_d['type'].split()[0]}", underline=True, color=RGBColor(192,0,0))
                        add_word_check(doc, "N (Pact. from Sap)", val_t, STRUTS_DB.get(st_d['type'], {}).get('allow', 999), "KN")
                        
                    doc.add_paragraph()
                    add_heading_14(doc, "- Check for Lower Soldier:")
                    add_eq(doc, f"- Each Strongback tied with two tie rods at spacing {sb.get('tie_h',0)*100:.0f}cm")
                    t_force = sb.get('tie_force_total', sb.get('tie_T_single', 0.0) * 2.0)
                    add_eq_highlight(doc, f"- Assign load on Soldier= ", f"{t_force:.2f} KN")
                    add_word_check(doc, "Check for Moment", sb['waler_M'], SECTIONS_DB[sb['waler_sec']]['Mall'], "KN.m")
                    
                    doc.add_paragraph()
                    add_heading_14(doc, "- Check for Tie rod Ø15mm at angle 45°:")
                    add_word_check(doc, "Tie Force", sb.get('tie_T_single', 0.0), 90.00, "KN")
                    
                    if 'img_ld_single' in sb:
                        doc.add_page_break()
                        add_heading_14(doc, "- Strongback System Diagrams:")
                        doc.add_paragraph()
                        
                        add_two_images_side_by_side(
                            doc, 
                            sb['img_ld_single'], "Assigned Load Diagram", 
                            sb['img_ax_single'], "Axial Force Diagram", 
                            width_cm=8.5
                        )
                        doc.add_paragraph()
                        
                        add_two_images_side_by_side(
                            doc, 
                            sb['img_sh_single'], "Shear Force Diagram", 
                            sb['img_mo_single'], "Bending Moment Diagram", 
                            width_cm=8.5
                        )
                        doc.add_paragraph()
                        
                        p_rx = doc.add_paragraph()
                        p_rx.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        p_rx.add_run().add_picture(io.BytesIO(sb['img_rx_single']), width=Cm(10.0))
                        add_centered_text(doc, "Reactions Diagram", size=12)

                # ----------------- WIND & TILTING -----------------
                if conf.get('cat') == 'vertical' and conf.get('tilting', {}).get('active'):
                    td = conf['tilting']
                    doc.add_page_break()
                    
                    add_heading_14(doc, f"{chk_counter}. Check of Tilting System:")
                    chk_counter += 1
                    doc.add_paragraph()
                    
                    if os.path.exists("Wind_Maps.pdf"):
                        pdf_doc = fitz.open("Wind_Maps.pdf")
                        for page_num in range(len(pdf_doc)):
                            pix = pdf_doc.load_page(page_num).get_pixmap(dpi=150)
                            img = crop_white_margins(Image.open(io.BytesIO(pix.tobytes("png"))))
                            img_stream = io.BytesIO()
                            img.save(img_stream, format="PNG")
                            img_stream.seek(0)
                            target_width = 17.5
                            target_height = target_width / (img.width / img.height)
                            max_h = 19.5 if page_num == 0 else 22.5
                            if target_height > max_h: 
                                target_height = max_h
                                target_width = target_height * (img.width / img.height)
                            p_wm = doc.add_paragraph()
                            p_wm.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            p_wm.add_run().add_picture(img_stream, width=Cm(target_width), height=Cm(target_height))
                            img.close()
                            doc.add_page_break()
                        pdf_doc.close()
                    
                    add_heading_14(doc, "- Loads from Wind Pressure:")
                    add_eq(doc, "Refer to Regulations and Standards the Saudi Building Code (SBC) 2018", bold=True, color=RGBColor(192,0,0))
                    add_eq(doc, f"V = {td['v_wind']:.0f} m/s\t\t\t\t\tSBC 301 (2018)")
                    add_eq(doc, "G = 0.85\t\t\t\t\tSBC 301 (2018)")
                    add_eq(doc, "Cf = 1.30\t\t\t\t\tSBC 301 (2018)")
                    add_eq(doc, "I = 1.0\t\t\t\t\t\tSBC 301 (2018)")
                    add_eq(doc, "kd = 0.85\t\t\t\t\tSBC 301 (2018)")
                    add_eq(doc, "kzt = 1.0\t\t\t\t\tSBC 301 (2018)")
                    
                    p_kz = doc.add_paragraph()
                    force_ltr_left(p_kz)
                    p_kz.add_run("kz  =  ").font.name = 'Arial'
                    r_kz = p_kz.add_run(f"{td['kz_wind']:.2f}")
                    r_kz.font.highlight_color = WD_COLOR_INDEX.YELLOW
                    
                    qz = 0.613 * td['kz_wind'] * 1.00 * 0.85 * (td['v_wind'] ** 2)
                    add_eq(doc, "\n- Wind Pressure (qz) = 0.613 kz kzt kd V²   (N/M²)")
                    p2 = doc.add_paragraph()
                    force_ltr_left(p2)
                    r_qz = p2.add_run(f"qz = 0.613 x {td['kz_wind']:.2f} x 1.00 x 0.85 x ({td['v_wind']:.0f})²  =  {qz:.2f} N/M2")
                    r_qz.font.highlight_color = WD_COLOR_INDEX.YELLOW
                    r_qz.font.bold = True
                    
                    Af = td['h_panel'] * td['wp_panel']
                    add_eq(doc, "\n- Wind Force (F) = qz G Cf Af")
                    F_wind = (qz / 1000) * 0.850 * 1.300 * Af
                    add_eq(doc, f"F = {qz/1000:.3f} x 0.850 x 1.300 x {Af:.2f}  =  {F_wind:.2f} KN")
                    
                    p4 = doc.add_paragraph()
                    force_ltr_left(p4)
                    r_wd = p4.add_run(f"Distributed Wind Load/m'  =  {F_wind:.2f} / {td['h_panel']:.2f}  =  {td['w_dist']:.2f} KN/M")
                    r_wd.font.highlight_color = WD_COLOR_INDEX.YELLOW
                    r_wd.font.bold = True

                    if td.get('bkt_data') and td['bkt_data']['active']:
                        bkt = td['bkt_data']
                        doc.add_page_break()
                        add_heading_14(doc, "- Access Bracket Forces:")
                        add_eq(doc, f"Total Vertical Load (P) = {bkt['LL'] * bkt['L1'] * td['wp_panel']:.2f} KN")
                        add_eq(doc, f"Couple Force (F) = {bkt['F']:.2f} KN")

                    doc.add_page_break()
                    add_eq(doc, "Strut Axial Forces (N):", underline=True, bold=True)
                    if 'struts' in td:
                        for st_c in td['struts']: 
                            add_eq(doc, f"• For {st_c['type'].split()[0]} @ Y = {st_c['y']:.2f}m :")
                            add_word_check(doc, "Axial Force (N)", abs(st_c['N']), st_c['allow'], "KN")
                        
                    if 'img_w' in td:
                        add_two_images_side_by_side(
                            doc, 
                            td['img_w'], "Assigned Wind Load Diagram", 
                            td['img_n'], "Axial Force Diagram", 
                            width_cm=8.5
                        )
                        
                        if td.get('block_data') and td['block_data']['active']:
                            b_data = td['block_data']
                            add_heading_14(doc, "Check of Concrete Blocks & Expansion Bolts for double base Plate:")
                            
                            lines = [
                                {'text': "1- Concrete block dimensions:", 'bold': True},
                                {'text': f"- Concrete block: ({b_data['L']:.2f} X {b_data['W']:.2f} X {b_data['H']:.2f}) m"},
                                {'text': f"- Weight of Block = {b_data['L']:.2f} X {b_data['W']:.2f} X {b_data['H']:.2f} X {b_data['gamma']/10:.2f} = {b_data['W_block']/10:.3f} ton = {b_data['W_block']:.2f} kN"},
                                {'text': f"- Reaction in Y-direction = {b_data['max_ry']:.2f} kN"},
                                {'text': f"- Total Weight = {b_data['W_block']:.2f} - {b_data['max_ry']:.2f} = {b_data['N_eff']:.2f} kN\n"},
                                
                                {'text': "a. Check of Stability", 'bold': True, 'underline': True},
                                {'text': f"- M_st = {b_data['N_eff']:.2f} x {b_data['L']/2:.3f} = {b_data['M_st']:.2f} kN.m"},
                                {'text': f"- Reaction in X-direction = {b_data['max_rx']:.2f} kN"},
                                {'text': f"- M_ov = {b_data['max_rx']:.2f} x {b_data['H']/2:.3f} = {b_data['M_ov']:.2f} kN.m"},
                                {'text': f"- F.O.S = M_st / M_ov = {b_data['FOS_ov']:.2f} > 1.50   SAFE" if b_data['FOS_ov']>=1.5 else f"- F.O.S = {b_data['FOS_ov']:.2f} < 1.50   UNSAFE", 'bold': True, 'color': RGBColor(0,128,0) if b_data['FOS_ov']>=1.5 else RGBColor(255,0,0)},
                                
                                {'text': "\nb. Check Sliding:", 'bold': True, 'underline': True},
                                {'text': f"- Reaction in X-direction = ΣFx = {b_data['max_rx']:.2f} kN"},
                                {'text': f"- μ : Coefficient of friction = {b_data['mu']:.2f}"},
                                {'text': f"- ΣFy = μ X Total weight = {b_data['N_eff']*b_data['mu']:.2f} kN"},
                                {'text': f"- F.O.S = ΣFy / ΣFx = {b_data['FOS_sl']:.2f} > 1.50   SAFE" if b_data['FOS_sl']>=1.5 else f"- F.O.S = {b_data['FOS_sl']:.2f} < 1.50   UNSAFE", 'bold': True, 'color': RGBColor(0,128,0) if b_data['FOS_sl']>=1.5 else RGBColor(255,0,0)},
                                {'text': "\n2- Expansion bolts of Double base Plate:", 'bold': True, 'underline': True}
                            ]
                            add_text_and_image_side_by_side(doc, lines, td['img_r'], img_width_cm=8.5)
                            
                        else:
                            add_heading_14(doc, "Base Connections & Anchor Bolts:")
                            lines = [{'text': "Reactions Output:", 'bold': True}]
                            add_text_and_image_side_by_side(doc, lines, td['img_r'], img_width_cm=8.5)

                        max_rx_base = max(td.get('rx1', 0), td.get('rx2', 0))
                        max_ry_base = max(td.get('ry1', 0), td.get('ry2', 0))
                        
                        add_word_check(doc, "Revit Pin Shear Check", td.get('max_n', 0), 80.00, "KN")
                        add_word_check(doc, "Max Shear per Bolt (Worst Base)", max_rx_base/2, 29.50, "KN")
                        add_word_check(doc, "Max Tension per Bolt (Worst Base)", max_ry_base/2, 15.10, "KN")
                        
                        p_red = doc.add_paragraph()
                        p_red.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        p_red.paragraph_format.line_spacing = 1.5
                        run_red = p_red.add_run("According to these Loads Use for each Double Base Plate\n2 (HUS3 Screw anchor) M 14X115\nWhich are mentioned in attached reference.\nor use any equivalent bolts.")
                        run_red.font.name = 'Arial'
                        run_red.font.size = Pt(12)
                        run_red.font.color.rgb = RGBColor(255, 0, 0)

            has_double_sided_or_column = any((c.get('cat') == 'vertical' and not c.get('strongback', {}).get('active')) for c in configs)
            
            if "Vertical" in sys_cat and has_double_sided_or_column and os.path.exists("Hilti_Bolt.pdf"):
                doc.add_page_break()
                append_pdf_stream_to_word("Hilti_Bolt.pdf", doc, is_path=True, max_width_cm=17.5, max_height_cm=24.0, add_border=False)
                
            out = io.BytesIO()
            doc.save(out)
            st.success("✅ Smart Analysis Complete! Calculation Sheet generated successfully.")
            st.download_button("⬇️ Download Calculation Sheet", data=out.getvalue(), file_name=file_name, mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
            import matplotlib.pyplot as plt
            plt.close('all')
            gc.collect()
