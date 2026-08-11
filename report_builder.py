# report_builder.py

import os
import io
import fitz
import numpy as np
from PIL import Image, ImageOps, ImageChops
from docx import Document
from docx.shared import Cm, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_COLOR_INDEX
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import nsdecls, qn
from docx.oxml import parse_xml, OxmlElement
from config import SECTIONS_DB, STRUTS_DB
from math_solver import solve_beam_advanced, get_element_safety_details

def crop_white_margins(img):
    if img.mode != "RGB": 
        img = img.convert("RGB")
    bg = Image.new("RGB", img.size, (255, 255, 255))
    diff = ImageChops.difference(img, bg)
    bbox = diff.getbbox()
    if bbox:
        pad = 20
        return img.crop((max(0, bbox[0]-pad), max(0, bbox[1]-pad), min(img.width, bbox[2]+pad), min(img.height, bbox[3]+pad)))
    return img

def append_pdf_stream_to_word(pdf_source, doc_obj, is_path=False, max_width_cm=17.5, max_height_cm=24.0, add_border=True, reduce_first_page=False, crop_white=False):
    pdf_document = fitz.open(pdf_source) if is_path else fitz.open(stream=pdf_source, filetype="pdf")
    try:
        for page_num in range(len(pdf_document)):
            pix = pdf_document.load_page(page_num).get_pixmap(dpi=150)
            img = Image.open(io.BytesIO(pix.tobytes("png")))
            
            if crop_white: 
                img = crop_white_margins(img)
            if add_border: 
                img = ImageOps.expand(img, border=4, fill='#A9A9A9')
                
            img_width, img_height = img.size
            aspect_ratio = img_width / img_height
            target_width = max_width_cm
            current_max_h = 21.5 if (reduce_first_page and page_num == 0) else max_height_cm
            target_height = target_width / aspect_ratio
            
            if target_height > current_max_h: 
                target_height = current_max_h
                target_width = target_height * aspect_ratio
            if target_width > max_width_cm: 
                target_width = max_width_cm
                target_height = target_width / aspect_ratio
                
            img_stream = io.BytesIO()
            img.save(img_stream, format="PNG")
            img_stream.seek(0)
            img.close()
            
            p = doc_obj.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            p.add_run().add_picture(img_stream, width=Cm(target_width), height=Cm(target_height))
    finally:
        pdf_document.close()

def force_ltr_left(p):
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(12)
    pPr = p._element.get_or_add_pPr()
    bidi = pPr.find(qn('w:bidi'))
    if bidi is not None: 
        bidi.set(qn('w:val'), '0')
    else: 
        bidi = OxmlElement('w:bidi')
        bidi.set(qn('w:val'), '0')
        pPr.append(bidi)
        
    jc = pPr.find(qn('w:jc'))
    if jc is not None: 
        jc.set(qn('w:val'), 'left')
    else: 
        jc = OxmlElement('w:jc')
        jc.set(qn('w:val'), 'left')
        pPr.append(jc)

def add_heading_14(doc, text):
    p = doc.add_paragraph()
    force_ltr_left(p)
    run = p.add_run(text)
    run.font.name = 'Arial'
    run.font.size = Pt(14)
    run.font.bold = True
    run.underline = True

def add_eq(doc, text, italic=False, underline=False, bold=False, color=None):
    p = doc.add_paragraph()
    force_ltr_left(p)
    run = p.add_run(text)
    run.font.name = 'Arial'
    run.font.size = Pt(12)
    run.italic = italic
    run.underline = underline
    run.bold = bold
    if color: 
        run.font.color.rgb = color

def add_eq_highlight(doc, pre_text, high_text):
    p = doc.add_paragraph()
    force_ltr_left(p)
    r1 = p.add_run(pre_text)
    r1.font.name = 'Arial'
    r1.font.size = Pt(12)
    r2 = p.add_run(high_text)
    r2.font.name = 'Arial'
    r2.font.size = Pt(12)
    r2.font.bold = True
    r2.font.highlight_color = WD_COLOR_INDEX.YELLOW

def add_word_check(doc, title, act, allw, unit, extra_txt=""):
    p = doc.add_paragraph()
    force_ltr_left(p)
    
    if title:
        r1 = p.add_run(f"• {title}:\n")
        r1.font.name = 'Arial'
        r1.font.size = Pt(12)
        r1.font.bold = True
        
    r2 = p.add_run(f"  Max = {act:.2f} {unit}   <   {allw:.2f} {unit}   ")
    r2.font.name = 'Arial'
    r2.font.size = Pt(12)
    
    res = p.add_run("SAFE" if act <= allw else "UNSAFE")
    res.font.name = 'Arial'
    res.font.size = Pt(12)
    res.font.bold = True
    res.font.color.rgb = RGBColor(255, 0, 0)
    
    if extra_txt: 
        r3 = p.add_run(f"\n  ({extra_txt})")
        r3.font.name = 'Arial'
        r3.font.size = Pt(12)
        r3.italic = True

def add_fos_check(doc, title, act, min_allw):
    p = doc.add_paragraph()
    force_ltr_left(p)
    
    r1 = p.add_run(f"- {title} = {act:.2f} > {min_allw:.2f}   ")
    r1.font.name = 'Arial'
    r1.font.size = Pt(12)
    r1.font.bold = True
    r1.font.highlight_color = WD_COLOR_INDEX.YELLOW
    
    res = p.add_run("SAFE" if act >= min_allw else "UNSAFE")
    res.font.name = 'Arial'
    res.font.size = Pt(12)
    res.font.bold = True
    res.font.color.rgb = RGBColor(255, 0, 0)
    res.font.highlight_color = WD_COLOR_INDEX.YELLOW

def insert_blue_banner(doc_obj, text, font_size=14):
    table = doc_obj.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.rows[0].cells[0]
    cell.width = Cm(18.5) 
    tcPr = cell._tc.get_or_add_tcPr()
    tcPr.append(parse_xml(r'<w:vAlign {} w:val="center"/>'.format(nsdecls('w'))))
    tcPr.append(parse_xml(r'<w:shd {} w:fill="4A7BBB"/>'.format(nsdecls('w'))))
    
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    run = p.add_run(text)
    run.font.name = 'Arial'
    run.font.size = Pt(font_size)
    run.font.bold = True
    run.font.color.rgb = RGBColor(255, 255, 255)

def add_centered_text(doc_obj, text, size=12, color=RGBColor(0,0,0)):
    p = doc_obj.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(12)
    run = p.add_run(text)
    run.font.name = 'Arial'
    run.font.size = Pt(size)
    run.font.bold = False
    run.font.color.rgb = color

def add_two_images_side_by_side(doc_obj, img1_bytes, txt1, img2_bytes, txt2, width_cm=8.5):
    table = doc_obj.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    cell1 = table.rows[0].cells[0]
    cell1.width = Cm(width_cm)
    p1 = cell1.paragraphs[0] if len(cell1.paragraphs) > 0 else cell1.add_paragraph()
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p1.paragraph_format.line_spacing = 1.0
    p1.paragraph_format.space_after = Pt(0)
    p1.add_run().add_picture(io.BytesIO(img1_bytes), width=Cm(width_cm))
    
    p_txt1 = cell1.add_paragraph()
    p_txt1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_lbl1 = p_txt1.add_run(f"\n{txt1}")
    r_lbl1.font.name = 'Arial'
    r_lbl1.font.size = Pt(12)
    
    cell2 = table.rows[0].cells[1]
    cell2.width = Cm(width_cm)
    p2 = cell2.paragraphs[0] if len(cell2.paragraphs) > 0 else cell2.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.line_spacing = 1.0
    p2.paragraph_format.space_after = Pt(0)
    p2.add_run().add_picture(io.BytesIO(img2_bytes), width=Cm(width_cm))
    
    p_txt2 = cell2.add_paragraph()
    p_txt2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_lbl2 = p_txt2.add_run(f"\n{txt2}")
    r_lbl2.font.name = 'Arial'
    r_lbl2.font.size = Pt(12)

def add_text_and_image_side_by_side(doc_obj, text_lines_with_formatting, img_bytes, img_width_cm=8.5):
    table = doc_obj.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell_txt = table.rows[0].cells[0]
    cell_img = table.rows[0].cells[1]
    cell_txt.width = Cm(10.0)
    cell_img.width = Cm(img_width_cm)
    
    p_txt = cell_txt.paragraphs[0] if len(cell_txt.paragraphs) > 0 else cell_txt.add_paragraph()
    force_ltr_left(p_txt)
    
    for item in text_lines_with_formatting:
        run = p_txt.add_run(item.get('text', '') + "\n")
        run.font.name = 'Arial'
        run.font.size = Pt(11)
        run.font.bold = item.get('bold', False)
        if item.get('underline', False): run.font.underline = True
        if item.get('color'): run.font.color.rgb = item['color']
        
    p_img = cell_img.paragraphs[0] if len(cell_img.paragraphs) > 0 else cell_img.add_paragraph()
    p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_img.add_run().add_picture(io.BytesIO(img_bytes), width=Cm(img_width_cm))
    add_centered_text(cell_img, "\nReactions Diagram", size=12)

def perform_global_safety_check(configs_list):
    is_all_safe = True
    for conf in configs_list:
        if conf.get('is_panel_system') and conf.get('w', 0) > conf.get('panel_allowable', 999): 
            is_all_safe = False
            
        if conf.get('tie_h') and not conf.get('strongback', {}).get('active'):
            tie_load = conf.get('w', 0) * conf.get('tie_h', 0) * conf.get('tie_v', 0)
            if tie_load > 90.0: 
                is_all_safe = False
                
        if conf.get('bolt_h'):
            bolt_load = conf.get('w', 0) * conf.get('bolt_h', 0) * conf.get('bolt_v', 0)
            if bolt_load > 50.0: 
                is_all_safe = False
                
        if conf.get('s_sec'):
            s_M, s_V, s_D, s_aM, s_aV, s_aD = get_element_safety_details(conf, True)
            m_M, m_V, m_D, m_aM, m_aV, m_aD = get_element_safety_details(conf, False)
            
            if s_M > s_aM or s_V > s_aV or s_D > s_aD: 
                is_all_safe = False
            if m_M > m_aM or m_V > m_aV or m_D > m_aD: 
                is_all_safe = False
                
            if conf.get('t_allow') is not None and conf['t_allow'] < 900:
                prop_m = SECTIONS_DB[conf['m_sec']]
                _, _, _, _, R_m = solve_beam_advanced(
                    conf['m_L'], conf['m_sup'], conf['m_ld'], prop_m['E'], prop_m['I']
                )
                if np.max(R_m) > conf['t_allow']: 
                    is_all_safe = False
                    
        if conf.get('strongback', {}).get('active'):
            sb = conf['strongback']
            sb_m_v_safe = sb['M_v'] <= SECTIONS_DB[sb['sv']]['Mall']
            sb_v_v_safe = sb['V_v'] <= SECTIONS_DB[sb['sv']]['Qall']
            sb_m_h_safe = sb['M_h'] <= SECTIONS_DB[sb['sh']]['Mall']
            sb_v_h_safe = sb['V_h'] <= SECTIONS_DB[sb['sh']]['Qall']
            sb_tie_safe = sb['tie_T_single'] <= 90.0
            sb_waler_safe = sb['waler_M'] <= SECTIONS_DB[sb['waler_sec']]['Mall']
            sb_pin_safe = sb['max_diag_force'] <= 80.0
            
            if not (sb_m_v_safe and sb_v_v_safe and sb_m_h_safe and sb_v_h_safe and sb_tie_safe and sb_waler_safe and sb_pin_safe):
                is_all_safe = False
                
            for d in sb['diags']:
                d_force = next((abs(e['N_ax']) for e in sb['elements'] if e['type'] == 'truss' and e['sec'] == d['type'].split()[0]), 0)
                if d_force > d['allow']: 
                    is_all_safe = False
                    
        if conf.get('cat') == 'vertical' and conf.get('tilting', {}).get('active'):
            td = conf['tilting']
            if not td.get('length_safe', True): 
                is_all_safe = False
                
            if 'struts' in td:
                for st_c in td['struts']:
                    if abs(st_c['N']) > st_c['allow']: 
                        is_all_safe = False
                        
            max_rx = max(td.get('rx1', 0), td.get('rx2', 0))
            max_ry = max(td.get('ry1', 0), td.get('ry2', 0))
            
            if td.get('max_n', 0) > 80.0: is_all_safe = False
            if (max_ry / 2) > 15.1: is_all_safe = False
            if (max_rx / 2) > 29.5: is_all_safe = False
                
            if td.get('block_data') and td['block_data'].get('active'):
                if not td['block_data'].get('safe', False):
                    is_all_safe = False
                    
    return is_all_safe
