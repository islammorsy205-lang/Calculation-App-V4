# ==============================================================================
# BRIDGE MASTER - THE COMMERCIAL GRADE ENGINE (Ultimate Edition)
# ==============================================================================
# 🌟 Features: True 2D XZ Plane, Strict DXF Parsing via TABLE_ANALYSIS Layer,
# Parallel Isolated Session States, Smart Z-Snapping for Supports,
# Full calculation tables for Word Reports, Isolated Load Diagrams.
# J1, J2 Support Indexing, Strict Strut Filtering, and Load Replication.
# ==============================================================================

import streamlit as st
import numpy as np
import pandas as pd
import io
import os
import re
import math
import tempfile
import time
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
from matplotlib.patches import Polygon
from PIL import Image, ImageChops
from docx import Document
from docx.shared import Cm, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

from report_builder import insert_blue_banner, add_eq, append_pdf_stream_to_word
from config import SECTIONS_DB, STRUTS_DB

# ---------------------------------------------------------
# Dependency Checks
# ---------------------------------------------------------
try:
    import ezdxf
except ImportError:
    st.error("⚠️ مكتبة 'ezdxf' غير موجودة! برجاء تثبيتها: pip install ezdxf")
    ezdxf = None

try:
    import pytesseract
except ImportError:
    pytesseract = None


# =========================================================
# 0. Helper Functions & Styles
# =========================================================
def apply_plot_styles():
    """
    تطبيق تنسيقات الرسم البياني الهندسية الدقيقة والخطوط
    """
    matplotlib.rcParams['font.family'] = 'sans-serif'
    matplotlib.rcParams['font.sans-serif'] = ['Arial', 'Helvetica', 'DejaVu Sans']
    matplotlib.rcParams['axes.linewidth'] = 0.3
    matplotlib.rcParams['font.size'] = 7


def get_short_name(sec_name):
    """
    استخراج الاسم القصير للقطاع وحذف التفاصيل بين الأقواس لتنظيم الرسم
    """
    return re.sub(r'\s*\(.*?\)', '', sec_name).strip()


def crop_image_bbox(img_bytes):
    """
    قص حواف الصور الناتجة بشكل صارم (Aggressive Cropping)
    لإزالة أي بكسل أبيض حول الدياجرام ليأخذ أفضل Fit في تقرير الوورد
    """
    img = Image.open(io.BytesIO(img_bytes)).convert("RGBA")
    bg = Image.new("RGBA", img.size, (255, 255, 255, 0))
    diff = ImageChops.difference(img, bg)
    bbox = diff.getbbox()
    if bbox:
        padding = 5
        bbox = (
            max(0, bbox[0] - padding),
            max(0, bbox[1] - padding),
            min(img.size[0], bbox[2] + padding),
            min(img.size[1], bbox[3] + padding)
        )
        img = img.crop(bbox)
    
    out = io.BytesIO()
    img.save(out, format='PNG')
    return out.getvalue()


def safe_render_fig(fig):
    """
    دالة لحفظ وإخراج الصورة بأمان وبأعلى جودة للطباعة
    """
    try:
        plt.subplots_adjust(left=0.01, right=0.99, top=0.99, bottom=0.01)
        buf = io.BytesIO()
        fig.savefig(buf, format='png', dpi=400, bbox_inches='tight', pad_inches=0.0, transparent=True)
        return crop_image_bbox(buf.getvalue())
    finally:
        plt.close(fig)


def draw_reaction_arrow(ax, node_x, node_z, force_mag, axis_nx, axis_nz):
    """
    رسم أسهم ردود الأفعال بدقة مع تمييز الشد (أحمر) والضغط (أزرق)
    """
    if abs(force_mag) < 0.001: 
        return
        
    arr_L = 0.5  
    sgn = np.sign(force_mag)
    dx = sgn * axis_nx
    dz = sgn * axis_nz
    
    start_x = node_x - arr_L * dx
    start_z = node_z - arr_L * dz
    
    if force_mag >= 0:
        arr_c = 'blue'
    else:
        arr_c = 'red'
        
    ax.arrow(
        start_x, start_z, arr_L * dx, arr_L * dz, 
        length_includes_head=True, head_width=0.08, head_length=0.12, 
        fc=arr_c, ec=arr_c, lw=0.8, zorder=5
    )
    
    ax.text(
        start_x - 0.15 * dx, start_z - 0.15 * dz, 
        f"{force_mag:+.2f}", color=arr_c, fontsize=7, 
        fontname='Arial', ha='center', va='center'
    )


def eval_seg_point(seg, s_val):
    """
    حساب الإحداثيات المطلقة (X, Z) لأي نقطة على أي Segment بناءً على المسافة S
    """
    if seg.get('is_divided'):
        actual_s = s_val + seg.get('parent_offset', 0.0)
        return eval_seg_point(seg['parent_seg'], actual_s)

    L = seg.get('L', 0.0)
    s_val = min(max(s_val, 0.0), L)
    
    if L > 1e-6:
        ratio = s_val / L
    else:
        ratio = 0.0
    
    if seg.get('is_dxf') or 'abs_p1' in seg:
        p1 = seg.get('abs_p1', (0,0))
        p2 = seg.get('abs_p2', (0,0))
        px = p1[0] + ratio * (p2[0] - p1[0])
        pz = p1[1] + ratio * (p2[1] - p1[1])
        th = math.atan2(p2[1] - p1[1], p2[0] - p1[0])
        return px, pz, th
    
    return 0.0, 0.0, 0.0


def get_closest_segment_exact(pt, segs):
    """
    إيجاد أقرب Segment لأي نقطة مرسومة للالتحام المغناطيسي الصارم
    """
    min_d = 9999.0
    best_idx = 0
    best_s = 0.0
    px = pt[0]
    pz = pt[1]
    
    for idx, seg in enumerate(segs):
        if seg.get('is_divided'):
            temp_seg = seg['parent_seg']
            L_orig = seg['parent_seg'].get('L', 0.0)
        else:
            temp_seg = seg
            L_orig = seg.get('L', 0.0)
            
        if 'abs_p1' in temp_seg:
            p1 = np.array(temp_seg['abs_p1'])
            p2 = np.array(temp_seg['abs_p2'])
            v = p2 - p1
            w = np.array([px, pz]) - p1
            c2 = np.dot(v, v)
            
            if c2 > 1e-6:
                ratio = np.dot(w, v) / c2
            else:
                ratio = 0.0
                
            ratio = max(0.0, min(1.0, ratio))
            proj = p1 + ratio * v
            d = np.linalg.norm(np.array([px, pz]) - proj)
            
            if d < min_d:
                min_d = d
                best_idx = idx
                best_s = ratio * L_orig
                if seg.get('is_divided'):
                    best_s -= seg.get('parent_offset', 0.0)
                    
    return min_d, best_idx, best_s


def get_shifted_coords_along_segment(px, pz, ds, segs):
    """
    الزحزحة الموازية لتحريك النهايز على نفس خط السولجر دون مغادرته
    """
    if abs(ds) < 1e-4: 
        return px, pz
        
    d_min, best_idx, best_s = get_closest_segment_exact((px, pz), segs)
    
    if d_min > 0.5: 
        return px + ds, pz
    
    seg = segs[best_idx]
    new_s = max(0.0, min(best_s + ds, seg.get('L', 0.0)))
    nx, nz, _ = eval_seg_point(seg, new_s)
    
    return nx, nz


# =========================================================
# 💡 STRICT STRUT FILTERING LOGIC
# =========================================================
def get_valid_strut_names():
    """الفلتر الصارم لأسماء النهايز حسب الأولوية المطلوبة"""
    if not STRUTS_DB: return ["PPH"]
    valid_struts = []
    for s_name in STRUTS_DB.keys():
        name_u = s_name.upper()
        base_name = s_name.split('(')[0].strip()
        
        # 1. إلغاء TILT و MMP تماماً
        if "TILT" in name_u or "MMP" in name_u: continue
        # 2. إلغاء أي قطاع ينتهي بـ 1 أو 3
        if base_name.endswith('1') or base_name.endswith('3'): continue
        
        # 3. تحديد الأولوية (PPS أولاً، ثم PPH)
        priority = 99
        if "PPS" in name_u: priority = 1
        elif "PPH" in name_u: priority = 2
            
        valid_struts.append({'name': s_name, 'pri': priority})
        
    valid_struts.sort(key=lambda x: x['pri'])
    return [x['name'] for x in valid_struts] if valid_struts else ["PPH"]


def get_optimal_strut_section(req_length, req_axial_force):
    """
    البحث عن أفضل قطاع آمن للنهايز من ضمن القائمة المفلترة فقط
    """
    valid_names = get_valid_strut_names()
    valid_struts = []
    
    for s_name in valid_names:
        s_props = STRUTS_DB.get(s_name, {})
        m = re.search(r'\((\d+\.\d+):(\d+\.\d+)m\)', s_name)
        if m:
            min_L = float(m.group(1))
            max_L = float(m.group(2))
            
            if min_L <= req_length <= max_L:
                if s_props.get('allow', 0.0) >= abs(req_axial_force):
                    valid_struts.append({
                        'name': s_name, 
                        'allowable': s_props.get('allow', 0.0)
                    })
                    
    if not valid_struts: 
        return None
        
    # اختيار الأكثر أماناً والأقل وزناً
    valid_struts.sort(key=lambda x: x['allowable']) 
    return valid_struts[0]['name']


# =========================================================
# 1. THE STRICT DXF PARSER (TABLE_ANALYSIS & Smart Extraction)
# =========================================================
def parse_dxf_bridge_cases(file_bytes):
    """
    مُحلل الـ DXF المتطور: يقرأ الأريات والقطاعات.
    💡 يحتوي على خوارزمية التماثل (Symmetry Auto-Detector) لحل مشكلة القطاعات 
    المجهولة في الجانب الأيسر وتوقيع الأحمال عليها بامتياز.
    """
    if ezdxf is None: 
        return None
        
    tmp_path = ""
    try:
        try: 
            dxf_str = file_bytes.decode('utf-8')
        except UnicodeDecodeError: 
            dxf_str = file_bytes.decode('cp1252', errors='ignore')
            
        with tempfile.NamedTemporaryFile(delete=False, suffix=".dxf", mode='w', encoding='utf-8') as tmp:
            tmp.write(dxf_str)
            tmp_path = tmp.name
            
        doc = ezdxf.readfile(tmp_path)
        msp = doc.modelspace()
        
        layer_table = 'TABLE_ANALYSIS'
        layer_supp = 'SUPPORT'
        layer_text = 'TEXT_DATA'
        layer_frame = 'FRAME'
        layer_strut = 'PUSH_PULL'
        
        table_boxes = []
        
        for e in msp:
            if e.dxftype() in ['LWPOLYLINE', 'POLYLINE'] and e.dxf.layer.upper() == layer_table:
                if e.dxftype() == 'LWPOLYLINE':
                    points = e.get_points('xy')
                else:
                    points = [v.dxf.location for v in e.vertices]
                    
                xs = [p[0] / 1000.0 for p in points]
                zs = [p[1] / 1000.0 for p in points]
                
                table_boxes.append({
                    'min_x': min(xs), 'max_x': max(xs),
                    'min_z': min(zs), 'max_z': max(zs),
                    'cx': (min(xs) + max(xs)) / 2.0,
                    'cz': (min(zs) + max(zs)) / 2.0
                })
                
        if not table_boxes:
            all_xs = []
            all_zs = []
            for e in msp:
                if e.dxftype() == 'LINE' and e.dxf.layer.upper() == layer_frame:
                    x1, z1 = e.dxf.start.x / 1000.0, e.dxf.start.y / 1000.0
                    x2, z2 = e.dxf.end.x / 1000.0, e.dxf.end.y / 1000.0
                    if math.hypot(x2 - x1, z2 - z1) >= 0.01:
                        all_xs.extend([x1, x2])
                        all_zs.extend([z1, z2])
            
            if all_xs and all_zs:
                table_boxes.append({
                    'min_x': min(all_xs) - 1.0, 'max_x': max(all_xs) + 1.0,
                    'min_z': min(all_zs) - 1.0, 'max_z': max(all_zs) + 1.0,
                    'cx': (min(all_xs) + max(all_xs)) / 2.0,
                    'cz': (min(all_zs) + max(all_zs)) / 2.0
                })
                
        table_boxes.sort(key=lambda b: b['min_x'])
        
        cases_raw = []
        for i, box in enumerate(table_boxes):
            cases_raw.append({
                'box': box, 'frames': [], 'struts': [], 
                'supports': [], 'cut_points': [], 's_texts': [], 'a_texts': []
            })
            
        for e in msp:
            layer = e.dxf.layer.upper()
            dxftype = e.dxftype()
            x_cad = 0.0
            z_cad = 0.0 
            is_valid_point = False
            
            if dxftype in ['POINT', 'CIRCLE']:
                if dxftype == 'POINT':
                    x_cad = e.dxf.location.x / 1000.0
                    z_cad = e.dxf.location.y / 1000.0
                else:
                    x_cad = e.dxf.center.x / 1000.0
                    z_cad = e.dxf.center.y / 1000.0
                is_valid_point = True
                
            elif dxftype == 'INSERT':
                x_cad = e.dxf.insert.x / 1000.0
                z_cad = e.dxf.insert.y / 1000.0
                is_valid_point = True
                
            if is_valid_point:
                for c in cases_raw:
                    b = c['box']
                    if b['min_x'] <= x_cad <= b['max_x'] and b['min_z'] <= z_cad <= b['max_z']:
                        norm_x = x_cad - b['cx']
                        norm_z = z_cad - b['cz']
                        if layer == layer_supp: 
                            c['supports'].append({'x': norm_x, 'z': norm_z, 'type': 'Roller', 'angle': 0.0})
                        elif layer == layer_text: 
                            c['cut_points'].append({'x': norm_x, 'z': norm_z})
                        break

            elif dxftype in ['TEXT', 'MTEXT']:
                x_cad = e.dxf.insert.x / 1000.0
                z_cad = e.dxf.insert.y / 1000.0
                
                for c in cases_raw:
                    b = c['box']
                    if b['min_x'] <= x_cad <= b['max_x'] and b['min_z'] <= z_cad <= b['max_z']:
                        if layer == layer_text:
                            if dxftype == 'MTEXT': txt = e.text
                            else: txt = e.dxf.text
                                
                            # تنظيف التيكست من شوائب الكاد تماماً
                            txt = re.sub(r'\\[A-Za-z0-9]+;', '', txt)
                            txt = txt.upper().replace('\n', '').replace('\r', '').replace(' ', '')
                            
                            s_m = re.search(r'S(\d+)=([\d\.]+)', txt)
                            s_lbl = re.search(r'S(\d+)(?!=)', txt)
                            a_m = re.search(r'A(\d+)=([\d\.]+)', txt)
                            
                            norm_x = x_cad - b['cx']
                            norm_z = z_cad - b['cz']
                            
                            if s_m: 
                                c['s_texts'].append({'idx': int(s_m.group(1)), 'val': float(s_m.group(2)), 'x': norm_x, 'z': norm_z})
                            elif s_lbl:
                                c['s_texts'].append({'idx': int(s_lbl.group(1)), 'val': 0.0, 'x': norm_x, 'z': norm_z})
                                
                            if a_m: 
                                c['a_texts'].append({'idx': int(a_m.group(1)), 'val': float(a_m.group(2)), 'x': norm_x, 'z': norm_z})
                        break
                        
            elif dxftype in ['LINE', 'LWPOLYLINE', 'POLYLINE']:
                if dxftype != 'LINE': entities = list(e.virtual_entities())
                else: entities = [e]
                    
                for sub_e in entities:
                    if sub_e.dxftype() == 'LINE':
                        x1 = sub_e.dxf.start.x / 1000.0
                        z1 = sub_e.dxf.start.y / 1000.0
                        x2 = sub_e.dxf.end.x / 1000.0
                        z2 = sub_e.dxf.end.y / 1000.0
                        
                        if math.hypot(x2 - x1, z2 - z1) < 0.01:
                            continue
                            
                        mid_x = (x1 + x2) / 2.0
                        mid_z = (z1 + z2) / 2.0
                        
                        for c in cases_raw:
                            b = c['box']
                            if b['min_x'] <= mid_x <= b['max_x'] and b['min_z'] <= mid_z <= b['max_z']:
                                norm_x1 = x1 - b['cx']
                                norm_z1 = z1 - b['cz']
                                norm_x2 = x2 - b['cx']
                                norm_z2 = z2 - b['cz']
                                
                                if layer == layer_frame: 
                                    c['frames'].append({'x1': norm_x1, 'z1': norm_z1, 'x2': norm_x2, 'z2': norm_z2})
                                elif layer == layer_strut: 
                                    c['struts'].append({'x1': norm_x1, 'z1': norm_z1, 'x2': norm_x2, 'z2': norm_z2})
                                break

        processed_cases = []
        for c_idx, c in enumerate(cases_raw):
            if not c['frames']: 
                continue 
            
            base_segments = []
            for i, line in enumerate(c['frames']):
                L = math.hypot(line['x2'] - line['x1'], line['z2'] - line['z1'])
                base_segments.append({
                    'name': f"F{i+1}", 'master_idx': i, 
                    'type': 'Straight Line', 'Shape Type': 'Straight Line', 
                    'L': L, 'is_dxf': True, 
                    'abs_p1': (line['x1'], line['z1']), 
                    'abs_p2': (line['x2'], line['z2'])
                })
                
            # إزالة التطابق Overlapping
            unique_segments = []
            for seg in base_segments:
                is_dup = False
                for u_seg in unique_segments:
                    d_p1 = math.hypot(seg['abs_p1'][0] - u_seg['abs_p1'][0], seg['abs_p1'][1] - u_seg['abs_p1'][1])
                    d_p2 = math.hypot(seg['abs_p2'][0] - u_seg['abs_p2'][0], seg['abs_p2'][1] - u_seg['abs_p2'][1])
                    d_p1_rev = math.hypot(seg['abs_p1'][0] - u_seg['abs_p2'][0], seg['abs_p1'][1] - u_seg['abs_p2'][1])
                    d_p2_rev = math.hypot(seg['abs_p2'][0] - u_seg['abs_p1'][0], seg['abs_p2'][1] - u_seg['abs_p1'][1])
                    if (d_p1 < 0.05 and d_p2 < 0.05) or (d_p1_rev < 0.05 and d_p2_rev < 0.05):
                        is_dup = True
                        break
                if not is_dup: 
                    unique_segments.append(seg)
                    
            base_segments = unique_segments
            
            dxf_areas = []
            # 1. التسمية الأساسية من التيكست
            for s_txt in c['s_texts']:
                min_d, best_idx, _ = get_closest_segment_exact((s_txt['x'], s_txt['z']), base_segments)
                if min_d < 4.0:  # 💡 زيادة مسافة البحث لاصطياد التيكست البعيد
                    seg_final_name = f"S{s_txt['idx']}"
                    base_segments[best_idx]['name'] = seg_final_name
                    
                    a_txt = None
                    for a in c['a_texts']:
                        if a['idx'] == s_txt['idx']:
                            a_txt = a
                            break
                            
                    if a_txt and s_txt['val'] > 1e-4:
                        if not any(d['segment'] == seg_final_name for d in dxf_areas):
                            dxf_areas.append({
                                'seg_idx': best_idx, 
                                'segment': seg_final_name, 
                                'length': s_txt['val'], 
                                'area': a_txt['val']
                            })
                            
            # 💡 2. خوارزمية التماثل السحرية (Symmetry Auto-Detector) لحل مشكلة السكاشن المجهولة يساراً
            all_xs = []
            for seg in base_segments:
                all_xs.extend([seg['abs_p1'][0], seg['abs_p2'][0]])
            if all_xs:
                center_x = (min(all_xs) + max(all_xs)) / 2.0
                named_segs = [seg for seg in base_segments if seg['name'].startswith('S') and seg['name'] not in ["S30", "S31", "S40", "S41"]]
                
                for n_seg in named_segs:
                    p1_x, p1_z = n_seg['abs_p1']
                    p2_x, p2_z = n_seg['abs_p2']
                    m_p1_x = center_x - (p1_x - center_x)
                    m_p2_x = center_x - (p2_x - center_x)

                    for u_seg in base_segments:
                        if u_seg['name'].startswith('F'): # يتم تطبيقها فقط على القطاعات غير المسماة
                            u_p1_x, u_p1_z = u_seg['abs_p1']
                            u_p2_x, u_p2_z = u_seg['abs_p2']

                            d1 = math.hypot(u_p1_x - m_p1_x, u_p1_z - p1_z) + math.hypot(u_p2_x - m_p2_x, u_p2_z - p2_z)
                            d2 = math.hypot(u_p1_x - m_p2_x, u_p1_z - p2_z) + math.hypot(u_p2_x - m_p1_x, u_p2_z - p1_z)

                            if min(d1, d2) < 1.0: # 1 متر سماحية في الكاد
                                u_seg['name'] = n_seg['name'] # تم تسمية القطاع الشمال بنجاح!
            
            unlabeled = []
            for idx, seg in enumerate(base_segments):
                if seg['name'].startswith('F'):
                    mid_x = (seg['abs_p1'][0] + seg['abs_p2'][0]) / 2.0
                    mid_z = (seg['abs_p1'][1] + seg['abs_p2'][1]) / 2.0
                    unlabeled.append({'idx': idx, 'mx': mid_x, 'mz': mid_z})
                    
            if unlabeled:
                unlabeled.sort(key=lambda item: item['mz'])
                
                bottom_cands = [u for u in unlabeled if u['mz'] - unlabeled[0]['mz'] < 0.5]
                bottom_cands.sort(key=lambda item: item['mx'])
                if len(bottom_cands) > 0: base_segments[bottom_cands[0]['idx']]['name'] = "S30"
                if len(bottom_cands) > 1: base_segments[bottom_cands[-1]['idx']]['name'] = "S31"
                        
                top_cands = [u for u in unlabeled if unlabeled[-1]['mz'] - u['mz'] < 0.5]
                top_cands.sort(key=lambda item: item['mx'])
                if len(top_cands) > 0 and base_segments[top_cands[0]['idx']]['name'].startswith('F'): base_segments[top_cands[0]['idx']]['name'] = "S40"
                if len(top_cands) > 1 and base_segments[top_cands[-1]['idx']]['name'].startswith('F'): base_segments[top_cands[-1]['idx']]['name'] = "S41"

            # 💡 3. التوزيع الشامل للأحمال على جميع السكاشن ذات المسمى المتطابق
            initial_loads = []
            for area_item in dxf_areas:
                s_name = area_item['segment']
                w_val = (area_item['area'] * 25.0 * 1.30) / area_item['length']
                
                # إيجاد جميع القطاعات التي تحمل نفس الاسم (اليمين والشمال معاً بفضل خوارزمية التماثل)
                matching_indices = [idx_m for idx_m, seg_m in enumerate(base_segments) if seg_m['name'] == s_name]
                
                if matching_indices:
                    t_mode = 'Single Segment' if len(matching_indices) == 1 else 'Multiple Segments'
                    initial_loads.append({
                        'seg_idx': matching_indices[0], 
                        'category': 'Dead Load', 
                        'type': 'Uniform', 
                        'dir': 'Global Z (Vertical)', 
                        'target_mode': t_mode, 
                        'target_segs_idx': matching_indices, 
                        'start': 0.0, 
                        'end': base_segments[matching_indices[0]]['L'], 
                        'w1': -abs(w_val), 'w2': -abs(w_val), 'loc': 0.0,
                        'is_auto': True
                    })

            # توليد اللايف لود الافتراضي وتجميعه في خانة واحدة
            base_ll_w = -abs(2.90 * 1.30)
            valid_ll_indices = [idx for idx, seg in enumerate(base_segments) if seg['name'] not in ["S30", "S31"]]
            if valid_ll_indices:
                initial_loads.append({
                    'seg_idx': valid_ll_indices[0], 
                    'category': 'Live Load', 
                    'type': 'Uniform', 
                    'dir': 'Global Z (Vertical)', 
                    'target_mode': 'Multiple Segments', 
                    'target_segs_idx': valid_ll_indices, 
                    'start': 0.0, 'end': base_segments[valid_ll_indices[0]]['L'], 
                    'w1': base_ll_w, 'w2': base_ll_w, 'loc': 0.0,
                    'is_auto': True
                })

            strut_opts = get_valid_strut_names()
            struts_mapped = []
            for line in c['struts']:
                if line['z1'] > line['z2']: 
                    tx, tz, bx, bz = line['x1'], line['z1'], line['x2'], line['z2']
                else: 
                    tx, tz, bx, bz = line['x2'], line['z2'], line['x1'], line['z1']
                    
                struts_mapped.append({
                    'tx': tx, 'tz': tz, 'bx': bx, 'bz': bz, 
                    'sec': strut_opts[0] if strut_opts else "Unknown"
                })
                
            if c['supports']:
                c['supports'].sort(key=lambda sp: sp['x'])
                c['supports'][0]['type'] = 'Hinged'
                
            processed_cases.append({
                'title': f"Table {c_idx+1}", 
                'segments': base_segments, 
                'struts': struts_mapped, 
                'supports': c['supports'], 
                'cut_points': c['cut_points'], 
                'dxf_areas': dxf_areas, 
                'loads': initial_loads
            })
            
        return processed_cases
        
    except Exception as e:
        st.error(f"DXF Parsing Error: {e}")
        return None
    finally:
        if tmp_path and os.path.exists(tmp_path):
            try: 
                os.remove(tmp_path)
            except: 
                pass
            # =========================================================
# 2. Meshing & FEA Matrix Engine (True 2D Frame in XZ Plane)
# =========================================================
def perform_smart_division(base_segments, supports, struts, cut_points=[]):
    """
    التقطيع الذكي (Smart Division) لغرض رسم المخرجات بشكل دقيق وتقطيع الفريمات
    عند أماكن الركائز والنهايز ونقاط التحميل.
    """
    cut_points_dict = {}
    for i, seg in enumerate(base_segments):
        cut_points_dict[i] = {0.0, seg.get('L', 0.0)}
        
    for sp in supports:
        d_min, w_seg, w_s = get_closest_segment_exact((sp['x'], sp['z']), base_segments)
        if d_min < 0.30: 
            cut_points_dict[w_seg].add(min(max(w_s, 0.0), base_segments[w_seg]['L']))
            
    for st in struts:
        dt, wt_seg, wt_s = get_closest_segment_exact((st['tx'], st['tz']), base_segments)
        if dt < 0.30: 
            cut_points_dict[wt_seg].add(min(max(wt_s, 0.0), base_segments[wt_seg]['L']))
            
        db, wb_seg, wb_s = get_closest_segment_exact((st['bx'], st['bz']), base_segments)
        if db < 0.30: 
            cut_points_dict[wb_seg].add(min(max(wb_s, 0.0), base_segments[wb_seg]['L']))
            
    for cp in cut_points:
        d_min, w_seg, w_s = get_closest_segment_exact((cp['x'], cp['z']), base_segments)
        if d_min < 0.30: 
            cut_points_dict[w_seg].add(min(max(w_s, 0.0), base_segments[w_seg]['L']))

    divided_segments = []
    sub_letters = "abcdefghijklmnopqrstuvwxyz"
    
    for m_idx, s_vals_set in sorted(cut_points_dict.items()):
        master_seg = base_segments[m_idx]
        sorted_s = sorted(list(s_vals_set))
        num_sub = len(sorted_s) - 1
        
        for k in range(num_sub):
            s_start = sorted_s[k]
            s_end = sorted_s[k+1]
            
            if s_end - s_start < 1e-4: 
                continue
                
            if num_sub == 1:
                sub_name = master_seg['name']
            else:
                sub_name = f"{master_seg['name']}-{sub_letters[k % 26]}"
                
            new_seg = master_seg.copy()
            new_seg.update({
                'name': sub_name, 
                'is_divided': True, 
                'parent_seg': master_seg, 
                'parent_offset': s_start, 
                'L': s_end - s_start, 
                'master_idx': m_idx
            })
            divided_segments.append(new_seg)
            
    return divided_segments


def build_chain_mesh(segments, seg_sections, loads, struts, supports, cut_points=[], mesh_size=0.50):
    """
    بناء مصفوفة العقد (Nodes) والعناصر (Elements) وتجميع الأحمال لعمل الـ FEA.
    يأخذ في الاعتبار التقسيمات الصغيرة (Mesh Size) لزيادة دقة النتائج.
    """
    nodes = []
    elements = []
    nodal_loads = []
    node_tol = 0.01 
    
    def get_or_add_node(x, z):
        for i, n in enumerate(nodes):
            if abs(n[0] - x) < node_tol and abs(n[1] - z) < node_tol: 
                return i
        nodes.append([x, z])
        return len(nodes) - 1

    support_injections = {i: [] for i in range(len(segments))}
    supports_list_out = []
    
    for sup in supports:
        sx = sup['x']
        sz = sup.get('z', sup.get('y', 0.0))
        min_d, w_seg, w_s = get_closest_segment_exact((sx, sz), segments)
        
        if min_d < 0.30: 
            support_injections[w_seg].append(w_s)
            
        nid = get_or_add_node(sx, sz)
        supports_list_out.append({
            'node': nid, 
            'type': sup.get('type', 'Roller'), 
            'angle': sup.get('angle', 0.0)
        })

    for cp in cut_points:
        cx = cp['x']
        cz = cp.get('z', cp.get('y', 0.0))
        min_d, w_seg, w_s = get_closest_segment_exact((cx, cz), segments)
        if min_d < 0.30: 
            support_injections[w_seg].append(w_s)
        get_or_add_node(cx, cz)

    for st_idx, st_item in enumerate(struts):
        tx = st_item['tx']
        tz = st_item.get('tz', st_item.get('ty', 0.0))
        bx = st_item['bx']
        bz = st_item.get('bz', st_item.get('by', 0.0))
        
        dt, wt_seg, wt_s = get_closest_segment_exact((tx, tz), segments)
        if dt < 0.30: 
            support_injections[wt_seg].append(wt_s)
            tx, tz, _ = eval_seg_point(segments[wt_seg], wt_s)
            
        db, wb_seg, wb_s = get_closest_segment_exact((bx, bz), segments)
        if db < 0.30: 
            support_injections[wb_seg].append(wb_s)
            bx, bz, _ = eval_seg_point(segments[wb_seg], wb_s)
            
        top_node = get_or_add_node(tx, tz)
        bot_node = get_or_add_node(bx, bz)
        
        elements.append({
            'type': 'truss', 
            'group': 'strut', 
            'sec': st_item.get('sec', 'Unknown'),
            'n1': bot_node, 
            'n2': top_node, 
            'strut_idx': st_idx, 
            'E': 21000000.0, 
            'A': 0.001
        })

    for i, seg in enumerate(segments):
        L = seg['L']
        key_s_vals = [0.0, L]
        key_s_vals.extend(support_injections[i])
        
        for ld in loads:
            if ld.get('seg_idx') == i: 
                key_s_vals.extend([ld['start'], ld['end']])
                
        num_subdivisions = max(1, int(np.ceil(L / mesh_size)))
        mesh_points = np.linspace(0, L, num_subdivisions + 1)
        key_s_vals.extend(mesh_points)
            
        keys = []
        for k in key_s_vals:
            val = min(max(round(k, 4), 0.0), round(L, 4))
            if val not in keys:
                keys.append(val)
        keys.sort()
        
        node_indices = []
        for s in keys:
            px, pz, _ = eval_seg_point(seg, s)
            node_indices.append(get_or_add_node(px, pz))
        
        m_idx = seg.get('master_idx', i)
        if m_idx < len(seg_sections):
            sec_props = seg_sections[m_idx]
        else:
            sec_props = seg_sections[0]
        
        for j in range(len(keys)-1):
            n1 = node_indices[j]
            n2 = node_indices[j+1]
            if n1 == n2: 
                continue 
                
            s_mid = (keys[j] + keys[j+1]) / 2.0
            _, _, th_mid = eval_seg_point(seg, s_mid)
            c_t = np.cos(th_mid)
            s_t = np.sin(th_mid)
            p_x1 = p_z1 = p_x2 = p_z2 = 0.0
            
            for ld in loads:
                if ld.get('seg_idx') == i and ld.get('type') != 'Point Load':
                    if ld['start'] - 1e-4 <= s_mid <= ld['end'] + 1e-4:
                        L_ld = max(ld['end'] - ld['start'], 1e-5)
                        wa = ld['w1'] + (ld['w2'] - ld['w1']) * (keys[j] - ld['start']) / L_ld
                        wb = ld['w1'] + (ld['w2'] - ld['w1']) * (keys[j+1] - ld['start']) / L_ld
                        
                        dir_str = ld.get('dir', 'Global Z (Vertical)').upper()
                        if 'Z' in dir_str or 'Y' in dir_str:
                            p_x1 += wa * s_t
                            p_z1 += wa * c_t
                            p_x2 += wb * s_t
                            p_z2 += wb * c_t
                        elif 'X' in dir_str:
                            p_x1 += wa * c_t
                            p_z1 -= wa * s_t
                            p_x2 += wb * c_t
                            p_z2 -= wb * s_t
                        else:
                            p_z1 += wa
                            p_z2 += wb
                        
            elements.append({
                'type': 'frame', 
                'group': 'segment', 
                'sec': sec_props['name'],
                'n1': n1, 
                'n2': n2, 
                'px1': p_x1, 
                'py1': p_z1, 
                'px2': p_x2, 
                'py2': p_z2, 
                'E': sec_props['E'] * 10000.0, 
                'A': sec_props['A'], 
                'I': sec_props['I'] / 100000000.0,
                'seg_idx': i, 
                'L': keys[j+1] - keys[j], 
                'th_mid': th_mid
            })
            
        for ld in loads:
            if ld.get('seg_idx') == i and ld.get('type') == 'Point Load':
                px, pz, th_pt = eval_seg_point(seg, ld['start'])
                nid = get_or_add_node(px, pz)
                dir_str = ld.get('dir', 'Global Z (Vertical)').upper()
                
                if 'Z' in dir_str or 'Y' in dir_str: 
                    nodal_loads.append({'node': nid, 'Fx': 0.0, 'Fz': ld['w1']})
                elif 'X' in dir_str: 
                    nodal_loads.append({'node': nid, 'Fx': ld['w1'], 'Fz': 0.0})
                else: 
                    c_pt = np.cos(th_pt)
                    s_pt = np.sin(th_pt)
                    nodal_loads.append({'node': nid, 'Fx': -ld['w1'] * s_pt, 'Fz': ld['w1'] * c_pt})

    return nodes, elements, nodal_loads, supports_list_out


def solve_fea_engine(nodes, elements, nodal_loads, supports_list):
    """
    محرك الـ FEA الدقيق 100% باستخدام مصفوفة الجساءة (Stiffness Method).
    """
    NDOF = len(nodes) * 3
    K = np.zeros((NDOF, NDOF))
    F = np.zeros(NDOF)
    
    for el in elements:
        n1 = el['n1']
        n2 = el['n2']
        x1 = nodes[n1][0]
        z1 = nodes[n1][1]
        x2 = nodes[n2][0]
        z2 = nodes[n2][1]
        
        L = np.hypot(x2 - x1, z2 - z1)
        if L < 1e-5: 
            continue
            
        c = (x2 - x1) / L
        s = (z2 - z1) / L
        el['L'] = L
        el['c'] = c
        el['s'] = s
        
        E = el['E']
        A = el['A']
        I = el.get('I', 0.00005)
        
        T = np.array([
            [c, s, 0, 0, 0, 0], 
            [-s, c, 0, 0, 0, 0], 
            [0, 0, 1, 0, 0, 0], 
            [0, 0, 0, c, s, 0], 
            [0, 0, 0, -s, c, 0], 
            [0, 0, 0, 0, 0, 1]
        ])
        
        if el['type'] == 'truss':
            k_loc = np.zeros((6, 6))
            k_loc[0, 0] = E * A / L
            k_loc[3, 3] = E * A / L
            k_loc[0, 3] = -E * A / L
            k_loc[3, 0] = -E * A / L
        else:
            k_loc = np.array([
                [E*A/L, 0, 0, -E*A/L, 0, 0], 
                [0, 12*E*I/L**3, 6*E*I/L**2, 0, -12*E*I/L**3, 6*E*I/L**2],
                [0, 6*E*I/L**2, 4*E*I/L, 0, -6*E*I/L**2, 2*E*I/L], 
                [-E*A/L, 0, 0, E*A/L, 0, 0],
                [0, -12*E*I/L**3, -6*E*I/L**2, 0, 12*E*I/L**3, -6*E*I/L**2], 
                [0, 6*E*I/L**2, 2*E*I/L, 0, -6*E*I/L**2, 4*E*I/L]
            ])
            
            px1 = el.get('px1', 0.0)
            py1 = el.get('py1', 0.0)
            px2 = el.get('px2', 0.0)
            py2 = el.get('py2', 0.0)
            
            f_loc = np.array([
                (2*px1 + px2) * L / 6.0, 
                (7*py1 + 3*py2) * L / 20.0, 
                (3*py1 + 2*py2) * L**2 / 60.0, 
                (px1 + 2*px2) * L / 6.0, 
                (3*py1 + 7*py2) * L / 20.0, 
                -(2*py1 + 3*py2) * L**2 / 60.0
            ])
            f_glob = T.T @ f_loc
            
            dof = [3*n1, 3*n1+1, 3*n1+2, 3*n2, 3*n2+1, 3*n2+2]
            for r in range(6): 
                F[dof[r]] += f_glob[r]
            
        k_glob = T.T @ k_loc @ T
        dof = [3*n1, 3*n1+1, 3*n1+2, 3*n2, 3*n2+1, 3*n2+2]
        
        for r in range(6):
            for col in range(6): 
                K[dof[r], dof[col]] += k_glob[r, col]
                
    for nl in nodal_loads:
        F[3*nl['node']] += nl.get('Fx', 0.0)
        F[3*nl['node']+1] += nl.get('Fz', nl.get('Fy', 0.0))
        
    net_load_z = abs(np.sum(F[1::3]))

    K_orig = K.copy()
    fixed_dofs = []
    K_pen = 1e12
    
    for sup in supports_list:
        n = sup['node']
        t = sup['type']
        a = sup.get('angle', 0.0)
        
        if t == 'Fixed': 
            fixed_dofs.extend([3*n, 3*n+1, 3*n+2])
        elif t == 'Hinged': 
            fixed_dofs.extend([3*n, 3*n+1])
        elif t == 'Roller':
            rad = np.radians(a)
            nx = -np.sin(rad)
            nz = np.cos(rad) 
            
            K[3*n, 3*n] += K_pen * nx**2
            K[3*n+1, 3*n+1] += K_pen * nz**2
            K[3*n, 3*n+1] += K_pen * nx * nz
            K[3*n+1, 3*n] += K_pen * nx * nz

    free_dof = []
    for i in range(NDOF):
        if i not in fixed_dofs:
            free_dof.append(i)
            
    K_ff = K[np.ix_(free_dof, free_dof)]
    F_f = F[free_dof]
    
    U = np.zeros(NDOF)
    try: 
        U[free_dof] = np.linalg.solve(K_ff, F_f)
    except np.linalg.LinAlgError: 
        U[free_dof] = np.linalg.lstsq(K_ff, F_f, rcond=None)[0]
    
    R_reactions = K_orig @ U - F 
    
    for el in elements:
        if el.get('L', 0) < 1e-5: 
            continue
            
        n1 = el['n1']
        n2 = el['n2']
        c = el['c']
        s = el['s']
        L = el['L']
        E = el['E']
        A = el['A']
        I = el.get('I', 0.00005)
        
        u_glob = U[[3*n1, 3*n1+1, 3*n1+2, 3*n2, 3*n2+1, 3*n2+2]]
        T = np.array([
            [c, s, 0, 0, 0, 0], 
            [-s, c, 0, 0, 0, 0], 
            [0, 0, 1, 0, 0, 0], 
            [0, 0, 0, c, s, 0], 
            [0, 0, 0, -s, c, 0], 
            [0, 0, 0, 0, 0, 1]
        ])
        u_loc = T @ u_glob
        el['internal'] = {'u_loc': u_loc}
        
        if el['type'] == 'truss':
            N_val = (E * A / L) * (u_loc[3] - u_loc[0])
            xs = np.linspace(0, L, 51)
            el['internal'].update({
                'N': np.full_like(xs, N_val), 
                'V': np.zeros_like(xs), 
                'M': np.zeros_like(xs), 
                'D': np.zeros_like(xs), 
                'x': xs
            })
        else:
            k_loc = np.array([
                [E*A/L, 0, 0, -E*A/L, 0, 0], 
                [0, 12*E*I/L**3, 6*E*I/L**2, 0, -12*E*I/L**3, 6*E*I/L**2], 
                [0, 6*E*I/L**2, 4*E*I/L, 0, -6*E*I/L**2, 2*E*I/L], 
                [-E*A/L, 0, 0, E*A/L, 0, 0], 
                [0, -12*E*I/L**3, -6*E*I/L**2, 0, 12*E*I/L**3, -6*E*I/L**2], 
                [0, 6*E*I/L**2, 2*E*I/L, 0, -6*E*I/L**2, 4*E*I/L]
            ])
            
            px1 = el.get('px1', 0.0)
            py1 = el.get('py1', 0.0)
            px2 = el.get('px2', 0.0)
            py2 = el.get('py2', 0.0)
            
            f_loc = np.array([
                (2*px1 + px2)*L/6.0, 
                (7*py1 + 3*py2)*L/20.0, 
                (3*py1 + 2*py2)*L**2/60.0, 
                (px1 + 2*px2)*L/6.0, 
                (3*py1 + 7*py2)*L/20.0, 
                -(2*py1 + 3*py2)*L**2/60.0
            ])
            
            f_end = k_loc @ u_loc - f_loc
            xs = np.linspace(0, L, 51)
            
            N_arr = np.zeros_like(xs)
            V_arr = np.zeros_like(xs)
            M_arr = np.zeros_like(xs)
            D_arr = np.zeros_like(xs)
            
            v1 = u_loc[1]
            th1 = u_loc[2]
            v2 = u_loc[4]
            th2 = u_loc[5]
            
            for i, x in enumerate(xs):
                N_arr[i] = -f_end[0] - (px1*x + (px2-px1)*x**2/(2*L))
                V_arr[i] = f_end[1] + (py1*x + (py2-py1)*x**2/(2*L))
                M_arr[i] = -f_end[2] + f_end[1]*x + py1*x**2/2.0 + (py2-py1)*x**3/(6*L)
                
                xi = x / L if L > 0 else 0
                N1_shp = 1.0 - 3.0*xi**2 + 2.0*xi**3
                N2_shp = L * (xi - 2.0*xi**2 + xi**3)
                N3_shp = 3.0*xi**2 - 2.0*xi**3
                N4_shp = L * (-xi**2 + xi**3)
                
                v_x = v1*N1_shp + th1*N2_shp + v2*N3_shp + th2*N4_shp
                w_avg = (py1 + py2) / 2.0
                v_load = (w_avg * x**2 * (L - x)**2) / (24.0 * E * I) if (E * I) != 0 else 0
                
                D_arr[i] = (v_x + v_load) * 1000.0 
                
            el['internal'].update({
                'N': N_arr, 
                'V': V_arr, 
                'M': M_arr, 
                'D': D_arr, 
                'x': xs
            })
            
    return U, R_reactions, net_load_z


# =====================================================================
# 🧠 THE BEAST OPTIMIZER (Strict Snapping, Forced Cantilevers, DFS)
# =====================================================================
def run_bridge_optimizer(base_segments, working_segments, active_seg_sections, ui_struts, ui_loads, target_rxn, spacings_str, auto_mesh_size, is_symmetric, opt_mode, combo_factors, status_text, progress_bar):
    """
    محسن الكباري الذكي (The Beast):
    يعتمد مسافات الإسناد وتعديلات النهايز ويطبق (Strict Bottom-Snapping).
    تمت إضافة حماية ضد الـ Timeout وإجبار الكابولي وتوقيع الدعامات على السولجر السفلي فقط.
    """
    try: 
        spacings_raw = spacings_str.split(',')
        spacings = []
        for x in spacings_raw:
            spacings.append(float(x.strip()))
        spacings.sort(reverse=True)
    except Exception as e: 
        return False, None, None, "❌ Format error in spacings."
    
    if not base_segments: 
        return False, None, None, "❌ No base segments found."
        
    # 💡 1. إيجاد أقل منسوب (Z-min) للوصول للسولجر السفلي الحقيقي وتجاهل العلوية
    min_z = min(min(seg['abs_p1'][1], seg['abs_p2'][1]) for seg in base_segments if seg.get('Shape Type') == 'Straight Line')
            
    bottom_xs = []
    for seg in base_segments:
        if seg.get('Shape Type') == 'Straight Line':
            p1 = seg['abs_p1']
            p2 = seg['abs_p2']
            if abs(p1[1] - min_z) < 0.5 or abs(p2[1] - min_z) < 0.5: 
                bottom_xs.extend([p1[0], p2[0]])
                
    if bottom_xs:
        soldier_min_x = min(bottom_xs)
        soldier_max_x = max(bottom_xs)
    else:
        soldier_min_x = min(min(seg['abs_p1'][0], seg['abs_p2'][0]) for seg in base_segments if seg.get('Shape Type') == 'Straight Line')
        soldier_max_x = max(max(seg['abs_p1'][0], seg['abs_p2'][0]) for seg in base_segments if seg.get('Shape Type') == 'Straight Line')
        
    def get_z_on_bottom_chord(test_x):
        """إسقاط ذكي رأسي لإيجاد قيمة Z على أسفل سولجر فقط (أفقي أو مائل) بصرامة"""
        best_z = 9999.0
        found = False
        for seg in base_segments:
            if seg.get('Shape Type') == 'Straight Line':
                p1 = seg['abs_p1']
                p2 = seg['abs_p2']
                min_px = min(p1[0], p2[0])
                max_px = max(p1[0], p2[0])
                
                if min_px - 0.20 <= test_x <= max_px + 0.20:
                    if abs(max_px - min_px) < 1e-5: 
                        z_val = min(p1[1], p2[1])
                    else:
                        ratio = (test_x - p1[0]) / (p2[0] - p1[0])
                        z_val = p1[1] + ratio * (p2[1] - p1[1])
                        
                    # الالتصاق التام بالسولجر السفلي
                    if abs(p1[1] - min_z) < 0.5 or abs(p2[1] - min_z) < 0.5:
                        best_z = min(best_z, z_val)
                        found = True
        return best_z if found else min_z 
            
    center_x = (soldier_min_x + soldier_max_x) / 2.0
    half_width = (soldier_max_x - soldier_min_x) / 2.0
    
    test_combined_loads = []
    for i, ld in enumerate(ui_loads):
        t_mode = ld.get('target_mode', 'Single Segment')
        target_base_indices = []
        
        if t_mode == "Single Segment": 
            target_base_indices.append(ld.get('seg_idx', 0))
        elif t_mode == "Multiple Segments": 
            target_base_indices.extend(ld.get('target_segs_idx', []))
        else: 
            target_base_indices.extend(range(len(base_segments)))
            
        cat = ld.get('category', 'Dead Load')
        fac = combo_factors.get(cat, 1.0)
        w1 = ld.get('w1', 0.0) * fac
        w2 = ld.get('w2', ld.get('w1', 0.0)) * fac if ld.get('type') == 'Trapezoidal' else w1
        loc_m = ld.get('loc', 0.0)
        
        target_working_indices = []
        for w_idx, w_seg in enumerate(working_segments):
            if w_seg.get('master_idx', 0) in target_base_indices:
                target_working_indices.append(w_idx)
        
        for s_idx_num in target_working_indices:
            w_len = float(working_segments[s_idx_num].get('L', 0.0))
            start_val = min(loc_m, w_len) if ld.get('type') == 'Point Load' else 0.0
            end_val = start_val if ld.get('type') == 'Point Load' else w_len
                
            test_combined_loads.append({
                'seg_idx': s_idx_num, 
                'category': cat, 
                'type': ld.get('type', 'Uniform'), 
                'dir': ld.get('dir', 'Global Z (Vertical)'), 
                'start': start_val, 
                'end': end_val, 
                'w1': w1, 
                'w2': w2
            })

    def run_trial(test_supps, dynamic_struts):
        nodes_t, elems_t, nloads_t, slist_t = build_chain_mesh(
            working_segments, active_seg_sections, test_combined_loads, 
            dynamic_struts, test_supps, [], mesh_size=auto_mesh_size
        )
        
        U, R, net_load = solve_fea_engine(nodes_t, elems_t, nloads_t, slist_t)
        
        ry_list = [R[3*sup['node']+1] for sup in slist_t]
            
        max_ry = max(ry_list) if ry_list else 0
        min_ry = min(ry_list) if ry_list else 0
        
        soldier_safe = True
        for i, sec in enumerate(active_seg_sections):
            for el in elems_t:
                if el.get('group') == 'segment' and el.get('seg_idx') == i:
                    m_arr = el.get('internal', {}).get('M', [0])
                    v_arr = el.get('internal', {}).get('V', [0])
                    # التأكد الصارم من أمان السولجر (عزم وقص)
                    if np.max(np.abs(m_arr)) > sec['Mall'] or np.max(np.abs(v_arr)) > sec['Qall']: 
                        soldier_safe = False
                        break
                        
        struts_safe = True
        upgraded_struts = []
        for el in elems_t:
            if el['type'] == 'truss':
                N_max = np.max(np.abs(el.get('internal', {}).get('N', [0])))
                # التأكد الصارم من أمان النهايز بناء على القائمة المفلترة
                opt_sec = get_optimal_strut_section(el.get('L', 0.0), N_max)
                if not opt_sec: 
                    struts_safe = False
                    upgraded_struts.append(el.get('sec'))
                else: 
                    upgraded_struts.append(opt_sec)
        
        return max_ry, min_ry, net_load, soldier_safe, struts_safe, upgraded_struts

    dummy_supps = [
        {'x': soldier_min_x, 'z': min_z, 'type': 'Hinged', 'angle': 0.0}, 
        {'x': soldier_max_x, 'z': min_z, 'type': 'Hinged', 'angle': 0.0}
    ]
    _, _, total_system_load, _, _, _ = run_trial(dummy_supps, ui_struts)
    
    if target_rxn > 1e-4:
        min_required_props = max(2, int(math.ceil(total_system_load / target_rxn))) 
    else:
        min_required_props = 2

    # 💡 3. محرك بناء الشبكات (DFS) مع إجبار الكوابيل على الأطراف وتأمين التماثل
    valid_grids = []
    if is_symmetric:
        def build_sym_grids(current_grid):
            # الكابولي إجباري يمين ويسار الشبكة بين 0.15م و 1.50م
            cantilever = half_width - current_grid[-1]
            if 0.15 <= cantilever <= 1.50:
                full_grid = set(current_grid)
                for x in current_grid:
                    if x > 1e-4: full_grid.add(-x)
                        
                sym_coords = []
                for x in sorted(list(full_grid)):
                    cx = round(center_x + x, 3)
                    cx = max(soldier_min_x, min(soldier_max_x, cx)) # Clamp لضمان البقاء داخل الحدود
                    sym_coords.append(cx)
                valid_grids.append(tuple(sym_coords))
                
            if cantilever < 0.15: return
                
            for s in spacings: 
                build_sym_grids(current_grid + [current_grid[-1] + s])
                
        build_sym_grids([0.0])
        for s in spacings: 
            build_sym_grids([s / 2.0])
            
    else:
        def build_asym_grids(current_grid):
            # إجبار وجود كابولي يمين السولجر السفلي
            cantilever_right = soldier_max_x - current_grid[-1]
            if 0.15 <= cantilever_right <= 1.50:
                asym_coords = [round(x, 3) for x in current_grid if soldier_min_x - 0.05 <= x <= soldier_max_x + 0.05]
                valid_grids.append(tuple(asym_coords))
                
            if cantilever_right < 0.15: return
                
            for s in spacings: 
                build_asym_grids(current_grid + [current_grid[-1] + s])
                
        # إجبار وجود كابولي يسار السولجر السفلي كبداية (0.15م لـ 1.50م)
        cantilever_opts = np.arange(0.15, 1.51, 0.10)
        for lc in cantilever_opts: 
            build_asym_grids([soldier_min_x + lc])

    filtered_grids = [list(g) for g in set(valid_grids) if len(g) >= min_required_props]
            
    if not filtered_grids: 
        return False, None, None, f"❌ Impossible to optimize! Requires at least {min_required_props} props while respecting cantilevers."

    grids_by_props = {}
    for g in filtered_grids:
        p_count = len(g)
        if p_count not in grids_by_props: grids_by_props[p_count] = []
        grids_by_props[p_count].append(g)

    shift_options = [0.0, 0.10, -0.10, 0.20, -0.20]
    
    # 💡 4. حماية ضد التهنيج (Timeout Logic)
    if "Quick" in opt_mode:
        max_time = 180.0  # 3 دقائق للبحث السريع
    else:
        max_time = 600.0  # 10 دقائق للبحث العميق
        
    start_time = time.time()
    
    best_fallback_grid = None
    best_fallback_struts = ui_struts
    best_fallback_score = 999999.0
    
    trials_count = 0
    total_estimated_trials = len(filtered_grids) * len(shift_options)
    sorted_p_keys = sorted(list(grids_by_props.keys()))
    
    timeout_reached = False
    
    for p_count in sorted_p_keys:
        if timeout_reached: break
        
        for actual_coords in grids_by_props[p_count]:
            if time.time() - start_time > max_time: 
                timeout_reached = True
                break
            
            if is_symmetric:
                cantilever_L = half_width - (max(actual_coords) - center_x) 
            else:
                cantilever_L = soldier_max_x - max(actual_coords)
                
            excluded_zone_start = soldier_max_x - (cantilever_L / 3.0)
            excluded_zone_start_left = soldier_min_x + (cantilever_L / 3.0)
            
            # توقيع الدعامات على السولجر السفلي الملتصق حصرياً
            test_supps = []
            for gx in actual_coords:
                gy = get_z_on_bottom_chord(gx)
                test_supps.append({'x': gx, 'z': round(gy, 3), 'type': 'Hinged', 'angle': 0.0})
                
            for shift_val in shift_options:
                if time.time() - start_time > max_time: 
                    timeout_reached = True
                    break
                
                shifted_struts = []
                for strut in ui_struts:
                    new_strut = strut.copy()
                    nz_b_old = strut.get('bz', strut.get('by', 0.0))
                    nz_t_old = strut.get('tz', strut.get('ty', 0.0))
                    
                    nx_b, nz_b = get_shifted_coords_along_segment(strut['bx'], nz_b_old, shift_val, base_segments)
                    nx_t, nz_t = get_shifted_coords_along_segment(strut['tx'], nz_t_old, shift_val, base_segments)
                    
                    if nx_b > excluded_zone_start or nx_b < excluded_zone_start_left: 
                        nx_b, nz_b, nx_t, nz_t = strut['bx'], nz_b_old, strut['tx'], nz_t_old
                        
                    new_sec = get_valid_strut_names()[0] if STRUTS_DB else "Unknown"
                        
                    new_strut.update({'bx': nx_b, 'bz': nz_b, 'tx': nx_t, 'tz': nz_t, 'sec': new_sec})
                    shifted_struts.append(new_strut)
                
                max_ry, min_ry, _, soldier_safe, struts_safe, upg_secs = run_trial(test_supps, shifted_struts)
                
                if not struts_safe and len(upg_secs) == len(shifted_struts):
                    for idx_st in range(len(shifted_struts)): 
                        shifted_struts[idx_st]['sec'] = upg_secs[idx_st]
                    max_ry, min_ry, _, soldier_safe, struts_safe, _ = run_trial(test_supps, shifted_struts) 
                
                trials_count += 1
                if trials_count % 15 == 0:
                    ratio_prog = min(1.0, trials_count / float(total_estimated_trials))
                    progress_bar.progress(ratio_prog)
                    status_text.markdown(f"**⏳ Search:** Grid **{p_count} Props** | Best Rxn So Far: **{best_fallback_score:.2f} kN**")
                
                # النجاح المثالي والتأكد أن كل العناصر آمنة 100%
                if max_ry <= target_rxn and min_ry >= 0.5 and soldier_safe and struts_safe:
                    progress_bar.progress(1.0)
                    status_text.empty()
                    return True, test_supps, shifted_struts, f"✅ BOOM! Safe Grid Found: Max Rxn = {max_ry:.2f} kN. Props = {p_count}."
                    
                # تخزين أفضل نتيجة كخطة بديلة (Fallback)
                if max_ry < best_fallback_score and soldier_safe:
                    best_fallback_score = max_ry
                    best_fallback_grid = test_supps
                    best_fallback_struts = shifted_struts
                        
    # 💡 التعديل السحري: إرجاع أفضل حل تم التوصل إليه دائماً حتى لو الوقت خلص
    progress_bar.empty()
    status_text.empty()
    
    if best_fallback_grid:
        msg = f"⚠️ Best possible solution applied: Max Rxn = {best_fallback_score:.2f} kN. Please review visually."
        # نرجعه كـ True عشان الواجهة تقبله وتوقعه على الرسم
        return True, best_fallback_grid, best_fallback_struts, msg
            
    return False, None, None, f"❌ Failed! Cannot satisfy basic stability (Uplift)."
# =========================================================
# 3. Plotting Engine & Word Report Generator
# =========================================================
def draw_base_geometry(ax, nodes, elements, supports_list, segments, show_names=False, show_dimensions=False):
    """
    رسم الهندسة الأساسية للمنشأ (قطاعات، نهايز، ركائز، أسماء، ومسافات).
    """
    for el in elements:
        n1 = nodes[el['n1']]
        n2 = nodes[el['n2']]
        if el['type'] == 'truss':
            ax.plot([n1[0], n2[0]], [n1[1], n2[1]], color='red', linestyle='-', linewidth=0.8, zorder=1)
        else:
            ax.plot([n1[0], n2[0]], [n1[1], n2[1]], color='royalblue', linestyle='-', linewidth=1.5, zorder=1)
            
    # 💡 ترتيب الدعامات من اليسار إلى اليمين بناءً على إحداثي X لترقيمها (J1, J2, ...)
    sorted_sups = sorted(supports_list, key=lambda s: nodes[s['node']][0])
    
    for i, sup in enumerate(sorted_sups):
        x = nodes[sup['node']][0]
        z = nodes[sup['node']][1]
        t = sup['type']
        ang_rad = math.radians(sup.get('angle', 0.0))
        c_a = math.cos(ang_rad)
        s_a = math.sin(ang_rad)
        
        def rot(px, pz): 
            return x + (px - x)*c_a - (pz - z)*s_a, z + (px - x)*s_a + (pz - z)*c_a
        
        if t == 'Fixed':
            ax.add_patch(Polygon([rot(x-0.1, z-0.1), rot(x+0.1, z-0.1), rot(x+0.1, z+0.1), rot(x-0.1, z+0.1)], facecolor='none', edgecolor='limegreen', lw=1.0, zorder=5))
            l1, l2 = rot(x-0.1, z), rot(x+0.1, z)
            ax.plot([l1[0], l2[0]], [l1[1], l2[1]], color='limegreen', lw=1.0, zorder=4)
        elif t == 'Hinged':
            ax.add_patch(Polygon([rot(x, z), rot(x+0.12, z-0.15), rot(x-0.12, z-0.15)], facecolor='none', edgecolor='limegreen', lw=1.0, zorder=5))
            l1, l2 = rot(x-0.17, z-0.15), rot(x+0.17, z-0.15)
            ax.plot([l1[0], l2[0]], [l1[1], l2[1]], color='limegreen', lw=1.0, zorder=4)
        elif t == 'Roller':
            ax.add_patch(Polygon([rot(x, z), rot(x+0.12, z-0.15), rot(x-0.12, z-0.15)], facecolor='none', edgecolor='limegreen', lw=1.0, zorder=5))
            ax.add_patch(plt.Circle(rot(x, z-0.19), 0.04, facecolor='none', edgecolor='limegreen', lw=1.0, zorder=5))
            l1, l2 = rot(x-0.17, z-0.23), rot(x+0.17, z-0.23)
            ax.plot([l1[0], l2[0]], [l1[1], l2[1]], color='limegreen', lw=1.0, zorder=4)

        # 💡 رسم المسمى (J1, J2) تحت الدعامة الخضراء مباشرة
        if show_names:
            lbl_x, lbl_z = rot(x, z - 0.45)
            ax.text(lbl_x, lbl_z, f"J{i+1}", color='green', fontsize=8, fontweight='bold', ha='center', va='center', zorder=10)

    if show_names and segments:
        for i, seg in enumerate(segments):
            mx, mz, mth = eval_seg_point(seg, seg.get('L', 0)/2.0)
            rot_deg = math.degrees(mth)
            if rot_deg > 90: 
                rot_deg -= 180
            elif rot_deg < -90: 
                rot_deg += 180
                
            clean_name = seg.get('name', f"S{i+1}").split('-')[0]
            ax.text(
                mx - math.sin(mth)*0.3, mz + math.cos(mth)*0.3, 
                clean_name, color='dimgray', fontsize=8, 
                ha='center', va='center', rotation=rot_deg, fontname='Arial'
            )

    if show_dimensions and len(supports_list) > 1:
        sup_xs = []
        for sup in supports_list:
            sup_xs.append(nodes[sup['node']][0])
            
        sup_xs = sorted(list(set(sup_xs)))
        if len(sup_xs) > 1:
            dim_z = min([nodes[sup['node']][1] for sup in supports_list]) - 0.85
            ax.plot([sup_xs[0], sup_xs[-1]], [dim_z, dim_z], color='gray', lw=0.6, zorder=1)
            for i in range(len(sup_xs)):
                ax.plot([sup_xs[i], sup_xs[i]], [dim_z - 0.1, dim_z + 0.1], color='gray', lw=0.6, zorder=1)
                if i < len(sup_xs) - 1:
                    dist = sup_xs[i+1] - sup_xs[i]
                    mid_x = (sup_xs[i] + sup_xs[i+1]) / 2.0
                    ax.text(mid_x, dim_z + 0.05, f"{dist:.2f}m", color='gray', fontsize=7, ha='center')


def get_live_preview_image(nodes, elements, supports_list, loads, segments, cat_filter=None):
    """
    رسم الهندسة بالإضافة إلى مضلعات الأحمال بشكل منفصل.
    """
    apply_plot_styles()
    fig_ld, ax_ld = plt.subplots(figsize=(9, 5.5))
    ax_ld.set_aspect('equal', adjustable='datalim')
    ax_ld.axis('off')
    
    draw_base_geometry(ax_ld, nodes, elements, supports_list, segments, show_names=True, show_dimensions=True)
    
    color_map = {
        'Dead Load': ('blue', 0.15), 
        'Live Load': ('red', 0.15), 
        'Wind Load': ('green', 0.20)
    }
    
    for ld in loads:
        if cat_filter and ld.get('category', 'Dead Load') != cat_filter:
            continue
            
        i = ld.get('seg_idx', 0)
        if i >= len(segments): 
            continue
            
        w1 = ld.get('w1', 0.0)
        w2 = ld.get('w2', 0.0)
        
        if abs(w1) < 1e-4 and abs(w2) < 1e-4: 
            continue
        
        ld_color, ld_alpha = color_map.get(ld.get('category', 'Dead Load'), ('blue', 0.15))
        
        num_pts = max(10, int((ld.get('end', 0) - ld.get('start', 0)) / 0.1))
        s_vals = np.linspace(ld.get('start', 0), ld.get('end', 0), num_pts)
        
        poly_pts = []
        top_pts = []
        
        for sv in s_vals:
            px, pz, th = eval_seg_point(segments[i], sv)
            L_load = max(1e-5, (ld.get('end', 0) - ld.get('start', 0)))
            w_val = (w1 + (w2 - w1) * (sv - ld.get('start', 0)) / L_load) * 0.05
            
            poly_pts.append((px, pz))
            
            dir_str = ld.get('dir', 'Global Z (Vertical)').upper()
            if 'Z' in dir_str or 'Y' in dir_str: 
                top_pts.append((px, pz + abs(w_val))) 
            elif 'X' in dir_str: 
                top_pts.append((px - w_val, pz))
            else:
                c = math.cos(th)
                s = math.sin(th)
                top_pts.append((px + s * abs(w_val), pz + c * abs(w_val)))
                
        poly_pts.extend(top_pts[::-1])
        
        if len(poly_pts) > 2:
            ax_ld.add_patch(Polygon(poly_pts, facecolor=ld_color, edgecolor='none', alpha=ld_alpha, zorder=2))
            ax_ld.plot([p[0] for p in top_pts], [p[1] for p in top_pts], color=ld_color, lw=0.8, zorder=3)
            
            num_arrows = max(3, len(top_pts) // 2)
            step = max(1, len(top_pts) // num_arrows)
            for k in range(0, len(top_pts), step):
                dx = poly_pts[k][0] - top_pts[k][0]
                dz = poly_pts[k][1] - top_pts[k][1]
                ax_ld.arrow(
                    top_pts[k][0], top_pts[k][1], dx, dz, 
                    length_includes_head=True, head_width=0.06, head_length=0.1, 
                    fc=ld_color, ec=ld_color, lw=0.6, zorder=4
                )
            
    return safe_render_fig(fig_ld)


def plot_sap2000_diagrams(nodes, elements, R_reactions, scales, supports_list, loads, segments):
    """توليد وتغليف كافة الرسومات التحليلية المطلوبة مع عزل كل فئة وجدولة ترتيبها."""
    apply_plot_styles()
    figs_dict = {}
    
    # رسومات الأحمال المعزولة
    has_dl = any(ld.get('category') == 'Dead Load' for ld in loads if abs(ld.get('w1', 0)) > 1e-4)
    has_ll = any(ld.get('category') == 'Live Load' for ld in loads if abs(ld.get('w1', 0)) > 1e-4)
    has_wl = any(ld.get('category') == 'Wind Load' for ld in loads if abs(ld.get('w1', 0)) > 1e-4)
    
    if has_dl: figs_dict['DL'] = get_live_preview_image(nodes, elements, supports_list, loads, segments, cat_filter='Dead Load')
    if has_ll: figs_dict['LL'] = get_live_preview_image(nodes, elements, supports_list, loads, segments, cat_filter='Live Load')
    if has_wl: figs_dict['WL'] = get_live_preview_image(nodes, elements, supports_list, loads, segments, cat_filter='Wind Load')
    
    # رسم ردود الأفعال (Reactions)
    fig_r, ax_r = plt.subplots(figsize=(9, 5.5))
    ax_r.set_aspect('equal', adjustable='datalim')
    ax_r.axis('off')
    
    draw_base_geometry(ax_r, nodes, elements, supports_list, segments, show_names=True, show_dimensions=True)
    
    for sup in supports_list:
        n = sup['node']
        ang = math.radians(sup.get('angle', 0.0))
        Rx = R_reactions[3*n]
        Rz = R_reactions[3*n+1]
        x = nodes[n][0]
        z = nodes[n][1]
        
        R_loc_x = Rx * math.cos(ang) + Rz * math.sin(ang)
        R_loc_z = -Rx * math.sin(ang) + Rz * math.cos(ang)
        
        if sup['type'] == 'Roller': 
            draw_reaction_arrow(ax_r, x, z, R_loc_z, -math.sin(ang), math.cos(ang))
        else:
            draw_reaction_arrow(ax_r, x, z, R_loc_x, math.cos(ang), math.sin(ang))
            draw_reaction_arrow(ax_r, x, z, R_loc_z, -math.sin(ang), math.cos(ang))
            
    figs_dict['R'] = safe_render_fig(fig_r)
    
    # دالة فرعية لرسومات الـ N, V, M
    def create_force_plot(val_key, scale, c_pos, c_neg):
        fig_f, ax_f = plt.subplots(figsize=(9, 5.5))
        ax_f.set_aspect('equal', adjustable='datalim')
        ax_f.axis('off')
        
        draw_base_geometry(ax_f, nodes, elements, supports_list, segments, show_names=True)
        
        for el in elements:
            n1 = el['n1']
            n2 = el['n2']
            x1 = nodes[n1][0]
            z1 = nodes[n1][1]
            x2 = nodes[n2][0]
            z2 = nodes[n2][1]
            c = el.get('c', 1.0)
            s = el.get('s', 0.0)
            xs = el.get('internal', {}).get('x', [])
            vals = el.get('internal', {}).get(val_key, [])
            
            if len(vals) == 0 or np.all(np.abs(vals) < 1e-6): 
                continue
                
            plot_vals = -vals if val_key != 'N' else vals
            px = x1 - s * plot_vals * scale + c * xs
            pz = z1 + c * plot_vals * scale + s * xs
            
            for k in range(len(px)-1):
                ax_f.plot([px[k], px[k+1]], [pz[k], pz[k+1]], color=c_pos if vals[k] >= 0 else c_neg, lw=0.8)
                
            ax_f.plot([x1, px[0]], [z1, pz[0]], color=c_pos if vals[0]>=0 else c_neg, lw=0.8)
            ax_f.plot([x2, px[-1]], [z2, pz[-1]], color=c_pos if vals[-1]>=0 else c_neg, lw=0.8)
            
            mid = len(vals)//2
            if abs(vals[mid]) > 0.1: 
                ax_f.text(px[mid], pz[mid], f"{vals[mid]:+.2f}", fontsize=6, color=c_pos if vals[mid]>=0 else c_neg, ha='center', va='center')
                
        return safe_render_fig(fig_f)

    figs_dict['N'] = create_force_plot('N', scales['N'], 'blue', 'red')
    figs_dict['V'] = create_force_plot('V', scales['V'], 'blue', 'red')
    figs_dict['M'] = create_force_plot('M', scales['M'], 'blue', 'red')
    
    return figs_dict


def set_cell_bg_and_font(cell, bg_color_hex, font_color, bold=False):
    """مساعد برمجي لتلوين وتنسيق خلايا جدول الوورد بدقة Arial 12"""
    tcPr = cell._element.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), bg_color_hex)
    tcPr.append(shd)
    
    for p in cell.paragraphs:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for r in p.runs:
            r.font.name = 'Arial'
            r.font.size = Pt(12)
            r.font.bold = bold
            r.font.color.rgb = font_color


def generate_multi_case_report(cases_data, proj_info):
    """
    توليد النوتة الحسابية المعتمدة بالشكل القياسي للاستشاري (شاملة الغلاف والفهرس والمواصفات).
    تستخدم نفس أسلوب تقارير الأسقف مع دمج حسابات كباري الـ DXF وتنقيتها.
    """
    import fitz  
    
    if os.path.exists("Acrow_Template.docx"):
        doc = Document("Acrow_Template.docx")
    else:
        doc = Document()
        
    def force_ltr_left(p):
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        pPr = p._element.get_or_add_pPr()
        bidi = OxmlElement('w:bidi')
        bidi.set(qn('w:val'), '0')
        pPr.append(bidi)

    def add_line(text, bold=False, size=12, italic=False, color=None, align='left'):
        p = doc.add_paragraph()
        if align == 'center':
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        else:
            force_ltr_left(p)
        r = p.add_run(text)
        r.font.name = 'Arial'
        r.font.size = Pt(size)
        r.font.bold = bold
        r.font.italic = italic
        r.font.rtl = False
        if color:
            r.font.color.rgb = color
        return p

    # ==========================================
    # 1. COVER PAGE & REPLACEMENTS
    # ==========================================
    proj_name = proj_info.get("proj_name", "")
    contractor = proj_info.get("contractor", "")
    calc_sub = proj_info.get("calc_sub", "")
    sys_name = proj_info.get("sys_name", "Acrow Bridge Systems")
    proj_no = proj_info.get("proj_no", "")
    date_val = proj_info.get("date_val", "")
    calc_by = proj_info.get("calc_by", "")
    chk_by = proj_info.get("chk_by", "")
    cover_img = proj_info.get("cover_img", "")
    data_sheets = proj_info.get("data_sheets", [])
    ref_code = proj_info.get("ref_code", "SBC")

    def remove_hardcoded_prefix(p):
        if p.text and "CALCULATION SHEET FOR" in p.text.upper():
            clean_text = re.sub(r'(?i)CALCULATION SHEET FOR\s*', '', p.text)
            if p.runs:
                font_name, font_size, font_bold = p.runs[0].font.name, p.runs[0].font.size, p.runs[0].font.bold
                font_color = p.runs[0].font.color.rgb if p.runs[0].font.color else None
                for r in p.runs: r.text = ""
                p.runs[0].text = clean_text
                p.runs[0].font.name, p.runs[0].font.size, p.runs[0].font.bold = font_name, font_size, font_bold
                if font_color: p.runs[0].font.color.rgb = font_color
            else:
                p.text = clean_text

    for p in doc.paragraphs: 
        remove_hardcoded_prefix(p)
        
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                for p in cell.paragraphs: 
                    remove_hardcoded_prefix(p)

    replacements = {
        "[PROJECT_NAME]": proj_name,
        "[CONTRACTOR]": contractor,
        "[CALC_SUBJECT]": calc_sub,
        "[SYSTEM_NAME]": sys_name,
        "[PROJ_NO]": proj_no,
        "[DATE]": date_val,
        "[CALC_BY]": calc_by,
        "[CHK_BY]": chk_by,
        "[REV]": "00"
    }

    for p in doc.paragraphs:
        if "[COVER_IMAGE]" in p.text:
            for r in p.runs: r.text = r.text.replace("[COVER_IMAGE]", "")
            if cover_img and cover_img != "No images found." and os.path.exists(cover_img):
                p.add_run().add_picture(cover_img, width=Cm(15.0))
                
        for k, v in replacements.items():
            if k in p.text:
                for r in p.runs:
                    if k in r.text: r.text = r.text.replace(k, str(v))
                if k in p.text: p.text = p.text.replace(k, str(v))

    for sec in doc.sections:
        for hf in [sec.header, sec.first_page_header, sec.footer, sec.first_page_footer]:
            if hf:
                for tbl in hf.tables:
                    for row in tbl.rows:
                        for cell in row.cells:
                            for p in cell.paragraphs:
                                for k, v in replacements.items():
                                    if k in p.text:
                                        for r in p.runs:
                                            if k in r.text: r.text = r.text.replace(k, str(v))
                                        if k in p.text: p.text = p.text.replace(k, str(v))
                                        for r in p.runs: r.font._element.set(qn('w:ascii'), 'Arial')

    doc.add_page_break()

    # ==========================================
    # 2. INDEX PAGE
    # ==========================================
    def get_pdf_page_count_safe(pdf_path):
        try:
            pdf_d = fitz.open(pdf_path)
            c = len(pdf_d)
            pdf_d.close()
            return c
        except:
            return 1

    ds_page_map = {}
    current_page = 4 
    
    if data_sheets:
        for f in data_sheets:
            if os.path.exists(f):
                p_count = get_pdf_page_count_safe(f)
                bname = os.path.basename(f).replace('.pdf', '')
                ds_page_map[bname] = f"{current_page}" if p_count == 1 else f"{current_page}-{current_page + p_count - 1}"
                current_page += p_count
                    
    design_pdf = "Design_Loads_BS.pdf" if "BS" in ref_code and os.path.exists("Design_Loads_BS.pdf") else ("Design_Loads_ACI.pdf" if "ACI" in ref_code and os.path.exists("Design_Loads_ACI.pdf") else None)
    design_loads_page = ""
    
    if design_pdf: 
        p_count = get_pdf_page_count_safe(design_pdf)
        design_loads_page = f"{current_page}" if p_count == 1 else f"{current_page}-{current_page + p_count - 1}"
        current_page += p_count

    insert_blue_banner(doc, "INDEX OF CONTENTS", font_size=16)
    
    def add_idx_head(text):
        p = doc.add_paragraph()
        force_ltr_left(p)
        p.paragraph_format.line_spacing = 1.0
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(4)
        r = p.add_run(text)
        r.font.name = 'Arial'
        r.font.size = Pt(14)
        r.bold = True
        r.underline = True

    def add_idx_item(text):
        p = doc.add_paragraph()
        force_ltr_left(p)
        p.paragraph_format.line_spacing = 1.0
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(text)
        r.font.name = 'Arial'
        r.font.size = Pt(11)

    if ds_page_map:
        add_idx_head("1. Formwork Materials Technical Data:")
        idx_c = 1
        for k, v in ds_page_map.items():
            clean_k = re.sub(r'(?i)data\s*sheet\s*for\s*', '', k).strip()
            clean_k = re.sub(r'(?i)datasheet\s*for\s*', '', clean_k).strip()
            add_idx_item(f"   1.{idx_c} Data Sheet for {clean_k} ........................................ Page {v}")
            idx_c += 1
        
    if design_loads_page:
        add_idx_head("2. Design Loads of Bridge Deck Slab:")
        add_idx_item(f"   2.1 Design Loads .................................................... Page {design_loads_page}")
        
    add_idx_head("3. Design Data of Formwork:")
    for i_idx, case in enumerate(cases_data):
        add_idx_item(f"   3.{i_idx+1} {case['title']}")
            
    doc.add_page_break()

    # ==========================================
    # 3. REGULATIONS & DATA SHEETS
    # ==========================================
    insert_blue_banner(doc, "REGULATIONS AND STANDARDS", font_size=16)
    doc.add_paragraph()
    if "BS" in ref_code: 
        add_eq(doc, "1- BS 5975-1996: FORMWORK FOR CONCRETE")
        add_eq(doc, "2- BS 5975-2008: FORMWORK FOR CONCRETE")
        add_eq(doc, "3- FORMWORK A GUIDE TO A GOOD PRACTICE")
        add_eq(doc, "4- WISA®-FORM PLYWOOD.")
        add_eq(doc, "5- THE SAUDI BUILDING CODE (SBC) 2024")
    else: 
        add_eq(doc, "1- ACI 347R-14 ....... GUIDE TO FORMWORK FOR CONCRETE.")
        add_eq(doc, "2- ACI SP-4 ......... FORMWORK FOR CONCRETE.")
        add_eq(doc, "3- WISA®-FORM PLYWOOD.")
        add_eq(doc, "4- THE SAUDI BUILDING CODE (SBC) 2024")
    
    if data_sheets:
        doc.add_page_break()
        insert_blue_banner(doc, "FORMWORK MATERIALS TECHNICAL DATA", font_size=14)
        for f in data_sheets:
            if os.path.exists(f): 
                append_pdf_stream_to_word(f, doc, is_path=True, max_width_cm=17.5, max_height_cm=24.0, add_border=True, reduce_first_page=True)
    
    if design_pdf: 
        doc.add_page_break()
        insert_blue_banner(doc, "DESIGN LOADS OF BRIDGE DECK SLAB", font_size=14)
        append_pdf_stream_to_word(design_pdf, doc, is_path=True, max_width_cm=17.5, max_height_cm=24.0, add_border=True, reduce_first_page=True)

    # ==========================================
    # 5. DESIGN DATA OF FORMWORK
    # ==========================================
    doc.add_page_break()
    insert_blue_banner(doc, "DESIGN DATA OF FORMWORK", font_size=16)

    for case in cases_data:
        tbl_id_clean = str(case['title'].upper()).replace("TABLE", "").strip()
        add_line(f"Bridge deck slab table: Table T{tbl_id_clean}", bold=True, size=14)
        doc.add_paragraph()
        
        if 'calc_details' in case and len(case['calc_details']) > 0:
            add_line("LOAD CALCULATION DETAILS (Dead Load from DXF Areas):", bold=True, size=12)
            add_line("Equation Used: Load W (kN/m) = [Area (m2) × Density (kN/m3) × Loaded Width (m)] / Length (m)", bold=False, size=10, italic=True)
            
            table_ld = doc.add_table(rows=len(case['calc_details'])+1, cols=4)
            table_ld.alignment = WD_TABLE_ALIGNMENT.CENTER
            
            hdr_cells = table_ld.rows[0].cells
            headers = ["Segment", "Length (m)", "Area (m2)", "Load W (kN/m)"]
            
            for i, text in enumerate(headers):
                hdr_cells[i].text = text
                tcPr = hdr_cells[i]._element.get_or_add_tcPr()
                shd = OxmlElement('w:shd')
                shd.set(qn('w:val'), 'clear')
                shd.set(qn('w:color'), 'auto')
                shd.set(qn('w:fill'), '1F497D')
                tcPr.append(shd)
                for p in hdr_cells[i].paragraphs:
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for r in p.runs:
                        r.font.name = 'Arial'
                        r.font.size = Pt(12)
                        r.font.bold = True
                        r.font.color.rgb = RGBColor(255,255,255)
            
            for i, r_data in enumerate(case['calc_details']):
                row_cells = table_ld.rows[i+1].cells
                row_cells[0].text = str(r_data['segment'])
                row_cells[1].text = f"{r_data['length']:.2f}"
                row_cells[2].text = f"{r_data['area']:.2f}"
                row_cells[3].text = f"{r_data['load_w']:.2f}"
                for cell in row_cells:
                    tcPr = cell._element.get_or_add_tcPr()
                    shd = OxmlElement('w:shd')
                    shd.set(qn('w:fill'), 'F2F2F2')
                    tcPr.append(shd)
                    for p in cell.paragraphs:
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        for r in p.runs:
                            r.font.name = 'Arial'
                            r.font.size = Pt(12)
                            r.font.color.rgb = RGBColor(0,0,0)
                    
            doc.add_paragraph()
            
            # طباعة سطر الـ Live Load الصريح الخالي من الشوائب
            ll_w = 0.0
            for ld in case.get('loads', []):
                if ld.get('category') == 'Live Load':
                    ll_w = abs(ld.get('w1', 0.0))
                    break
            
            if ll_w > 0:
                current_tab_width = case.get('loaded_width_curr', 1.30)
                base_ll = ll_w / current_tab_width if current_tab_width > 0 else 0.0
                add_line("Live Load:", bold=True, size=12)
                add_line(f"- W1 = live load x Loaded Width for Load Calculation = {base_ll:.2f} x {current_tab_width:.2f} = {ll_w:.3f} kN/m.", size=12)
            
            doc.add_paragraph()

        # طباعة الدياجرامات الـ 7 بتنسيق هندسي شيك (أحمر، تحته خط، Not Bold)
        diagrams_map = [
            ('DL', "Dead Load Distribution Diagram:"),
            ('LL', "Live Load Distribution Diagram:"),
            ('WL', "Wind Load Distribution Diagram:"),
            ('M', "Bending Moment Diagram due to (DL+LL):"),
            ('N', "Normal Force Diagram due to (DL+LL):"),
            ('V', "Shear Force Diagram due to (DL+LL):"),
            ('R', "Reaction Diagram:")
        ]
        
        for key, label in diagrams_map:
            if key in case.get('img_bufs', {}):
                p_title = add_line(label, bold=False, size=12, color=RGBColor(192, 0, 0))
                if p_title.runs: 
                    p_title.runs[0].underline = True
                
                p_img = doc.add_paragraph()
                p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
                r_img = p_img.add_run()
                r_img.add_picture(io.BytesIO(case['img_bufs'][key]), width=Cm(16.5))
                doc.add_paragraph()
                
        doc.add_page_break()
                
    out = io.BytesIO()
    doc.save(out)
    return out
# =========================================================
# 6. Main Streamlit UI (Dynamic Extractor & Interactive Editor)
# =========================================================
def render_bridge_module(proj_info):
    """
    الواجهة الرئيسية للمحرك الهندسي التفاعلي، تشمل رفع الـ DXF 
    وبناء التابات المستقلة ببيانات معزولة، والأوبتيميزر، ووضع الإدخال اليدوي.
    """
    st.markdown("## 🌉 Bridge Formwork (True 2D DXF + Advanced Editor)")
    mode = st.radio(
        "Select Input Mode:", 
        ["1. Multi-Case DXF Auto-Extractor 🪄", "2. Single-Case Manual Builder 🛠️"], 
        horizontal=True
    )
    st.markdown("---")

    if "DXF" in mode:
        st.info("💡 **Smart Engine:** Upload DXF. Frames, Struts, and Geometry are extracted instantly in XZ plane!")
        
        # كثافة الخرسانة كمعامل عام للمشروع
        conc_density = st.number_input("Concrete Density (kN/m3)", value=25.0, step=0.5)
        
        uploaded_dxf = st.file_uploader("📥 Upload DXF File (.dxf)", type=['dxf'])
        
        if uploaded_dxf and st.button("🚀 Process DXF & Extract Data", type="primary", use_container_width=True):
            with st.spinner("Parsing DXF true 2D geometry, Mirroring Symmetry, & Extracting Areas..."):
                cases_data = parse_dxf_bridge_cases(uploaded_dxf.getvalue())
                
            if cases_data:
                st.session_state.bridge_cases = cases_data
                st.success(f"✅ Successfully extracted {len(cases_data)} structural case(s)!")
                st.rerun()
            else: 
                st.error("❌ Failed to parse DXF. Please ensure layers are correct (TABLE_ANALYSIS, FRAME, SUPPORT, TEXT_DATA).")

        if 'bridge_cases' in st.session_state:
            global_sec = {
                'name': "Soldier U100", 'E': 2100.0, 
                'A': 34.3 / 10000.0, 'I': 412.0 / 100000000.0, 
                'Mall': 13.1, 'Qall': 100.8
            }
            all_cases_ready = []
            
            tabs = st.tabs([c['title'] for c in st.session_state.bridge_cases])
            
            for c_idx, tab in enumerate(tabs):
                case = st.session_state.bridge_cases[c_idx]
                
                # تهيئة المتغيرات الخاصة بالتاب الحالي إن لم تكن موجودة
                if 'loaded_width_curr' not in case:
                    case['loaded_width_curr'] = 1.30
                if 'base_live_load' not in case:
                    case['base_live_load'] = 2.90
                    
                with tab:
                    c_edit, c_view = st.columns([1.2, 1.8])
                    
                    with c_edit:
                        st.markdown("### 📐 Table Specifications")
                        
                        # 💡 حفظ القيم القديمة قبل الودجت عشان نقدر نكتشف أي تغيير حصل فيها
                        prev_lw = case['loaded_width_curr']
                        prev_bll = case['base_live_load']
                        
                        cw1, cw2 = st.columns(2)
                        case['loaded_width_curr'] = cw1.number_input("Loaded Width (m) for this Table", value=float(case['loaded_width_curr']), step=0.05, key=f"lw_curr_{c_idx}")
                        case['base_live_load'] = cw2.number_input("Base Live Load (kN/m2)", value=float(case['base_live_load']), step=0.10, key=f"bll_curr_{c_idx}")
                        
                        # 💡 اكتشاف تغيير قيمة الـ Loaded Width أو الـ Base Live Load أوتوماتيك
                        values_changed = (abs(case['loaded_width_curr'] - prev_lw) > 1e-9) or (abs(case['base_live_load'] - prev_bll) > 1e-9)
                        
                        # 💡 زرار التحديث اليدوي (يفضل شغال كمان، مثلاً لو اتغير الـ Concrete Density)
                        manual_update_clicked = st.button("🔄 Update Auto-Loads", key=f"upd_auto_lds_{c_idx}", use_container_width=True)
                        
                        # 💡 يشتغل أوتوماتيك عند تغيير Loaded Width أو Base Live Load، أو يدوي عند الضغط ع الزرار
                        if manual_update_clicked or values_changed:
                            new_auto_loads = []
                            case['calc_details'] = []
                            
                            # 1. Dead Load
                            for area_item in case.get('dxf_areas', []):
                                s_name = area_item['segment']
                                w_val = (area_item['area'] * conc_density * case['loaded_width_curr']) / area_item['length']
                                matching_indices = [idx_m for idx_m, seg_m in enumerate(case['segments']) if seg_m['name'] == s_name]
                                
                                if matching_indices:
                                    t_mode = 'Single Segment' if len(matching_indices) == 1 else 'Multiple Segments'
                                    new_auto_loads.append({
                                        'seg_idx': matching_indices[0], 'category': 'Dead Load', 'type': 'Uniform', 
                                        'dir': 'Global Z (Vertical)', 'target_mode': t_mode, 'target_segs_idx': matching_indices, 
                                        'start': 0.0, 'end': case['segments'][matching_indices[0]]['L'], 
                                        'w1': -abs(w_val), 'w2': -abs(w_val), 'loc': 0.0, 'is_auto': True
                                    })
                                case['calc_details'].append({'segment': s_name, 'length': area_item['length'], 'area': area_item['area'], 'load_w': abs(w_val)})
                                
                            # 2. Live Load
                            valid_ll_indices = [idx for idx, seg in enumerate(case['segments']) if seg['name'] not in ["S30", "S31"]]
                            calculated_ll_magnitude = -abs(case['base_live_load'] * case['loaded_width_curr'])
                            if valid_ll_indices:
                                new_auto_loads.append({
                                    'seg_idx': valid_ll_indices[0], 'category': 'Live Load', 'type': 'Uniform', 
                                    'dir': 'Global Z (Vertical)', 'target_mode': 'Multiple Segments', 'target_segs_idx': valid_ll_indices, 
                                    'start': 0.0, 'end': case['segments'][valid_ll_indices[0]]['L'], 
                                    'w1': calculated_ll_magnitude, 'w2': calculated_ll_magnitude, 'loc': 0.0, 'is_auto': True
                                })
                            
                            manual_loads = [ld for ld in case.get('loads', []) if not ld.get('is_auto', False)]
                            case['loads'] = new_auto_loads + manual_loads
                            
                            # 💡 مسح الذاكرة المؤقتة للخانات عشان الواجهة تقرأ الأرقام الجديدة فوراً
                            for key in list(st.session_state.keys()):
                                if key.startswith(f"alw1_{c_idx}_") or key.startswith(f"alw2_{c_idx}_"):
                                    del st.session_state[key]
                                    
                            st.rerun()

                        st.markdown("### 🎛️ Global Load Factors")
                        c_f1, c_f2, c_f3 = st.columns(3)
                        fac_d = c_f1.number_input("DL Factor", value=1.00, step=0.1, key=f"f_d_{c_idx}")
                        fac_l = c_f2.number_input("LL Factor", value=1.00, step=0.1, key=f"f_l_{c_idx}")
                        fac_w = c_f3.number_input("WL Factor", value=1.00, step=0.1, key=f"f_w_{c_idx}")
                        combo_factors = {'Dead Load': fac_d, 'Live Load': fac_l, 'Wind Load': fac_w}
                        
                        # 💡 واجهة الدعامات المطابقة للتصميم المطلوب (J1, J2 مع زر ❌)
                        case['supports'].sort(key=lambda s: s['x'])
                        with st.expander(f"🔗 Edit Supports ({len(case['supports'])})", expanded=False):
                            for i, sup in enumerate(case['supports']):
                                st.markdown(f"**🟢 Support J{i+1}**")
                                c1, c2, c3, c4, c_del = st.columns([1, 1, 1.2, 1, 0.3])
                                sup['x'] = c1.number_input(f"J{i+1} X (m)", value=float(sup['x']), step=0.1, key=f"sx_{c_idx}_{i}")
                                sup['z'] = c2.number_input(f"J{i+1} Y (m)", value=float(sup.get('z', sup.get('y', 0.0))), step=0.1, key=f"sz_{c_idx}_{i}")
                                
                                type_opts = ["Hinged", "Roller", "Fixed"]
                                idx_type = type_opts.index(sup['type']) if sup['type'] in type_opts else 1
                                sup['type'] = c3.selectbox(f"J{i+1} Type", type_opts, index=idx_type, key=f"st_{c_idx}_{i}")
                                sup['angle'] = c4.number_input(f"J{i+1} Angle(°)", value=float(sup.get('angle',0.0)), step=15.0, key=f"sa_{c_idx}_{i}")
                                
                                c_del.markdown("<br>", unsafe_allow_html=True)
                                if c_del.button("❌", key=f"del_sup_{c_idx}_{i}"):
                                    case['supports'].pop(i)
                                    st.rerun()
                                    
                            if st.button("➕ Add Support", key=f"add_sup_{c_idx}"):
                                case['supports'].append({'x': 0.0, 'z': 0.0, 'type': 'Roller', 'angle': 0.0})
                                st.rerun()

                        if 'sec_overrides' not in case: 
                            case['sec_overrides'] = [global_sec.copy() for _ in range(len(case['segments']))]
                            
                        # 💡 توليد أسماء السكاشن المُميزة برقم الإندكس لحل مشكلة التكرار
                        seg_names = [s['name'] for s in case['segments']]
                        unique_seg_opts = [f"{idx} - {name}" for idx, name in enumerate(seg_names)]
                        
                        with st.expander("📏 Override Sections", expanded=False):
                            override_segs = st.multiselect("Select segments:", unique_seg_opts, key=f"ovr_seg_{c_idx}")
                            if override_segs:
                                rad_opt = st.radio("Override Profile:", ["Custom Section", "Acrow Beam S12"], key=f"ovr_rad_{c_idx}")
                                if rad_opt == "Custom Section":
                                    o1, o2, o3, o4 = st.columns(4)
                                    o_sec = {
                                        'name': "Custom", 'E': 2100.0, 
                                        'A': o1.number_input("A", value=50.0, key=f"oa_{c_idx}")/10000.0, 
                                        'I': o2.number_input("I", value=1200.0, key=f"oi_{c_idx}")/100000000.0, 
                                        'Mall': o3.number_input("Mall", value=30.0, key=f"om_{c_idx}"), 
                                        'Qall': o4.number_input("Qall", value=150.0, key=f"oq_{c_idx}")
                                    }
                                else:
                                    o_sec = {'name': "S12", 'E': 2100.0, 'A': 20.0/10000.0, 'I': 800.0/100000000.0, 'Mall': 15.0, 'Qall': 80.0}
                                    
                                for s_val in override_segs: 
                                    idx_seg = int(s_val.split(' - ')[0])
                                    case['sec_overrides'][idx_seg] = o_sec.copy()

                        # 💡 واجهة النهايز المفلترة مع أزرار المسح
                        with st.expander(f"📐 Edit Struts ({len(case['struts'])})", expanded=False):
                            strut_opts = get_valid_strut_names()
                            for i, stt in enumerate(case['struts']):
                                s1, s2, s3, s4, s5, s_del = st.columns([1,1,1,1,1.2,0.3])
                                stt['tx'] = s1.number_input("TX", value=float(stt['tx']), step=0.1, key=f"ttx_{c_idx}_{i}")
                                stt['tz'] = s2.number_input("TZ", value=float(stt.get('tz', stt.get('ty', 0.0))), step=0.1, key=f"tty_{c_idx}_{i}")
                                stt['bx'] = s3.number_input("BX", value=float(stt['bx']), step=0.1, key=f"tbx_{c_idx}_{i}")
                                stt['bz'] = s4.number_input("BZ", value=float(stt.get('bz', stt.get('by', 0.0))), step=0.1, key=f"tby_{c_idx}_{i}")
                                stt['sec'] = s5.selectbox("Sec", strut_opts, index=0, key=f"tsec_{c_idx}_{i}")
                                
                                s_del.markdown("<br>", unsafe_allow_html=True)
                                if s_del.button("❌", key=f"del_strut_{c_idx}_{i}"):
                                    case['struts'].pop(i)
                                    st.rerun()

                            if st.button("➕ Add Strut", key=f"add_strut_{c_idx}"):
                                case['struts'].append({'tx': 0.0, 'tz': 0.0, 'bx': 0.0, 'bz': 0.0, 'sec': strut_opts[0] if strut_opts else 'Unknown'})
                                st.rerun()

                        # 💡 عرض جميع الأحمال باستخدام unique_seg_opts لضمان القراءة الدقيقة للمتكرر
                        with st.expander(f"⬇️ Loads (Auto & Manual) - Total: {len(case.get('loads', []))}", expanded=False):
                            excel_text = st.text_area("📋 Paste Directly from Excel:", placeholder="S1 \t 25.5\nS2 \t 30.0", key=f"exc_{c_idx}")
                            if st.button("⚡ Assign Loads from Text", key=f"btn_exc_{c_idx}"):
                                for line in excel_text.split('\n'):
                                    s_m = re.search(r'(S\d+)', line.strip().upper())
                                    nums = re.findall(r'-?\d+\.?\d*', line.strip().upper())
                                    if s_m and nums:
                                        vals = [v for v in nums if v != s_m.group(1)[1:]]
                                        if vals:
                                            w_v = -abs(float(vals[-1]))
                                            s_name_extracted = s_m.group(1)
                                            matching_indices = [idx_m for idx_m, name_m in enumerate(seg_names) if name_m == s_name_extracted]
                                            
                                            if matching_indices:
                                                t_mode = 'Single Segment' if len(matching_indices) == 1 else 'Multiple Segments'
                                                case.setdefault('loads', []).append({
                                                    'seg_idx': matching_indices[0], 'category': 'Live Load', 'type': 'Uniform', 
                                                    'dir': 'Global Z (Vertical)', 'target_mode': t_mode, 'target_segs_idx': matching_indices, 
                                                    'start': 0.0, 'end': case['segments'][matching_indices[0]]['L'], 'w1': w_v, 'w2': w_v, 'loc': 0.0, 'is_auto': False
                                                })
                                st.rerun()

                            cat_opts = ["Dead Load", "Live Load", "Wind Load"]
                            type_opts = ["Uniform", "Trapezoidal", "Point Load"]
                            dir_opts = ["Global X (Horizontal)", "Global Z (Vertical)", "Local Z (Perpendicular)"]
                            t_mode_opts = ["Single Segment", "Multiple Segments", "All Segments"]
                            
                            for i, ld in enumerate(case.get('loads', [])):
                                title_prefix = "🤖 Auto" if ld.get('is_auto', False) else "✍️ Manual"
                                with st.expander(f"{title_prefix} Load {i+1} ({ld.get('category')} - {ld.get('type')})", expanded=False):
                                    c_l1, c_l2, c_l3, c_l4 = st.columns([1.5, 1.5, 1.5, 0.5])
                                    
                                    idx_cat = cat_opts.index(ld.get('category', 'Dead Load')) if ld.get('category') in cat_opts else 0
                                    ld['category'] = c_l1.selectbox("Category", cat_opts, index=idx_cat, key=f"alct_{c_idx}_{i}")
                                    
                                    idx_type = type_opts.index(ld.get('type', 'Uniform')) if ld.get('type') in type_opts else 0
                                    ld['type'] = c_l2.selectbox("Type", type_opts, index=idx_type, key=f"altp_{c_idx}_{i}")
                                    
                                    c_dir = ld.get('dir', 'Global Z (Vertical)')
                                    if 'Y' in c_dir: c_dir = 'Global Z (Vertical)'
                                    idx_dir = dir_opts.index(c_dir) if c_dir in dir_opts else 1
                                    ld['dir'] = c_l3.selectbox("Direction", dir_opts, index=idx_dir, key=f"aldr_{c_idx}_{i}")
                                    
                                    c_l4.markdown("<br>", unsafe_allow_html=True)
                                    if c_l4.button("❌", key=f"adel_ld_{c_idx}_{i}"): 
                                        case['loads'].pop(i)
                                        st.rerun()
                                    
                                    idx_mode = t_mode_opts.index(ld.get('target_mode', 'Single Segment')) if ld.get('target_mode') in t_mode_opts else 0
                                    ld['target_mode'] = st.radio("Apply To:", t_mode_opts, index=idx_mode, key=f"almode_{c_idx}_{i}", horizontal=True)
                                    
                                    # 💡 الاعتماد على unique_seg_opts لتجنب خطأ التكرار 
                                    if ld['target_mode'] == "Single Segment": 
                                        default_idx = ld.get('seg_idx', 0) if ld.get('seg_idx', 0) < len(unique_seg_opts) else 0
                                        s_val = st.selectbox("Target Seg", unique_seg_opts, index=default_idx, key=f"alsg_{c_idx}_{i}")
                                        ld['seg_idx'] = int(s_val.split(' - ')[0])
                                        
                                    elif ld['target_mode'] == "Multiple Segments": 
                                        safe_multi = [unique_seg_opts[idx] for idx in ld.get('target_segs_idx', []) if idx < len(unique_seg_opts)]
                                        sel_segs = st.multiselect("Target Segs", unique_seg_opts, default=safe_multi, key=f"alsm_{c_idx}_{i}")
                                        ld['target_segs_idx'] = [int(s.split(' - ')[0]) for s in sel_segs]
                                    
                                    sc1, sc2, sc3 = st.columns(3)
                                    ld['w1'] = sc1.number_input("W1 (kN/m)", value=float(ld.get('w1', 0.0)), step=1.0, key=f"alw1_{c_idx}_{i}")
                                    ld['w2'] = sc2.number_input("W2 (kN/m)", value=float(ld.get('w2', ld['w1'])) if ld['type'] == "Trapezoidal" else float(ld['w1']), step=1.0, key=f"alw2_{c_idx}_{i}")
                                    ld['loc'] = sc3.number_input("Location (m)", value=float(ld.get('loc', 0.0)), key=f"alloc_{c_idx}_{i}") if ld['type'] == "Point Load" else 0.0

                            if st.button("➕ Add Manual Load", key=f"add_mld_{c_idx}"): 
                                case.setdefault('loads', []).append({
                                    'seg_idx': 0, 'category': 'Live Load', 'type': 'Uniform', 
                                    'dir': 'Global Z (Vertical)', 'target_mode': 'Single Segment', 
                                    'target_segs_idx': [], 'start': 0.0, 'end': case['segments'][0]['L'], 
                                    'w1': -10.0, 'w2': -10.0, 'loc': 0.0, 'is_auto': False
                                })
                                st.rerun()
                        with st.expander("🤖 Bridge Generative AI Optimizer", expanded=False):
                            ai_rxn = st.number_input("Target Max Rxn (kN)", value=54.4, step=1.0, key=f"br_{c_idx}")
                            ai_spc = st.text_input("Spacings (m)", value="2.40, 2.10, 1.80, 1.50, 1.20, 0.90, 0.60", key=f"bs_{c_idx}")
                            is_sym = st.checkbox("Symmetric", value=True, key=f"bm_{c_idx}")
                            
                            opt_mode = st.radio("Optimization Depth:", ["Quick Search", "Deep Search"], index=0, key=f"opm_{c_idx}")
                            
                            if st.button("✨ Run Optimizer", type="primary", key=f"btn_opt_{c_idx}"):
                                p_bar = st.progress(0)
                                s_txt = st.empty()
                                with st.spinner("AI is optimizing (Anti-Freeze Active)..."):
                                    succ, r_sup, r_str, msg = run_bridge_optimizer(
                                        case['segments'], case['segments'], case['sec_overrides'], 
                                        case['struts'], case['loads'], ai_rxn, ai_spc, 0.25, 
                                        is_sym, opt_mode, combo_factors, s_txt, p_bar
                                    )
                                    if r_sup:
                                        case['supports'] = r_sup
                                        case['struts'] = r_str
                                        st.session_state.bridge_cases[c_idx] = case
                                        if succ: 
                                            st.success(msg)
                                        else: 
                                            st.warning(msg)
                                        time.sleep(1.5)
                                        st.rerun()
                                    else: 
                                        st.error(msg)

                    with c_view:
                        st.markdown("<h4 style='text-align: center;'>Live Geometry & Loads</h4>", unsafe_allow_html=True)
                        
                        expanded_loads = []
                        for ld in case.get('loads', []):
                            fac = combo_factors.get(ld.get('category', 'Dead Load'), 1.0)
                            t_mode = ld.get('target_mode', 'Single Segment')
                            
                            if t_mode == 'Single Segment': 
                                t_idx_list = [ld.get('seg_idx', 0)]
                            elif t_mode == 'Multiple Segments': 
                                t_idx_list = ld.get('target_segs_idx', [])
                            else: 
                                t_idx_list = list(range(len(case['segments'])))
                                
                            for s_idx in t_idx_list:
                                if s_idx >= len(case['segments']): 
                                    continue
                                f_ld = ld.copy()
                                f_ld['seg_idx'] = s_idx
                                f_ld['w1'] *= fac
                                f_ld['w2'] *= fac
                                
                                L_seg = case['segments'][s_idx].get('L', 0.0)
                                if f_ld['type'] == 'Point Load': 
                                    loc = min(f_ld.get('loc', 0.0), L_seg)
                                    f_ld['start'] = loc
                                    f_ld['end'] = loc
                                else: 
                                    f_ld['start'] = 0.0
                                    f_ld['end'] = L_seg
                                expanded_loads.append(f_ld)
                            
                        p_nodes, p_elems, p_nloads, p_supps = build_chain_mesh(
                            case['segments'], case['sec_overrides'], expanded_loads, 
                            case['struts'], case['supports'], case.get('cut_points', [])
                        )
                        
                        st.image(get_live_preview_image(p_nodes, p_elems, p_supps, expanded_loads, case['segments']), use_container_width=True)
                        
                        if st.button(f"🚀 Run FEA Analysis for {case['title']}", type="primary", use_container_width=True, key=f"btn_run_{c_idx}"):
                            with st.spinner(f"Solving FEA Matrix..."):
                                U, R, net_load = solve_fea_engine(p_nodes, p_elems, p_nloads, p_supps)
                                
                                st.session_state[f'fea_cache_{c_idx}'] = {
                                    'nodes': p_nodes, 'elements': p_elems, 'R': R, 
                                    'supports': p_supps, 'loads': expanded_loads
                                }
                                
                                safety_data = []
                                for i_seg, sec in enumerate(case['sec_overrides']):
                                    max_m = 0.0
                                    max_v = 0.0
                                    for el in p_elems:
                                        if el.get('group') == 'segment' and el.get('seg_idx') == i_seg:
                                            max_m = max(max_m, np.max(np.abs(el.get('internal', {}).get('M', [0]))))
                                            max_v = max(max_v, np.max(np.abs(el.get('internal', {}).get('V', [0]))))
                                            
                                    seg_name_clean = case['segments'][i_seg]['name'].split('-')[0]
                                    s_status = "SAFE" if (max_m <= sec['Mall'] and max_v <= sec['Qall']) else "UNSAFE ❌"
                                    safety_data.append({
                                        "Segment": seg_name_clean, 
                                        "M_max": f"{max_m:.2f} / {sec['Mall']:.2f}", 
                                        "V_max": f"{max_v:.2f} / {sec['Qall']:.2f}", 
                                        "Status": s_status
                                    })
                                st.session_state[f'safety_df_{c_idx}'] = safety_data

                        if f'fea_cache_{c_idx}' in st.session_state:
                            st.markdown("### 🎛️ Diagram Scales")
                            c_s1, c_s2, c_s3 = st.columns(3)
                            sc_n = c_s1.slider("Axial Scale", 0.001, 0.050, 0.010, step=0.001, key=f"scn_{c_idx}")
                            sc_v = c_s2.slider("Shear Scale", 0.001, 0.050, 0.010, step=0.001, key=f"scv_{c_idx}")
                            sc_m = c_s3.slider("Moment Scale", 0.001, 0.100, 0.010, step=0.001, key=f"scm_{c_idx}")
                            
                            cd = st.session_state[f'fea_cache_{c_idx}']
                            img_bufs = plot_sap2000_diagrams(
                                cd['nodes'], cd['elements'], cd['R'], 
                                {'N': sc_n, 'V': sc_v, 'M': sc_m}, 
                                cd['supports'], cd['loads'], case['segments']
                            )
                            case['img_bufs'] = img_bufs
                            case['safety_df'] = st.session_state[f'safety_df_{c_idx}']
                            
                            if 'DL' in img_bufs:
                                st.image(img_bufs['DL'], caption="Dead Load Distribution Diagram")
                            if 'LL' in img_bufs:
                                st.image(img_bufs['LL'], caption="Live Load Distribution Diagram")
                            if 'WL' in img_bufs:
                                st.image(img_bufs['WL'], caption="Wind Load Distribution Diagram")
                                
                            c_p1, c_p2 = st.columns(2)
                            c_p1.image(img_bufs['M'], caption="Bending Moment Diagram")
                            c_p2.image(img_bufs['N'], caption="Axial Force Diagram")
                            
                            c_p3, c_p4 = st.columns(2)
                            c_p3.image(img_bufs['V'], caption="Shear Force Diagram")
                            c_p4.image(img_bufs['R'], caption="Support Reactions Diagram")
                            
                            st.table(pd.DataFrame(case['safety_df']))
                            
                    all_cases_ready.append(case)
                    
            st.markdown("---")
            if st.button("📥 Download Multi-Case Word Report", type="primary", use_container_width=True):
                doc_out = generate_multi_case_report(all_cases_ready, proj_info)
                st.download_button(
                    "💾 Save Full Report", 
                    data=doc_out.getvalue(), 
                    file_name="Acrow_Bridge_Report.docx", 
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document", 
                    use_container_width=True
                )

    else:
        # =========================================================
        # 💡 Manual Builder Mode (Updated with J1, J2, Strut Filters, and Unique IDs)
        # =========================================================
        st.info("🛠️ **Manual Builder:** Define geometry in XZ plane directly.")
        if 'man_segs' not in st.session_state: 
            st.session_state.man_segs = [{'name': 'S1', 'L': 3.0, 'type': 'Straight Line'}]
        if 'man_sups' not in st.session_state: 
            st.session_state.man_sups = [{'x': 0.0, 'z': 0.0, 'type': 'Hinged', 'angle': 0.0}, {'x': 3.0, 'z': 0.0, 'type': 'Roller', 'angle': 0.0}]
        if 'man_strs' not in st.session_state: 
            st.session_state.man_strs = []
        if 'man_lds' not in st.session_state: 
            st.session_state.man_lds = []

        c_in, c_plot = st.columns([1.2, 1.8])
        with c_in:
            st.markdown("### 1. Segments")
            for i, seg in enumerate(st.session_state.man_segs):
                c1, c2, c_del = st.columns([1, 1.5, 0.3])
                seg['name'] = c1.text_input(f"Seg {i+1} Name", value=seg.get('name', f"S{i+1}"), key=f"msn_{i}")
                seg['L'] = c2.number_input(f"Length (m)", value=float(seg.get('L', 3.0)), step=0.1, key=f"msl_{i}")
                
                c_del.markdown("<br>", unsafe_allow_html=True)
                if c_del.button("❌", key=f"mdel_seg_{i}"):
                    if len(st.session_state.man_segs) > 1:
                        st.session_state.man_segs.pop(i)
                        st.rerun()
                        
                seg.update({
                    'master_idx': i, 
                    'abs_p1': (sum(s['L'] for s in st.session_state.man_segs[:i]), 0.0), 
                    'abs_p2': (sum(s['L'] for s in st.session_state.man_segs[:i]) + seg['L'], 0.0), 
                    'Shape Type': 'Straight Line'
                })
                
            if st.button("➕ Add Segment", key="madd_seg"): 
                st.session_state.man_segs.append({'name': f"S{len(st.session_state.man_segs)+1}", 'L': 3.0, 'type': 'Straight Line'})
                st.rerun()

            st.markdown("### 2. Supports & Struts")
            st.session_state.man_sups.sort(key=lambda s: s['x'])
            for i, sup in enumerate(st.session_state.man_sups):
                st.markdown(f"**🟢 Support J{i+1}**")
                c1, c2, c3, c4, c_del = st.columns([1, 1, 1.2, 1, 0.3])
                sup['x'] = c1.number_input(f"J{i+1} X (m)", value=float(sup.get('x',0)), key=f"msx_{i}")
                sup['z'] = c2.number_input(f"J{i+1} Y (m)", value=float(sup.get('z',0)), key=f"msz_{i}")
                
                t_opts = ["Hinged", "Roller", "Fixed"]
                idx_t = t_opts.index(sup['type']) if sup['type'] in t_opts else 1
                sup['type'] = c3.selectbox(f"J{i+1} Type", t_opts, index=idx_t, key=f"mst_{i}")
                sup['angle'] = c4.number_input(f"J{i+1} Angle(°)", value=float(sup.get('angle',0.0)), key=f"msa_{i}")
                
                c_del.markdown("<br>", unsafe_allow_html=True)
                if c_del.button("❌", key=f"mdel_sup_{i}"):
                    st.session_state.man_sups.pop(i)
                    st.rerun()
                    
            if st.button("➕ Add Support", key="madd_sup"): 
                st.session_state.man_sups.append({'x':0.0, 'z':0.0, 'type':'Roller', 'angle': 0.0})
                st.rerun()

            strut_opts = get_valid_strut_names()
            for i, ds in enumerate(st.session_state.man_strs):
                c1, c2, c3, c4, s_del = st.columns([1,1,1,1,0.3])
                ds['tx'] = c1.number_input("Top X", value=float(ds.get('tx',0)), key=f"mtx_{i}")
                ds['tz'] = c2.number_input("Top Z", value=float(ds.get('tz',3)), key=f"mtz_{i}")
                ds['bx'] = c3.number_input("Bot X", value=float(ds.get('bx',1)), key=f"mbx_{i}")
                ds['bz'] = c4.number_input("Bot Z", value=float(ds.get('bz',0)), key=f"mbz_{i}")
                ds['sec'] = st.selectbox(f"Strut {i+1} Sec", strut_opts, index=0, key=f"msec_{i}")
                
                s_del.markdown("<br>", unsafe_allow_html=True)
                if s_del.button("❌", key=f"mdel_str_{i}"):
                    st.session_state.man_strs.pop(i)
                    st.rerun()
                    
            if st.button("➕ Add Strut", key="madd_str"): 
                st.session_state.man_strs.append({'tx':0.0, 'tz':3.0, 'bx':1.0, 'bz':0.0, 'sec':strut_opts[0] if strut_opts else 'Unknown'})
                st.rerun()

            st.markdown("### 3. Loads & Factors")
            c_f1, c_f2, c_f3 = st.columns(3)
            fac_d = c_f1.number_input("Dead Load Factor", value=1.00, step=0.1, key="mmf_d")
            fac_l = c_f2.number_input("Live Load Factor", value=1.00, step=0.1, key="mmf_l")
            fac_w = c_f3.number_input("Wind Load Factor", value=1.00, step=0.1, key="mmf_w")
            m_combo_factors = {'Dead Load': fac_d, 'Live Load': fac_l, 'Wind Load': fac_w}

            cat_opts = ["Dead Load", "Live Load", "Wind Load"]
            type_opts = ["Uniform", "Trapezoidal", "Point Load"]
            dir_opts = ["Global X (Horizontal)", "Global Z (Vertical)", "Local Z (Perpendicular)"]
            t_mode_opts = ["Single Segment", "Multiple Segments", "All Segments"]
            
            # 💡 توليد أسماء مُميزة برقم الإندكس للـ Manual Builder لحل مشكلة تطابق الأسماء
            seg_names_man = [s['name'] for s in st.session_state.man_segs]
            unique_seg_opts_man = [f"{idx} - {name}" for idx, name in enumerate(seg_names_man)]

            with st.expander(f"⬇️ Manual Loads ({len(st.session_state.man_lds)})", expanded=True):
                for i, ld in enumerate(st.session_state.man_lds):
                    with st.expander(f"📥 Load {i+1} ({ld.get('category', 'Dead Load')})", expanded=False):
                        c_l1, c_l2, c_l3, c_l4 = st.columns([1.5, 1.5, 1.5, 0.5])
                        
                        idx_cat = cat_opts.index(ld.get('category', 'Dead Load')) if ld.get('category') in cat_opts else 0
                        ld['category'] = c_l1.selectbox("Category", cat_opts, index=idx_cat, key=f"mmlct_{i}")
                        
                        idx_type = type_opts.index(ld.get('type', 'Uniform')) if ld.get('type') in type_opts else 0
                        ld['type'] = c_l2.selectbox("Type", type_opts, index=idx_type, key=f"mmltp_{i}")
                        
                        c_dir = ld.get('dir', 'Global Z (Vertical)')
                        if 'Y' in c_dir: c_dir = 'Global Z (Vertical)'
                        idx_dir = dir_opts.index(c_dir) if c_dir in dir_opts else 1
                        ld['dir'] = c_l3.selectbox("Direction", dir_opts, index=idx_dir, key=f"mmldr_{i}")
                        
                        c_l4.markdown("<br>", unsafe_allow_html=True)
                        if c_l4.button("❌", key=f"mmdel_ld_{i}"): 
                            st.session_state.man_lds.pop(i)
                            st.rerun()

                        idx_mode = t_mode_opts.index(ld.get('target_mode', 'Single Segment')) if ld.get('target_mode') in t_mode_opts else 0
                        ld['target_mode'] = st.radio("Apply To:", t_mode_opts, index=idx_mode, key=f"mmlmode_{i}", horizontal=True)
                        
                        # استخدام unique_seg_opts_man لتحديد القطاعات بدقة وتجنب التكرار
                        if ld['target_mode'] == "Single Segment":
                            default_idx = ld.get('seg_idx', 0) if ld.get('seg_idx', 0) < len(unique_seg_opts_man) else 0
                            s_val = st.selectbox("Target Seg", unique_seg_opts_man, index=default_idx, key=f"mmlsg_{i}")
                            ld['seg_idx'] = int(s_val.split(' - ')[0])
                            
                        elif ld['target_mode'] == "Multiple Segments":
                            safe_multi = [unique_seg_opts_man[idx] for idx in ld.get('target_segs_idx', []) if idx < len(unique_seg_opts_man)]
                            sel_segs = st.multiselect("Target Segs", unique_seg_opts_man, default=safe_multi, key=f"mmlsm_{i}")
                            ld['target_segs_idx'] = [int(s.split(' - ')[0]) for s in sel_segs]

                        sc1, sc2, sc3 = st.columns(3)
                        ld['w1'] = sc1.number_input("W1 (kN/m)", value=float(ld.get('w1', -10.0)), step=1.0, key=f"mmlw1_{i}")
                        if ld['type'] == "Trapezoidal":
                            ld['w2'] = sc2.number_input("W2 (kN/m)", value=float(ld.get('w2', ld['w1'])), step=1.0, key=f"mmlw2_{i}")
                        else:
                            ld['w2'] = ld['w1']
                        
                        if ld['type'] == "Point Load":
                            ld['loc'] = sc3.number_input("Location (m)", value=float(ld.get('loc', 0.0)), key=f"mmlloc_{i}")
                        else:
                            ld['loc'] = 0.0

                if st.button("➕ Add Manual Load", key="mmadd_ld"): 
                    st.session_state.man_lds.append({
                        'seg_idx': 0, 'w1': -10.0, 'w2': -10.0, 'dir': 'Global Z (Vertical)', 
                        'category': 'Live Load', 'type': 'Uniform', 'target_mode': 'Single Segment', 
                        'target_segs_idx': [], 'loc': 0.0
                    })
                    st.rerun()

        with c_plot:
            st.markdown("<h4 style='text-align: center;'>Live Geometry & Loads</h4>", unsafe_allow_html=True)
            active_sections = [{'name': "Soldier U100", 'E': 2100.0, 'A': 34.3/10000.0, 'I': 412.0/100000000.0, 'Mall': 13.1, 'Qall': 100.8}] * len(st.session_state.man_segs)
            
            m_expanded_loads = []
            for ld in st.session_state.man_lds:
                fac = m_combo_factors.get(ld.get('category', 'Dead Load'), 1.0)
                t_mode = ld.get('target_mode', 'Single Segment')
                
                if t_mode == 'Single Segment': 
                    t_idx_list = [ld.get('seg_idx', 0)]
                elif t_mode == 'Multiple Segments': 
                    t_idx_list = ld.get('target_segs_idx', [])
                else: 
                    t_idx_list = list(range(len(st.session_state.man_segs)))
                
                for s_idx in t_idx_list:
                    if s_idx >= len(st.session_state.man_segs): 
                        continue
                    f_ld = ld.copy()
                    f_ld['seg_idx'] = s_idx
                    f_ld['w1'] *= fac
                    f_ld['w2'] *= fac
                    
                    L_seg = st.session_state.man_segs[s_idx].get('L', 0.0)
                    if f_ld['type'] == 'Point Load': 
                        loc = min(f_ld.get('loc', 0.0), L_seg)
                        f_ld['start'] = loc
                        f_ld['end'] = loc
                    else: 
                        f_ld['start'] = 0.0
                        f_ld['end'] = L_seg
                        
                    m_expanded_loads.append(f_ld)
            
            p_nodes, p_elems, p_nloads, p_supps = build_chain_mesh(
                st.session_state.man_segs, active_sections, m_expanded_loads, 
                st.session_state.man_strs, st.session_state.man_sups, []
            )
            
            st.image(get_live_preview_image(p_nodes, p_elems, p_supps, m_expanded_loads, st.session_state.man_segs), use_container_width=True)
            
            if st.button("🚀 Run FEA", type="primary", use_container_width=True, key="mrun_btn"):
                with st.spinner("Solving FEA..."):
                    U, R, net_load = solve_fea_engine(p_nodes, p_elems, p_nloads, p_supps)
                    st.session_state['man_fea_cache'] = {'nodes': p_nodes, 'elements': p_elems, 'R': R, 'supports': p_supps, 'loads': m_expanded_loads}
                    
                    safety_data = []
                    for i_seg, sec in enumerate(active_sections):
                        max_m, max_v = 0.0, 0.0
                        for el in p_elems:
                            if el.get('group') == 'segment' and el.get('seg_idx') == i_seg:
                                max_m = max(max_m, np.max(np.abs(el.get('internal', {}).get('M', [0]))))
                                max_v = max(max_v, np.max(np.abs(el.get('internal', {}).get('V', [0]))))
                        
                        seg_name_clean = st.session_state.man_segs[i_seg]['name'].split('-')[0]
                        s_status = "SAFE" if (max_m <= sec['Mall'] and max_v <= sec['Qall']) else "UNSAFE ❌"
                        safety_data.append({
                            "Segment": seg_name_clean, 
                            "M_max": f"{max_m:.2f} / {sec['Mall']:.2f}", 
                            "V_max": f"{max_v:.2f} / {sec['Qall']:.2f}", 
                            "Status": s_status
                        })
                    
                    st.session_state.man_safety_df = safety_data

            if 'man_fea_cache' in st.session_state:
                st.markdown("### 🎛️ Diagram Scales")
                c_s1, c_s2, c_s3 = st.columns(3)
                sc_n = c_s1.slider("Axial Scale", 0.001, 0.050, 0.010, step=0.001, key="mscn")
                sc_v = c_s2.slider("Shear Scale", 0.001, 0.050, 0.010, step=0.001, key="mscv")
                sc_m = c_s3.slider("Moment Scale", 0.001, 0.100, 0.010, step=0.001, key="mscm")
                
                cd = st.session_state['man_fea_cache']
                img_bufs = plot_sap2000_diagrams(cd['nodes'], cd['elements'], cd['R'], {'N': sc_n, 'V': sc_v, 'M': sc_m}, cd['supports'], cd['loads'], st.session_state.man_segs)
                st.session_state.man_case_data = [{'title': 'Manual Case', 'img_bufs': img_bufs, 'safety_df': st.session_state.man_safety_df}]
                
                if 'DL' in img_bufs: st.image(img_bufs['DL'], caption="Dead Load Distribution Diagram")
                if 'LL' in img_bufs: st.image(img_bufs['LL'], caption="Live Load Distribution Diagram")
                if 'WL' in img_bufs: st.image(img_bufs['WL'], caption="Wind Load Distribution Diagram")
                
                c_p1, c_p2 = st.columns(2)
                c_p1.image(img_bufs['M'], caption="Moment")
                c_p2.image(img_bufs['N'], caption="Axial")
                
                c_p3, c_p4 = st.columns(2)
                c_p3.image(img_bufs['V'], caption="Shear")
                c_p4.image(img_bufs['R'], caption="Reactions")
                
                st.table(pd.DataFrame(st.session_state.man_safety_df))
                
                st.markdown("---")
                if st.button("📥 Download Word Report", type="primary", use_container_width=True, key="mdown_btn"):
                    doc_out = generate_multi_case_report(st.session_state.man_case_data, proj_info)
                    st.download_button(
                        "💾 Save Manual Case Report", 
                        data=doc_out.getvalue(), 
                        file_name="Acrow_Bridge_Manual_Report.docx", 
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document", 
                        use_container_width=True,
                        key="mdown_btn_final"
                    )

# =========================================================
# Execution Entry Point
# =========================================================
if __name__ == "__main__":
    render_bridge_module({"proj_name": "Acrow Bridges", "date_val": "2026"})