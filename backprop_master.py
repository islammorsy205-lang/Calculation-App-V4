# backprop_master.py

import streamlit as st
import numpy as np
import io
import os
import re
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Cm, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_COLOR_INDEX
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

try:
    from math_solver import get_prop_allowable, get_scaffold_allowable
except ImportError:
    st.error("⚠️ لم يتم العثور على math_solver.py. برجاء التأكد من مسار الملفات.")
    def get_prop_allowable(*args): return 20.0
    def get_scaffold_allowable(*args): return 30.0

try:
    from report_builder import insert_blue_banner, add_eq, append_pdf_stream_to_word
except ImportError:
    pass

def get_shoring_capacity(t_nm, subtype, unb, req_ext):
    t_al = 20.0
    if t_nm == "Shorebrace Frame":
        t_al = 54.00
    elif t_nm == "Cup-lock":
        t_al = get_scaffold_allowable("Cup-lock", subtype, unb)
    elif t_nm == "Ring-lock":
        t_al = get_scaffold_allowable("Ring-lock", subtype, unb)
    elif t_nm == "Acrow Prop":
        t_al = get_prop_allowable(subtype, req_ext, True)
    return t_al

def plot_zone_system(conf):
    results = []
    W_attacking = conf['W_fresh']
    results.append({'level': 'Fresh Slab', 'attacking': W_attacking, 'capacity': 0, 'transferred': W_attacking})
    
    current_P = W_attacking
    for i, slab in enumerate(conf['existing_slabs']):
        if current_P <= 0: break
        avail_cap = (slab['sidl'] + slab['ll']) * (slab['strength'] / 100.0)
        absorbed = min(current_P, avail_cap)
        current_P = max(0, current_P - avail_cap)
        
        results.append({
            'level': f'Existing Slab {i+1}', 
            'attacking': results[-1]['transferred'], 
            'capacity': avail_cap, 
            'transferred': current_P,
            'sidl': slab['sidl'],
            'll': slab['ll'],
            'strength': slab['strength']
        })
        
    num_levels = len(results)
    fig, ax = plt.subplots(figsize=(8, num_levels * 1.5))
    
    y_pos = np.arange(num_levels, 0, -1) * 2
    
    for i, res in enumerate(results):
        y = y_pos[i]
        color = 'gray' if 'Existing' in res['level'] else 'blue'
        rect = plt.Rectangle((1, y-0.2), 6, 0.4, facecolor=color, alpha=0.5, edgecolor='black', linewidth=2)
        ax.add_patch(rect)
        ax.text(4, y, res['level'], ha='center', va='center', fontsize=12, fontweight='bold', color='white')
        
        if 'Existing' in res['level']:
            combined_text = f"SIDL:{res['sidl']:.2f} | L.L:{res['ll']:.2f} (kN/m²)\nStrength achieved: {res['strength']:.0f}%"
            ax.text(7.4, y + 0.25, combined_text, ha='right', va='bottom', fontsize=6, fontweight='normal', color='dimgray')
        
        if i < num_levels - 1 and res['transferred'] > 0:
            next_y = y_pos[i+1]
            ax.annotate('', xy=(4, next_y+0.2), xytext=(4, y-0.2),
                        arrowprops=dict(facecolor='red', shrink=0.05, width=4, headwidth=10))
            ax.text(4.2, (y + next_y)/2, f"{res['transferred']:.2f} kN/m²", color='red', fontsize=11, fontweight='bold')
            
            ax.plot([2.5, 2.5], [y-0.2, next_y+0.2], color='black', linewidth=3)
            ax.plot([5.5, 5.5], [y-0.2, next_y+0.2], color='black', linewidth=3)
            
    ax.set_xlim(0, 8.5) 
    ax.set_ylim(0, max(y_pos) + 1)
    ax.axis('off')
    plt.title("Load Transfer Diagram", fontsize=12, fontweight='bold')
    
    plt.tight_layout()
    img_buf = io.BytesIO()
    plt.savefig(img_buf, format='png', dpi=150, bbox_inches='tight')
    plt.close(fig)
    img_buf.seek(0)
    return img_buf

def generate_backprop_report(configs, proj_info):
    ref_code = proj_info.get('ref_code', 'BS')
    proj_name = proj_info.get('proj_name', '')
    contractor = proj_info.get('contractor', '')
    calc_sub = proj_info.get('calc_sub', '')
    sys_name = proj_info.get('sys_name', '')
    proj_no = proj_info.get('proj_no', '')
    calc_by = proj_info.get('calc_by', '')
    date_val = proj_info.get('date_val', '')
    chk_by = proj_info.get('chk_by', '')
    cover_img = proj_info.get('cover_img', None)
    data_sheets = proj_info.get('data_sheets', [])
    
    if os.path.exists("Acrow_Template.docx"):
        doc = Document("Acrow_Template.docx")
    else:
        doc = Document()
        
    def remove_hardcoded_prefix(p):
        if p.text and "CALCULATION SHEET FOR" in p.text.upper():
            for r in p.runs:
                if "CALCULATION SHEET FOR" in r.text.upper():
                    r.text = re.sub(r'(?i)CALCULATION SHEET FOR\s*', '', r.text)
            if "CALCULATION SHEET FOR" in p.text.upper():
                clean_text = re.sub(r'(?i)CALCULATION SHEET FOR\s*', '', p.text)
                if p.runs:
                    font_name = p.runs[0].font.name
                    font_size = p.runs[0].font.size
                    font_bold = p.runs[0].font.bold
                    font_color = p.runs[0].font.color.rgb if p.runs[0].font.color else None
                    for r in p.runs: r.text = ""
                    p.runs[0].text = clean_text
                    p.runs[0].font.name = font_name
                    p.runs[0].font.size = font_size
                    p.runs[0].font.bold = font_bold
                    if font_color: p.runs[0].font.color.rgb = font_color
                else:
                    p.text = clean_text

    for p in doc.paragraphs: remove_hardcoded_prefix(p)
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                for p in cell.paragraphs: remove_hardcoded_prefix(p)

    replacements = {
        "[PROJECT_NAME]": proj_name,
        "[CONTRACTOR]": contractor,
        "[CALC_SUBJECT]": calc_sub,
        "[SYSTEM_NAME]": sys_name
    }
    
    for p in doc.paragraphs:
        if "[COVER_IMAGE]" in p.text:
            for r in p.runs: r.text = r.text.replace("[COVER_IMAGE]", "")
            if cover_img and cover_img != "No images found." and os.path.exists(cover_img):
                p.add_run().add_picture(cover_img, width=Cm(15.0))
        for k, v in replacements.items():
            if k in p.text:
                for r in p.runs:
                    if k in r.text: r.text = r.text.replace(k, str(v))
                if k in p.text: p.text = p.text.replace(k, str(v))

    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for k, v in replacements.items():
                        if k in p.text:
                            for r in p.runs:
                                if k in r.text: r.text = r.text.replace(k, str(v))
                            if k in p.text: p.text = p.text.replace(k, str(v))
                            
    doc.add_page_break()
    
    # --- Headers and Footers ---
    for sec in doc.sections:
        for hf in [sec.header, sec.first_page_header, sec.footer, sec.first_page_footer]:
            if hf:
                for tbl in hf.tables:
                    for row in tbl.rows:
                        for cell in row.cells:
                            for p in cell.paragraphs:
                                for k, v in {"[PROJECT_NAME]": proj_name, "[CONTRACTOR]": contractor, "[PROJ_NO]": proj_no, "[DATE]": date_val, "[CALC_BY]": calc_by, "[CHK_BY]": chk_by, "[REV]": "00"}.items():
                                    if k in p.text:
                                        for r in p.runs:
                                            if k in r.text: r.text = r.text.replace(k, str(v))
                                        if k in p.text: p.text = p.text.replace(k, str(v))
                                        for r in p.runs: r.font._element.set(qn('w:ascii'), 'Arial')
                                        
    # --- Regulations ---
    insert_blue_banner(doc, "REGULATIONS AND STANDARDS", font_size=16)
    doc.add_paragraph()
    if "BS" in ref_code: 
        add_eq(doc, "1- BS 5975-1996: FORMWORK FOR CONCRETE")
        add_eq(doc, "2- BS 5975-2008: FORMWORK FOR CONCRETE")
        add_eq(doc, "3- FORMWORK A GUIDE TO A GOOD PRACTICE")
        add_eq(doc, "4- WISA®-FORM PLYWOOD.")
        add_eq(doc, "5- THE SAUDI BUILDING CODE (SBC) 2018")
    else: 
        add_eq(doc, "1- ACI 347R-14 ....... GUIDE TO FORMWORK FOR CONCRETE.")
        add_eq(doc, "2- ACI SP-4 ......... FORMWORK FOR CONCRETE.")
        add_eq(doc, "3- WISA®-FORM PLYWOOD.")
        add_eq(doc, "4- THE SAUDI BUILDING CODE (SBC) 2018")
    
    # --- Data Sheets ---
    if data_sheets:
        doc.add_page_break()
        insert_blue_banner(doc, "FORMWORK MATERIALS TECHNICAL DATA", font_size=14)
        for f in data_sheets:
            if os.path.exists(f): 
                append_pdf_stream_to_word(f, doc, is_path=True, max_width_cm=17.5, max_height_cm=24.0, add_border=True, reduce_first_page=True)
    
    # --- Design Loads ---
    design_pdf = "Design_Loads_BS.pdf" if "BS" in ref_code and os.path.exists("Design_Loads_BS.pdf") else ("Design_Loads_ACI.pdf" if "ACI" in ref_code and os.path.exists("Design_Loads_ACI.pdf") else None)
    if design_pdf: 
        doc.add_page_break()
        insert_blue_banner(doc, "DESIGN LOADS FOR SLAB", font_size=14)
        append_pdf_stream_to_word(design_pdf, doc, is_path=True, max_width_cm=17.5, max_height_cm=24.0, add_border=True, reduce_first_page=True)

    doc.add_page_break()
    insert_blue_banner(doc, "CALCULATION SHEET FOR RE-PROPPING (BACK-PROPPING)", font_size=16)
    
    def force_ltr_left(p):
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        pPr = p._element.get_or_add_pPr()
        bidi = OxmlElement('w:bidi')
        bidi.set(qn('w:val'), '0')
        pPr.append(bidi)
        
    def add_p(text, bold=False, underline=False, color=None, size=12, indent=0):
        p = doc.add_paragraph()
        force_ltr_left(p)  
        p.paragraph_format.line_spacing = 1.5
        if indent > 0:
            p.paragraph_format.left_indent = Cm(indent)
        r = p.add_run(text)
        r.font.name = 'Arial'
        r.font.size = Pt(size)
        r.font.bold = bold
        r.font.underline = underline
        r.font.rtl = False  
        if color:
            r.font.color.rgb = color
        return p

    # --- Actual Zone Calculations ---
    for z_idx, conf in enumerate(configs):
        if z_idx > 0: doc.add_page_break()
        
        add_p(f"Zone {z_idx+1} Calculation:", bold=True, size=14, color=RGBColor(0, 0, 128))
        add_p("Design Loads for Slabs Back-Propping", bold=True, underline=True, color=RGBColor(192, 0, 0), size=14)
        
        add_p("A. Dead load:", indent=1)
        add_p(f"-  O.W of Concrete (Concrete density = {conf['gamma_c']:.1f} KN/m³)", indent=2)
        add_p(f"-  O.W of Formwork = {conf['FW']:.2f} KN/m²", indent=2)
        
        add_p("B. Live load:", indent=1)
        add_p(f"-  Live load = {conf['LL']:.2f} KN/m²", indent=2)
        
        add_p("\nLoad Acting on Existing Lower Slabs while Casting of Fresh Concrete Slabs.", bold=True, underline=True)
        add_p(f"-  Slabs: {conf['ts_fresh'] * 1000:.0f} mm.", indent=1)
        
        W_slab_str = f"W Slab (KN/m²) = O.W of Slab + Live Load + O.W of Formwork"
        add_p(W_slab_str, bold=True, indent=1)
        calc_str = f"               = {conf['gamma_c']:.1f}X{conf['ts_fresh']:.2f} + {conf['LL']:.2f} + {conf['FW']:.2f} = {conf['W_fresh']:.2f} KN/m²"
        add_p(calc_str, bold=True, indent=1)

        current_transferred = conf['W_fresh']
        
        for i, slab in enumerate(conf['existing_slabs']):
            if current_transferred <= 0: break
            
            add_p(f"\n❖ Characteristic Surface Loads for Critical Zone (Existing Slab {i+1}):", bold=True, underline=True)
            add_p(f"-  Super-imposed Dead Load (SDL) = {slab['sidl']:.2f} KN/m²", indent=1)
            add_p(f"-  Live Load (L.L) = {slab['ll']:.2f} KN/m²", indent=1)
            
            unfactored = slab['sidl'] + slab['ll']
            add_p(f"\nTotal un-Factored Load (W) = {slab['sidl']:.2f} + {slab['ll']:.2f} = {unfactored:.2f} KN/m²")
            add_p(f"Assume the concrete reach {slab['strength']:.0f}% from its strength.")
            capacity = unfactored * (slab['strength'] / 100.0)
            add_p(f"Therefore: - Total resisting load = {unfactored:.2f} x {slab['strength']/100:.2f} = {capacity:.2f} KN/m²")
            
            add_p(f"\nRe-Shoring Check for the System loaded on Existing Slab {i+1}", bold=True)
            add_p(f"➢ Total Re-Shoring Loads from upper level = {current_transferred:.2f} KN/m²", indent=1)
            
            next_transferred = max(0, current_transferred - capacity)
            add_p("Therefore; -", bold=True)
            add_p(f"The Transferred Loads to the Lower Level = {current_transferred:.2f} - {capacity:.2f} = {next_transferred:.2f} KN/m²")
            
            if next_transferred > 0:
                grid_i = slab['shore']
                add_p(f"Max. Loaded Area \"Back Propped area at Level {i+1}\" = {grid_i['gx']:.2f}x{grid_i['gy']:.2f} = {grid_i['area']:.2f} m²", indent=1)
                load_leg_i = grid_i['area'] * next_transferred
                
                check_txt_i = f"Area Load on one leg of {grid_i['sys']} = {grid_i['area']:.2f} x {next_transferred:.2f} = {load_leg_i:.2f} KN < {grid_i['cap']:.2f} KN"
                
                p_check_i = doc.add_paragraph()
                force_ltr_left(p_check_i) 
                p_check_i.paragraph_format.line_spacing = 1.5
                p_check_i.paragraph_format.left_indent = Cm(1)
                
                r_ci = p_check_i.add_run(check_txt_i)
                r_ci.font.name = 'Arial'
                r_ci.font.size = Pt(12)
                r_ci.font.rtl = False
                
                r_resi = p_check_i.add_run("   SAFE" if load_leg_i <= grid_i['cap'] else "   UNSAFE")
                r_resi.font.name = 'Arial'
                r_resi.font.size = Pt(12)
                r_resi.font.bold = True
                r_resi.font.rtl = False
                r_resi.font.color.rgb = RGBColor(0, 128, 0) if load_leg_i <= grid_i['cap'] else RGBColor(255, 0, 0)
                
            current_transferred = next_transferred

        levels_propped = conf['levels_propped']
        add_p("\n=======================================================", bold=True)
        add_p(f"FINAL CONCLUSION FOR ZONE {z_idx+1}:", bold=True, color=RGBColor(0, 128, 0))
        add_p(f"Total Number of Floors Required to be Propped = {levels_propped} Floors", bold=True)
        
        if levels_propped == 1:
            add_p("(This means ONLY the main shoring under the fresh slab is needed. The first existing slab can safely carry the transferred load, and NO back-propping is required underneath it.)", color=RGBColor(100,100,100))
        else:
            add_p(f"(This includes 1 floor of main shoring directly under the fresh slab, plus {levels_propped - 1} floor(s) of back-propping underneath the existing slabs.)", color=RGBColor(100,100,100))
        add_p("=======================================================\n", bold=True)

        add_p("Load Path Diagram:", bold=True)
        p_img = doc.add_paragraph()
        p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pPr = p_img._element.get_or_add_pPr()
        bidi = OxmlElement('w:bidi')
        bidi.set(qn('w:val'), '0')
        pPr.append(bidi)
        
        p_img.add_run().add_picture(io.BytesIO(conf['img_buf'].read()), width=Cm(16.0))
            
    out = io.BytesIO()
    doc.save(out)
    return out

def render_backprop_module(proj_info):
    ref_code = proj_info.get('ref_code', 'BS')
    st.markdown("## 🏗️ Multi-Zone Slab Re-propping (Back-propping)")
    st.info("💡 **Independent Module:** Evaluate multiple independent zones of fresh slabs. Each floor can have a distinct shoring grid and system.")
    
    LL_const = 1.50 if "BS" in ref_code else 2.40
    FW_load = 0.50
    st.success(f"**Code Detected:** {ref_code}  →  Construction Live Load = {LL_const} kN/m², Formwork = {FW_load} kN/m²")
    
    val_num_zones = int(st.session_state.get("bp_num_zones", 1))
    num_zones = st.number_input("Number of Fresh Slab Zones to Check", min_value=1, max_value=5, value=val_num_zones, key="bp_num_zones")
    tabs = st.tabs([f"Zone {i+1}" for i in range(int(num_zones))])
    
    configs = []
    sys_opts = ["Acrow Prop", "Cup-lock", "Ring-lock", "Shorebrace Frame"]
    
    for idx, tab in enumerate(tabs):
        with tab:
            st.markdown(f"### Zone {idx+1} Properties")
            c1, c2 = st.columns(2)
            
            val_g = float(st.session_state.get(f'g_{idx}', 25.0))
            gamma_c = c1.number_input("Concrete Density (kN/m³)", value=val_g, step=0.5, key=f'g_{idx}')
            
            val_ts = float(st.session_state.get(f'ts_{idx}', 0.28))
            ts_fresh = c2.number_input("Fresh Slab Thickness (m)", value=val_ts, step=0.01, key=f'ts_{idx}')
            
            W_fresh = (gamma_c * ts_fresh) + LL_const + FW_load
            st.info(f"**Total Fresh Slab Load = {W_fresh:.2f} kN/m²**")
            
            st.markdown("---")
            
            val_nx = int(st.session_state.get(f'nx_{idx}', 2))
            num_exist = st.number_input("Number of Existing Slabs Below", min_value=1, value=val_nx, step=1, key=f'nx_{idx}')
            existing_slabs = []
            
            current_ui_transferred = W_fresh
            
            for j in range(int(num_exist)):
                st.markdown(f"#### Existing Slab {j+1}")
                ec1, ec2, ec3 = st.columns(3)
                
                val_ll = float(st.session_state.get(f'll_{idx}_{j}', 2.50))
                ll_des = ec1.number_input("Design L.L (kN/m²)", value=val_ll, step=0.5, key=f'll_{idx}_{j}')
                
                val_sidl = float(st.session_state.get(f'sidl_{idx}_{j}', 0.50))
                sidl_des = ec2.number_input("Design SIDL (kN/m²)", value=val_sidl, step=0.5, key=f'sidl_{idx}_{j}')
                
                val_str = float(st.session_state.get(f'str_{idx}_{j}', 80.0))
                strength = ec3.number_input("Strength Achieved (%)", value=val_str, step=5.0, key=f'str_{idx}_{j}')
                
                slab_capacity = (sidl_des + ll_des) * (strength / 100.0)
                next_ui_transferred = max(0, current_ui_transferred - slab_capacity)
                
                st.markdown(f"**Level {j+1} Back-propping Shoring (Props Under Existing Slab {j+1})**")
                ssc1, ssc2, ssc3 = st.columns(3)
                
                idx_sys = sys_opts.index(st.session_state.get(f'sj_{idx}_{j}', sys_opts[0])) if st.session_state.get(f'sj_{idx}_{j}') in sys_opts else 0
                sys_j = ssc1.selectbox("Shoring Type", sys_opts, index=idx_sys, key=f'sj_{idx}_{j}')
                
                val_gx = float(st.session_state.get(f'gxj_{idx}_{j}', 1.2))
                gx_j = ssc2.number_input("Grid X (m)", value=val_gx, step=0.1, key=f'gxj_{idx}_{j}')
                
                val_gy = float(st.session_state.get(f'gyj_{idx}_{j}', 1.2))
                gy_j = ssc3.number_input("Grid Y (m)", value=val_gy, step=0.1, key=f'gyj_{idx}_{j}')
                
                subtype_j, unb_j, ext_j = "", 1.5, 3.0
                if sys_j == "Cup-lock":
                    cup_opts = ["S355 (st.52)", "S235"]
                    idx_cj = cup_opts.index(st.session_state.get(f'cj_{idx}_{j}', cup_opts[0])) if st.session_state.get(f'cj_{idx}_{j}') in cup_opts else 0
                    subtype_j = ssc1.selectbox("Grade", cup_opts, index=idx_cj, key=f'cj_{idx}_{j}')
                    
                    val_cuj = float(st.session_state.get(f'cuj_{idx}_{j}', 1.5))
                    unb_j = ssc2.number_input("Unbraced (m)", value=val_cuj, key=f'cuj_{idx}_{j}')
                elif sys_j == "Ring-lock":
                    ring_opts = ["Ringlock 1.5\"", "Ringlock 2.0\""]
                    idx_rj = ring_opts.index(st.session_state.get(f'rj_{idx}_{j}', ring_opts[0])) if st.session_state.get(f'rj_{idx}_{j}') in ring_opts else 0
                    subtype_j = ssc1.selectbox("Size", ring_opts, index=idx_rj, key=f'rj_{idx}_{j}')
                    
                    val_ruj = float(st.session_state.get(f'ruj_{idx}_{j}', 1.5))
                    unb_j = ssc2.number_input("Unbraced (m)", value=val_ruj, key=f'ruj_{idx}_{j}')
                elif sys_j == "Acrow Prop":
                    try: 
                        from config import PROP_DB
                        prop_opts = list(PROP_DB.keys())
                        idx_pj = prop_opts.index(st.session_state.get(f'pj_{idx}_{j}', prop_opts[0])) if st.session_state.get(f'pj_{idx}_{j}') in prop_opts else 0
                        subtype_j = ssc1.selectbox("Prop Type", prop_opts, index=idx_pj, key=f'pj_{idx}_{j}')
                    except: 
                        subtype_j = "Prop No.2"
                        
                    val_pej = float(st.session_state.get(f'pej_{idx}_{j}', 3.0))
                    ext_j = ssc2.number_input("Extension (m)", value=val_pej, key=f'pej_{idx}_{j}')
                
                cap_j = get_shoring_capacity(sys_j, subtype_j, unb_j, ext_j)
                level_j_shore = {'sys': sys_j, 'gx': gx_j, 'gy': gy_j, 'area': gx_j*gy_j, 'cap': cap_j}
                
                if next_ui_transferred > 0:
                    actual_leg_load = (gx_j * gy_j) * next_ui_transferred
                    if actual_leg_load <= cap_j:
                        st.success(f"✅ **SAFE** | Load Transferred: **{next_ui_transferred:.2f} kN/m²** | Actual Leg Load: **{actual_leg_load:.2f} kN** < Leg Capacity: **{cap_j:.2f} kN**")
                    else:
                        st.error(f"❌ **UNSAFE** | Load Transferred: **{next_ui_transferred:.2f} kN/m²** | Actual Leg Load: **{actual_leg_load:.2f} kN** > Leg Capacity: **{cap_j:.2f} kN**")
                else:
                    st.success(f"✅ **NO SHORING REQUIRED** | Slab fully absorbed the load. (Transferred: 0.00 kN/m²)")
                
                current_ui_transferred = next_ui_transferred
                
                existing_slabs.append({
                    'll': ll_des, 'sidl': sidl_des, 'strength': strength, 'shore': level_j_shore
                })
                
            configs.append({
                'gamma_c': gamma_c, 'ts_fresh': ts_fresh, 'LL': LL_const, 'FW': FW_load, 'W_fresh': W_fresh,
                'existing_slabs': existing_slabs
            })
            
    st.markdown("---")
    if st.button("🚀 Calculate & Generate Detailed Left-Aligned Report", type="primary", use_container_width=True):
        
        for idx, conf in enumerate(configs):
            current_transferred = conf['W_fresh']
            levels_propped = 1 
            
            for slab in conf['existing_slabs']:
                if current_transferred <= 0: break
                capacity = (slab['sidl'] + slab['ll']) * (slab['strength'] / 100.0)
                next_transferred = max(0, current_transferred - capacity)
                if next_transferred > 0:
                    levels_propped += 1
                current_transferred = next_transferred
                
            conf['levels_propped'] = levels_propped
            conf['img_buf'] = plot_zone_system(conf)
            
            st.success(f"✅ **Zone {idx+1} Summary:** You need to prop **{levels_propped} Floors** in total. (1 Floor Main Shoring + {levels_propped-1} Floors Back-propping).")
            st.image(conf['img_buf'], use_container_width=False)
            
        with st.spinner("🔄 Building Complete Calculation Document..."):
            docx_out = generate_backprop_report(configs, proj_info)
            st.download_button("⬇️ Download Back-propping Calculation Sheet", 
                               data=docx_out.getvalue(), 
                               file_name="Back_Propping_Report.docx", 
                               mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")