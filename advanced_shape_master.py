# advanced_shape_master.py

import streamlit as st
import numpy as np
import pandas as pd
import io
import os
import re
import math
import tempfile
import time
from collections import deque
import matplotlib as mpl
import matplotlib.pyplot as plt
from matplotlib.patches import Polygon
from PIL import Image, ImageChops
from docx import Document
from docx.shared import Cm, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import streamlit.components.v1 as components

try:
    import ezdxf
except ImportError:
    st.error("⚠️ مكتبة 'ezdxf' غير موجودة! برجاء كتابة الأمر 'pip install ezdxf' في التيرمينال لتفعيل ميزة قراءة الكاد.")
    ezdxf = None

try:
    from config import SECTIONS_DB, STRUTS_DB
    from report_builder import insert_blue_banner, add_eq, append_pdf_stream_to_word
except ImportError:
    st.error("⚠️ برجاء التأكد من وجود ملفات config.py و report_builder.py")

# =========================================================
# 0. Helper Functions & Styles
# =========================================================
def apply_plot_styles():
    mpl.rcParams['font.family'] = 'sans-serif'
    mpl.rcParams['font.sans-serif'] = ['Arial', 'Helvetica', 'DejaVu Sans']
    mpl.rcParams['axes.linewidth'] = 0.3
    mpl.rcParams['font.size'] = 7
    mpl.rcParams['font.weight'] = 'normal'

def get_short_name(sec_name):
    return re.sub(r'\s*\(.*?\)', '', sec_name).strip()

def crop_image_bbox(img_bytes):
    img = Image.open(io.BytesIO(img_bytes)).convert("RGBA")
    bg = Image.new("RGBA", img.size, (255, 255, 255, 0))
    diff = ImageChops.difference(img, bg)
    bbox = diff.getbbox()
    if bbox:
        img = img.crop(bbox)
    out = io.BytesIO()
    img.save(out, format='PNG')
    return out.getvalue()

def safe_render_fig(fig):
    try:
        plt.subplots_adjust(left=0.01, right=0.99, top=0.99, bottom=0.01)
        buf = io.BytesIO()
        fig.savefig(buf, format='png', dpi=400, bbox_inches='tight', pad_inches=0.0, transparent=True)
        return crop_image_bbox(buf.getvalue())
    finally:
        plt.close(fig)

def draw_reaction_arrow(ax, node_x, node_y, force_mag, axis_nx, axis_ny):
    if abs(force_mag) < 0.001:
        return
    arr_L = 0.5  
    sgn = np.sign(force_mag)
    dx = sgn * axis_nx
    dy = sgn * axis_ny
    start_x = node_x - arr_L * dx
    start_y = node_y - arr_L * dy
    arr_c = 'blue' if force_mag >= 0 else 'red'
    
    ax.arrow(
        start_x, start_y, arr_L*dx, arr_L*dy, 
        length_includes_head=True, 
        head_width=0.08, head_length=0.12, 
        fc=arr_c, ec=arr_c, lw=0.8, zorder=5
    )
    ax.text(
        start_x - 0.15*dx, start_y - 0.15*dy, 
        f"{force_mag:+.2f}", 
        color=arr_c, fontsize=7, fontname='Arial', 
        ha='center', va='center'
    )

def eval_seg_point(seg, s_val, start_data=None):
    if seg.get('is_divided'):
        actual_s = s_val + seg.get('parent_offset', 0.0)
        return eval_seg_point(seg['parent_seg'], actual_s, start_data)

    L = seg.get('L', 0.0)
    s_val = min(max(s_val, 0.0), L)
    ratio = s_val / L if L > 1e-6 else 0.0
    is_dxf = seg.get('is_dxf', False)
    shape_type = seg.get('Shape Type', 'Straight Line')
    
    if is_dxf:
        if shape_type == 'Straight Line' and 'abs_p1' in seg:
            p1 = seg['abs_p1']
            p2 = seg['abs_p2']
            px = p1[0] + ratio * (p2[0] - p1[0])
            py = p1[1] + ratio * (p2[1] - p1[1])
            dx = p2[0] - p1[0]
            dy = p2[1] - p1[1]
            th = math.atan2(dy, dx)
            return px, py, th
            
        elif shape_type == 'Curve (Arc & Radius)' and 'abs_c' in seg:
            c = seg['abs_c']
            r = seg['abs_r']
            current_ang = seg['abs_sa'] + ratio * seg.get('sweep', 0)
            px = c[0] + r * math.cos(current_ang)
            py = c[1] + r * math.sin(current_ang)
            th = current_ang + math.pi/2
            return px, py, th
            
    if start_data:
        x0 = start_data.get('x0', 0)
        y0 = start_data.get('y0', 0)
        th0 = start_data.get('th0', 0)
        kappa = start_data.get('kappa', 0)
        if abs(kappa) < 1e-6: 
            x = x0 + s_val * math.cos(th0)
            y = y0 + s_val * math.sin(th0)
            th = th0
        else: 
            x = x0 + (math.sin(th0 + kappa * s_val) - math.sin(th0)) / kappa
            y = y0 - (math.cos(th0 + kappa * s_val) - math.cos(th0)) / kappa
            th = th0 + kappa * s_val
        return x, y, th
    
    return 0.0, 0.0, 0.0

def get_closest_segment_exact(pt, segs):
    min_d = 9999.0
    best_idx = 0
    best_s = 0.0
    pt = np.array(pt)
    
    for idx, seg in enumerate(segs):
        if seg.get('is_divided'):
            temp_seg = seg['parent_seg']
            L_orig = temp_seg.get('L', 0.0)
        else:
            temp_seg = seg
            L_orig = temp_seg.get('L', 0.0)
            
        if temp_seg.get('Shape Type') == 'Straight Line' and 'abs_p1' in temp_seg:
            p1 = np.array(temp_seg['abs_p1'])
            p2 = np.array(temp_seg['abs_p2'])
            v = p2 - p1
            w = pt - p1
            c2 = np.dot(v, v)
            ratio = np.dot(w, v) / c2 if c2 > 1e-6 else 0.0
            ratio = max(0.0, min(1.0, ratio))
            proj = p1 + ratio * v
            d = np.linalg.norm(pt - proj)
            if d < min_d:
                min_d = d
                best_idx = idx
                best_s = ratio * L_orig
                if seg.get('is_divided'):
                    best_s -= seg.get('parent_offset', 0.0)
                
        elif temp_seg.get('Shape Type') == 'Curve (Arc & Radius)' and 'abs_c' in temp_seg:
            c = np.array(temp_seg['abs_c'])
            r = temp_seg['abs_r']
            v = pt - c
            ang = math.atan2(v[1], v[0])
            sa = temp_seg['abs_sa']
            sweep = temp_seg['sweep']
            ang_norm = (ang - sa) % (2 * math.pi)
            if ang_norm > abs(sweep):
                ratio = 1.0 if ang_norm < math.pi else 0.0
            else:
                ratio = ang_norm / sweep if abs(sweep) > 1e-6 else 0.0
            ratio = max(0.0, min(1.0, ratio))
            current_ang = sa + ratio * sweep
            proj = c + r * np.array([math.cos(current_ang), math.sin(current_ang)])
            d = np.linalg.norm(pt - proj)
            if d < min_d:
                min_d = d
                best_idx = idx
                best_s = ratio * L_orig
                if seg.get('is_divided'):
                    best_s -= seg.get('parent_offset', 0.0)
                
    return min_d, best_idx, best_s

# 💡 التعديل الجذري لأولويات النهايز (PPS > PPH > Rest)
def get_strut_priority(name):
    name_u = name.upper()
    score = 100
    if "PPS" in name_u: score = 10
    elif "PPH" in name_u: score = 20
    elif "TILT" in name_u: score = 30
    elif "MNB" in name_u: score = 40
    elif "MIB" in name_u: score = 50
    
    is_odd = bool(re.search(r'(1|3)\s*\(', name)) or bool(re.search(r'(1|3)$', name.split('(')[0].strip()))
    if is_odd: score += 5
    return score

# =========================================================
# 1. THE DXF PARSER (Universal Rules)
# =========================================================
def extract_dxf_for_interactive(file_bytes):
    if ezdxf is None: return None
    tmp_path = ""
    try:
        try: dxf_str = file_bytes.decode('utf-8')
        except UnicodeDecodeError: dxf_str = file_bytes.decode('cp1252', errors='ignore')
            
        with tempfile.NamedTemporaryFile(delete=False, suffix=".dxf", mode='w', encoding='utf-8') as tmp:
            tmp.write(dxf_str)
            tmp_path = tmp.name
            
        doc = ezdxf.readfile(tmp_path)
        msp = doc.modelspace()
        
        frames, struts, supports = [], [], []
        
        for e in msp:
            lyr = e.dxf.layer.lower()
            dxftype = e.dxftype()
            
            if dxftype in ['POINT', 'CIRCLE', 'INSERT']:
                if dxftype == 'POINT': supports.append({'x': e.dxf.location.x, 'y': e.dxf.location.y})
                elif dxftype == 'CIRCLE': supports.append({'x': e.dxf.center.x, 'y': e.dxf.center.y})
                elif dxftype == 'INSERT': supports.append({'x': e.dxf.insert.x, 'y': e.dxf.insert.y})
                continue 

            entities = list(e.virtual_entities()) if dxftype in ['LWPOLYLINE', 'POLYLINE'] else [e]
            for sub_e in entities:
                sub_type = sub_e.dxftype()
                
                if "push" in lyr or "pull" in lyr:
                    if sub_type == 'LINE':
                        struts.append({'p1': [sub_e.dxf.start.x, sub_e.dxf.start.y], 'p2': [sub_e.dxf.end.x, sub_e.dxf.end.y]})
                else:
                    if sub_type == 'LINE':
                        frames.append({'type': 'line', 'x1': sub_e.dxf.start.x, 'y1': sub_e.dxf.start.y, 'x2': sub_e.dxf.end.x, 'y2': sub_e.dxf.end.y})
                    elif sub_type == 'ARC':
                        frames.append({'type': 'arc', 'c': [sub_e.dxf.center.x, sub_e.dxf.center.y], 'r': sub_e.dxf.radius, 'sa': math.radians(sub_e.dxf.start_angle), 'ea': math.radians(sub_e.dxf.end_angle)})
                        
        return {'frames': frames, 'struts': struts, 'supports': supports}
    except Exception as e: return None
    finally:
        if tmp_path and os.path.exists(tmp_path):
            try: os.remove(tmp_path)
            except: pass

def parse_dxf_to_data(file_bytes):
    raw_data = extract_dxf_for_interactive(file_bytes)
    if not raw_data or not raw_data['frames']: return None
        
    raw_frames = raw_data.get('frames', [])
    raw_struts = raw_data.get('struts', [])
    raw_supports = raw_data.get('supports', [])

    def get_min_x(f):
        if f['type'] == 'line': return min(f['x1'], f['x2'])
        return f['c'][0] - f['r']
    raw_frames.sort(key=get_min_x)

    base_segments = []
    for idx, f in enumerate(raw_frames):
        if f['type'] == 'line':
            p_start, p_end = (f['x1'], f['y1']), (f['x2'], f['y2'])
            if p_start[0] > p_end[0] + 1e-5 or (abs(p_start[0] - p_end[0]) < 1e-5 and p_start[1] > p_end[1]):
                p_start, p_end = p_end, p_start
            dx_line, dy_line = p_end[0]-p_start[0], p_end[1]-p_start[1]
            L, ang = math.hypot(dx_line, dy_line), math.degrees(math.atan2(dy_line, dx_line))
            base_segments.append({
                'name': f"S{idx+1}", 'master_idx': idx, 'type': 'Straight Line', 'Shape Type': 'Straight Line', 'L': L, 
                'start_angle': ang, 'smooth': False, 'is_dxf': True, 'abs_p1': p_start, 'abs_p2': p_end, 'kappa': 0.0, 'is_divided': False
            })
        elif f['type'] == 'arc':
            sa, ea = f['sa'], f['ea']
            if ea < sa: ea += 2 * math.pi
            sweep, L = ea - sa, f['r'] * (ea - sa)
            base_segments.append({
                'name': f"S{idx+1}", 'master_idx': idx, 'type': 'Curve (Arc & Radius)', 'Shape Type': 'Curve (Arc & Radius)', 'L': L, 
                'Radius (R) (m)': f['r'], 'Curvature Direction': "Arching Up ⤴ (Concave)",
                'start_angle': math.degrees(sa + math.pi/2), 'smooth': False, 'is_dxf': True, 
                'abs_c': list(f['c']), 'abs_r': f['r'], 'abs_sa': sa, 'abs_ea': ea, 'sweep': sweep, 'kappa': 1.0/f['r'], 'is_divided': False
            })

    struts_mapped = []
    for s in raw_struts:
        p1, p2 = s['p1'], s['p2']
        top_p, bot_p = (p1, p2) if p1[1] > p2[1] else (p2, p1)
        struts_mapped.append({
            'tx': top_p[0], 'ty': top_p[1], 
            'bx': bot_p[0], 'by': bot_p[1], 
            'sec': list(STRUTS_DB.keys())[0] if STRUTS_DB else "PPH 353 (1.5:3.5m)"
        })

    supps_mapped = []
    for sp in raw_supports:
        supps_mapped.append({'x': sp['x'], 'y': sp['y'], 'type': 'Hinged', 'angle': 0.0})

    return {'base_segments': base_segments, 'struts': struts_mapped, 'supports': supps_mapped}
# =========================================================
# 2. Smart Division & Dynamic Meshing Engine
# =========================================================
def get_approx_xy(segs, s_idx, s_val):
    if s_idx < 0 or s_idx >= len(segs): return 0.0, 0.0
    seg = segs[s_idx]
    if seg.get('is_dxf'):
        px, py, _ = eval_seg_point(seg, s_val)
        return px, py
    return 0.0, 0.0

def perform_smart_division(base_segments, supports, struts):
    cut_points_dict = {i: {0.0, seg['L']} for i, seg in enumerate(base_segments)}
    
    for sp in supports:
        d_min, w_seg, w_s = get_closest_segment_exact((sp['x'], sp['y']), base_segments)
        if d_min < 0.30:
            cut_points_dict[w_seg].add(min(max(w_s, 0.0), base_segments[w_seg]['L']))
            
    for st in struts:
        dt, wt_seg, wt_s = get_closest_segment_exact((st['tx'], st['ty']), base_segments)
        if dt < 0.30: cut_points_dict[wt_seg].add(min(max(wt_s, 0.0), base_segments[wt_seg]['L']))
        
        db, wb_seg, wb_s = get_closest_segment_exact((st['bx'], st['by']), base_segments)
        if db < 0.30: cut_points_dict[wb_seg].add(min(max(wb_s, 0.0), base_segments[wb_seg]['L']))

    divided_segments = []
    sub_letters = "abcdefghijklmnopqrstuvwxyz"
    
    for m_idx, s_vals_set in sorted(cut_points_dict.items()):
        master_seg = base_segments[m_idx]
        sorted_s = sorted(list(s_vals_set))
        num_sub = len(sorted_s) - 1
        
        for k in range(num_sub):
            s_start = sorted_s[k]
            s_end = sorted_s[k+1]
            if s_end - s_start < 1e-4: continue
            
            sub_name = master_seg['name'] if num_sub == 1 else f"{master_seg['name']}-{sub_letters[k % 26]}"
            new_seg = master_seg.copy()
            new_seg.update({
                'name': sub_name,
                'is_divided': True,
                'parent_seg': master_seg,
                'parent_offset': s_start,
                'L': s_end - s_start,
                'master_idx': m_idx
            })
            divided_segments.append(new_seg)
            
    return divided_segments

def build_chain_mesh(segments, seg_sections, loads, struts, supports, base_segments, mesh_size=0.25):
    nodes = []
    elements = []
    nodal_loads = []
    
    # مسافة اللحام 5 سم
    node_tol = 0.05 
    
    def get_or_add_node(x, y):
        for i, n in enumerate(nodes):
            if abs(n[0] - x) < node_tol and abs(n[1] - y) < node_tol:
                return i
        nodes.append([x, y])
        return len(nodes) - 1

    key_nodes = set()
    support_injections = {i: [] for i in range(len(base_segments))}
    strut_injections = {i: [] for i in range(len(base_segments))}
    supports_list_out = []
    
    for sup in supports:
        sx, sy = sup['x'], sup['y']
        d_min, w_seg, w_s = get_closest_segment_exact((sx, sy), base_segments)
        if d_min < 0.30:
            nx, ny, _ = eval_seg_point(base_segments[w_seg], w_s)
            support_injections[w_seg].append(w_s)
        else:
            nx, ny = sx, sy
            
        nid = get_or_add_node(nx, ny)
        supports_list_out.append({'node': nid, 'type': sup.get('type', 'Hinged'), 'angle': sup.get('angle', 0.0)})

    for st_item in struts:
        dt, wt_seg, wt_s = get_closest_segment_exact((st_item['tx'], st_item['ty']), base_segments)
        if dt < 0.30:
            strut_injections[wt_seg].append(wt_s)
            st_item['tx'], st_item['ty'], _ = eval_seg_point(base_segments[wt_seg], wt_s)
            
        db, wb_seg, wb_s = get_closest_segment_exact((st_item['bx'], st_item['by']), base_segments)
        if db < 0.30:
            strut_injections[wb_seg].append(wb_s)
            st_item['bx'], st_item['by'], _ = eval_seg_point(base_segments[wb_seg], wb_s)

    for i, seg in enumerate(segments):
        L = seg.get('L', 0.0)
        key_s_vals = [0.0, L]
        seg_m_idx = seg.get('master_idx')
        parent_off = seg.get('parent_offset', 0.0)

        for s_val in strut_injections.get(seg_m_idx, []):
            s_local = s_val - parent_off
            if -1e-4 <= s_local <= L + 1e-4:
                key_s_vals.append(max(0.0, min(L, s_local)))
                
        for s_val in support_injections.get(seg_m_idx, []):
            s_local = s_val - parent_off
            if -1e-4 <= s_local <= L + 1e-4:
                key_s_vals.append(max(0.0, min(L, s_local)))

        for ld in loads:
            if ld.get('seg_idx') == i:
                key_s_vals.extend([ld['start'], ld['end']])
                
        num_sub = max(1, int(np.ceil(L / mesh_size)))
        for p in np.linspace(0, L, num_sub+1): 
            key_s_vals.append(p)
            
        keys = sorted(list(set([min(max(round(k, 4), 0.0), round(L, 4)) for k in key_s_vals])))
        node_indices = []
        for s_val in keys:
            px, py, _ = eval_seg_point(seg, s_val)
            nid = get_or_add_node(px, py)
            node_indices.append(nid)
            if any(abs(s_val - round(kv, 4)) < 1e-3 for kv in key_s_vals): 
                key_nodes.add(nid)
            
        sec_props = seg_sections[i] if i < len(seg_sections) else seg_sections[0]
        
        for j in range(len(keys)-1):
            n1 = node_indices[j]
            n2 = node_indices[j+1]
            if n1 == n2: continue 
                
            s_mid = (keys[j] + keys[j+1]) / 2.0
            _, _, th_mid = eval_seg_point(seg, s_mid)
            c_t, s_t = np.cos(th_mid), np.sin(th_mid)
            p_x1, p_y1, p_x2, p_y2 = 0.0, 0.0, 0.0, 0.0
            
            for ld in loads:
                if ld.get('seg_idx') == i and ld.get('type') != 'Point Load':
                    if ld['start'] - 1e-4 <= s_mid <= ld['end'] + 1e-4:
                        L_ld = max(ld['end'] - ld['start'], 1e-5)
                        wa = ld['w1'] + (ld['w2'] - ld['w1']) * (keys[j] - ld['start']) / L_ld
                        wb = ld['w1'] + (ld['w2'] - ld['w1']) * (keys[j+1] - ld['start']) / L_ld
                        
                        dir_str = ld.get('dir', '')
                        if 'Global Z' in dir_str or 'Global Y' in dir_str:
                            p_x1 += wa * s_t; p_y1 += wa * c_t; p_x2 += wb * s_t; p_y2 += wb * c_t
                        elif 'Global X' in dir_str:
                            p_x1 += wa * c_t; p_y1 -= wa * s_t; p_x2 += wb * c_t; p_y2 -= wb * s_t
                        else:
                            p_y1 += wa; p_y2 += wb
                            
            elements.append({
                'type': 'frame', 'group': 'segment', 'sec': sec_props['name'],
                'n1': n1, 'n2': n2, 'px1': p_x1, 'py1': p_y1, 'px2': p_x2, 'py2': p_y2,
                'E': sec_props['E'] * 10000.0, 'A': sec_props['A'], 'I': sec_props['I'] / 100000000.0,
                'seg_idx': i, 's_start': keys[j], 's_end': keys[j+1], 'L': keys[j+1] - keys[j], 'th_mid': th_mid
            })
            
        for ld in loads:
            if ld.get('seg_idx') == i and ld.get('type') == 'Point Load':
                px, py, th_pt = eval_seg_point(seg, ld['start'])
                nid = get_or_add_node(px, py)
                dir_str = ld.get('dir', '')
                if 'Global Z' in dir_str or 'Global Y' in dir_str: nodal_loads.append({'node': nid, 'Fx': 0.0, 'Fy': ld['w1']})
                elif 'Global X' in dir_str: nodal_loads.append({'node': nid, 'Fx': ld['w1'], 'Fy': 0.0})
                else: 
                    c_pt, s_pt = np.cos(th_pt), np.sin(th_pt)
                    nodal_loads.append({'node': nid, 'Fx': -ld['w1']*s_pt, 'Fy': ld['w1']*c_pt})

    for st_idx, st_item in enumerate(struts):
        top_node = get_or_add_node(st_item['tx'], st_item['ty'])
        bot_node = get_or_add_node(st_item['bx'], st_item['by'])
        
        elements.append({
            'type': 'truss', 'group': 'strut', 'sec': st_item.get('sec', 'Unknown'),
            'n1': bot_node, 'n2': top_node, 'strut_idx': st_idx, 'E': 21000000.0, 'A': 0.001
        })

    display_nodes = set([s['node'] for s in supports_list_out])
    display_nodes.update(key_nodes) 

    return nodes, elements, nodal_loads, display_nodes, supports_list_out

# =========================================================
# 3. Advanced FEA Solver (Exact Matrix Engine)
# =========================================================
def solve_fea_engine(nodes, elements, nodal_loads, supports_list):
    NDOF = len(nodes) * 3
    K = np.zeros((NDOF, NDOF))
    F = np.zeros(NDOF)
    
    for el in elements:
        n1, n2 = el['n1'], el['n2']
        x1, y1 = nodes[n1][0], nodes[n1][1]
        x2, y2 = nodes[n2][0], nodes[n2][1]
        L = np.hypot(x2 - x1, y2 - y1)
        
        if L < 1e-5: 
            el['c'], el['s'], el['L'] = 1, 0, 1e-5
            el['internal'] = {'N': [0,0], 'V': [0,0], 'M': [0,0], 'x': [0, 1e-5]}
            continue
            
        c, s = (x2 - x1) / L, (y2 - y1) / L
        el['L'], el['c'], el['s'] = L, c, s
        E, A, I = el['E'], el['A'], el.get('I', 0.00005)
        
        T = np.array([
            [c, s, 0, 0, 0, 0], [-s, c, 0, 0, 0, 0], [0, 0, 1, 0, 0, 0],
            [0, 0, 0, c, s, 0], [0, 0, 0, -s, c, 0], [0, 0, 0, 0, 0, 1]
        ])
        
        if el['type'] == 'truss':
            k_loc = np.zeros((6, 6))
            k_loc[0,0] = E * A / L; k_loc[3,3] = E * A / L
            k_loc[0,3] = -E * A / L; k_loc[3,0] = -E * A / L
        else:
            k_loc = np.array([
                [E*A/L, 0, 0, -E*A/L, 0, 0], 
                [0, 12*E*I/L**3, 6*E*I/L**2, 0, -12*E*I/L**3, 6*E*I/L**2],
                [0, 6*E*I/L**2, 4*E*I/L, 0, -6*E*I/L**2, 2*E*I/L], 
                [-E*A/L, 0, 0, E*A/L, 0, 0],
                [0, -12*E*I/L**3, -6*E*I/L**2, 0, 12*E*I/L**3, -6*E*I/L**2], 
                [0, 6*E*I/L**2, 2*E*I/L, 0, -6*E*I/L**2, 4*E*I/L]
            ])
            px1, py1 = el.get('px1',0), el.get('py1',0)
            px2, py2 = el.get('px2',0), el.get('py2',0)
            f_loc = np.array([
                (2*px1 + px2)*L/6.0, (7*py1 + 3*py2)*L/20.0, (3*py1 + 2*py2)*L**2/60.0,
                (px1 + 2*px2)*L/6.0, (3*py1 + 7*py2)*L/20.0, -(2*py1 + 3*py2)*L**2/60.0
            ])
            f_glob = T.T @ f_loc
            dof = [3*n1, 3*n1+1, 3*n1+2, 3*n2, 3*n2+1, 3*n2+2]
            for r in range(6): F[dof[r]] += f_glob[r]
            
        k_glob = T.T @ k_loc @ T
        dof = [3*n1, 3*n1+1, 3*n1+2, 3*n2, 3*n2+1, 3*n2+2]
        for r in range(6):
            for col in range(6): K[dof[r], dof[col]] += k_glob[r, col]
                
    for nl in nodal_loads:
        F[3*nl['node']] += nl['Fx']
        F[3*nl['node']+1] += nl['Fy']
            
    # حساب إجمالي الحمل الرأسي للتأكد من الاتزان
    net_load_y = abs(np.sum(F[1::3]))

    K_orig = K.copy()
    fixed_dofs = []
    K_pen = 1e12
    
    for sup in supports_list:
        n, t, a = sup['node'], sup['type'], sup.get('angle', 0.0)
        if t == 'Fixed': fixed_dofs.extend([3*n, 3*n+1, 3*n+2])
        elif t == 'Hinged': fixed_dofs.extend([3*n, 3*n+1])
        elif t == 'Roller':
            rad = np.radians(a)
            nx, ny = -np.sin(rad), np.cos(rad) 
            K[3*n, 3*n] += K_pen*nx**2; K[3*n+1, 3*n+1] += K_pen*ny**2
            K[3*n, 3*n+1] += K_pen*nx*ny; K[3*n+1, 3*n] += K_pen*nx*ny

    free_dof = [i for i in range(NDOF) if i not in fixed_dofs]
    K_ff = K[np.ix_(free_dof, free_dof)]
    F_f = F[free_dof]
    
    U = np.zeros(NDOF)
    try: U[free_dof] = np.linalg.solve(K_ff, F_f)
    except np.linalg.LinAlgError: U[free_dof] = np.linalg.lstsq(K_ff, F_f, rcond=None)[0]
    
    R_reactions = K_orig @ U - F 
    
    for el in elements:
        if el.get('L', 0) < 1e-5: continue
        n1, n2 = el['n1'], el['n2']
        c, s, L, E, A, I = el['c'], el['s'], el['L'], el['E'], el['A'], el.get('I', 0.00005)
        
        u_glob = U[[3*n1, 3*n1+1, 3*n1+2, 3*n2, 3*n2+1, 3*n2+2]]
        T = np.array([
            [c, s, 0, 0, 0, 0], [-s, c, 0, 0, 0, 0], [0, 0, 1, 0, 0, 0],
            [0, 0, 0, c, s, 0], [0, 0, 0, -s, c, 0], [0, 0, 0, 0, 0, 1]
        ])
        u_loc = T @ u_glob
        el['internal'] = {'u_loc': u_loc}
        
        if el['type'] == 'truss':
            N_val = (E * A / L) * (u_loc[3] - u_loc[0])
            xs = np.linspace(0, L, 51)
            el['internal'].update({
                'N': np.full_like(xs, N_val), 'V': np.zeros_like(xs), 'M': np.zeros_like(xs),
                'x': xs
            })
        else:
            k_loc = np.array([
                [E*A/L, 0, 0, -E*A/L, 0, 0], [0, 12*E*I/L**3, 6*E*I/L**2, 0, -12*E*I/L**3, 6*E*I/L**2],
                [0, 6*E*I/L**2, 4*E*I/L, 0, -6*E*I/L**2, 2*E*I/L], [-E*A/L, 0, 0, E*A/L, 0, 0],
                [0, -12*E*I/L**3, -6*E*I/L**2, 0, 12*E*I/L**3, -6*E*I/L**2], [0, 6*E*I/L**2, 2*E*I/L, 0, -6*E*I/L**2, 4*E*I/L]
            ])
            px1, py1, px2, py2 = el.get('px1',0), el.get('py1',0), el.get('px2',0), el.get('py2',0)
            f_loc = np.array([
                (2*px1 + px2)*L/6.0, (7*py1 + 3*py2)*L/20.0, (3*py1 + 2*py2)*L**2/60.0,
                (px1 + 2*px2)*L/6.0, (3*py1 + 7*py2)*L/20.0, -(2*py1 + 3*py2)*L**2/60.0
            ])
            f_end = k_loc @ u_loc - f_loc
            
            xs = np.linspace(0, L, 51) 
            N_arr, V_arr, M_arr = np.zeros_like(xs), np.zeros_like(xs), np.zeros_like(xs)
            
            for i, x in enumerate(xs):
                N_arr[i] = -f_end[0] - (px1*x + (px2-px1)*x**2/(2*L))
                V_arr[i] = f_end[1] + (py1*x + (py2-py1)*x**2/(2*L))
                M_arr[i] = -f_end[2] + f_end[1]*x + py1*x**2/2.0 + (py2-py1)*x**3/(6*L)
                
            el['internal'].update({'N': N_arr, 'V': V_arr, 'M': M_arr, 'x': xs})
            
    return U, R_reactions, net_load_y

# =====================================================================
# 🧠 THE HEURISTIC OPTIMIZER ENGINE (PHYSICS-AWARE V14)
# =====================================================================
def run_auto_optimizer(base_segments, working_segments, active_seg_sections, ui_struts, ui_loads, target_rxn, spacings_str, view_plane, auto_mesh_size, dxf_v, is_symmetric):
    try:
        spacings = sorted([float(x.strip()) for x in spacings_str.split(',')], reverse=True)
    except: return False, None, "❌ Format error in spacings."
    
    if not base_segments: return False, None, "❌ No base segments found."
    
    min_y, max_x, min_x = 9999.0, -9999.0, 9999.0
    for seg in base_segments:
        if seg.get('is_dxf') and seg.get('Shape Type') == 'Straight Line':
            p1, p2 = seg['abs_p1'], seg['abs_p2']
            min_y = min(min_y, p1[1], p2[1])
            max_x = max(max_x, p1[0], p2[0])
            min_x = min(min_x, p1[0], p2[0])
        elif seg.get('is_dxf') and seg.get('Shape Type') == 'Curve (Arc & Radius)':
            c, r = seg['abs_c'], seg['abs_r']
            min_y = min(min_y, c[1]-r)
            max_x = max(max_x, c[0]+r)
            min_x = min(min_x, c[0]-r)
            
    if min_y == 9999.0: return False, None, "❌ Base geometry not found."
    
    center_x = (min_x + max_x) / 2.0
    bridge_width = max_x - min_x
    half_width = bridge_width / 2.0
    
    strut_xs = set()
    for ds in ui_struts:
        strut_xs.add(ds.get('bx', 0.0))
        strut_xs.add(ds.get('tx', 0.0))
        
    test_combined_loads = []
    fac_d = float(st.session_state.get("cmb_d", 1.00))
    fac_l = float(st.session_state.get("cmb_l", 1.00))
    fac_w = float(st.session_state.get("cmb_w", 1.00))
    combo_factors = {'Dead Load': fac_d, 'Live Load': fac_l, 'Wind Load': fac_w} 
    
    working_seg_names = [s.get('name', f"S{i+1}") for i, s in enumerate(working_segments)]
    base_seg_names = [s.get('name', f"S{i+1}") for i, s in enumerate(base_segments)]
    
    for i, ld in enumerate(ui_loads):
        t_mode = ld.get('target_mode', 'All Segments')
        target_base_indices = []
        if t_mode == "Single Segment":
            s_choice = st.session_state.get(f"ld_single_{i}", base_seg_names[0])
            if s_choice in base_seg_names: target_base_indices.append(base_seg_names.index(s_choice))
        elif t_mode == "Multiple Segments":
            raw_def_segs = st.session_state.get(f"ld_multi_{i}", ld.get('target_segs', []))
            safe_def_segs = [s for s in raw_def_segs if s in base_seg_names]
            target_base_indices = [base_seg_names.index(s) for s in safe_def_segs]
        else:
            target_base_indices = list(range(len(base_segments)))
            
        w1 = float(st.session_state.get(f"ld_w1_{i}_{dxf_v}", ld.get('w1', 0.0)))
        w2 = float(st.session_state.get(f"ld_w2_{i}_{dxf_v}", ld.get('w2', 0.0))) if ld.get('type') == 'Trapezoidal' else w1
        
        target_working_indices = [w_idx for w_idx, w_seg in enumerate(working_segments) if w_seg.get('master_idx', 0) in target_base_indices]
        
        for s_idx_num in target_working_indices:
            max_s = float(working_segments[s_idx_num].get('L', 0.0))
            test_combined_loads.append({
                'seg_idx': s_idx_num, 'category': ld['category'], 'type': ld['type'], 'dir': ld['dir'], 
                'start': 0.0, 'end': max_s if ld['type'] != 'Point Load' else 0.0, 
                'w1': w1 * combo_factors[ld['category']], 'w2': w2 * combo_factors[ld['category']]
            })

    def run_trial(test_supps):
        nodes_t, elems_t, nloads_t, _, slist_t = build_chain_mesh(
            working_segments, active_seg_sections, test_combined_loads,
            ui_struts, test_supps, base_segments, mesh_size=auto_mesh_size
        )
        U, R, net_load = solve_fea_engine(nodes_t, elems_t, nloads_t, slist_t)
        
        max_ry, min_ry = 0.0, 9999.0
        for sup in slist_t:
            ry = R[3*sup['node'] + 1] 
            if ry > max_ry: max_ry = ry
            if ry < min_ry: min_ry = ry
        return max_ry, min_ry, net_load

    # 💡 1. حساب إجمالي الحمل لإقصاء أي شبكة لا تفي بالعدد المطلوب (Mathematical Pruner)
    dummy_supps = [{'x': min_x, 'y': min_y, 'type': 'Hinged'}, {'x': max_x, 'y': min_y, 'type': 'Roller'}]
    _, _, total_system_load = run_trial(dummy_supps)
    min_required_props = max(2, int(math.ceil(total_system_load / target_rxn)))

    valid_grids = []
    
    if is_symmetric:
        def build_sym_grids(current_grid):
            cantilever = half_width - current_grid[-1]
            if 0.40 <= cantilever <= 1.50:
                full_grid = set(current_grid)
                for x in current_grid:
                    if x > 1e-4: full_grid.add(-x)
                actual_coords = [round(center_x + x, 3) for x in sorted(list(full_grid))]
                valid_grids.append(tuple(actual_coords))
            if cantilever < 0.40: return
            for s in spacings:
                build_sym_grids(current_grid + [current_grid[-1] + s])
        build_sym_grids([0.0])
        for s in spacings: build_sym_grids([s / 2.0])
        
    else:
        def build_asym_grids(current_grid):
            cantilever = max_x - current_grid[-1]
            if 0.40 <= cantilever <= 1.50:
                actual_coords = [round(x, 3) for x in current_grid if min_x - 0.05 <= x <= max_x + 0.05]
                valid_grids.append(tuple(actual_coords))
            if cantilever < 0.40: return
            for s in spacings:
                build_asym_grids(current_grid + [current_grid[-1] + s])
                
        left_cantilevers = np.arange(0.40, 1.51, 0.10)
        for lc in left_cantilevers:
            build_asym_grids([min_x + lc])

    unique_grids = list(set(valid_grids))
    
    # 💡 التصفية: إقصاء أي رصة بها عدد دعامات أقل من المسموح رياضياً لتجنب تضييع الوقت
    filtered_grids = [list(g) for g in unique_grids if len(g) >= min_required_props]

    if not filtered_grids:
        return False, None, f"❌ Impossible! Requires at least {min_required_props} props to handle {total_system_load:.1f} kN safely, but your spacings can't fit them. Decrease spacings."

    # 💡 2. التقييم العبقري: معاقبة الدعامات الزائدة + مكافأة ضخمة للرص تحت النهايز
    def grid_score(grid):
        prop_count = len(grid)
        prop_penalty = (prop_count - min_required_props) * 20.0 
        alignment_bonus = 0.0
        
        for sx in strut_xs:
            if grid:
                min_dist = min([abs(px - sx) for px in grid])
                if min_dist <= 0.15: alignment_bonus += 50.0   # مكافأة ضخمة جداً (أولوية قصوى)
                elif min_dist <= 0.40: alignment_bonus += 20.0
                elif min_dist <= 0.80: alignment_bonus += 5.0
                
        return prop_penalty - alignment_bonus

    filtered_grids.sort(key=grid_score)

    start_time = time.time()
    max_time = 90.0
    best_fallback_grid = None
    best_fallback_score = 999999.0
    trials_count = 0
    
    for actual_coords in filtered_grids:
        if time.time() - start_time > max_time: break
        
        test_supps = []
        for idx, gx in enumerate(actual_coords):
            stype = 'Hinged' if idx == 0 else 'Roller'
            test_supps.append({'x': gx, 'y': round(min_y, 3), 'type': stype, 'angle': 0.0})
            
        max_ry, min_ry, _ = run_trial(test_supps)
        trials_count += 1
        
        # 💡 النجاح المطلق
        if max_ry <= target_rxn and min_ry >= 0.5:
            sym_txt = "Symmetric" if is_symmetric else "Asymmetric"
            return True, test_supps, f"✅ Optimum {sym_txt} grid found! (0 Uplift, {len(test_supps)} Props. Tested {trials_count} combos)"
            
        # حفظ الطوارئ
        if min_ry >= 0.5:
            if max_ry < best_fallback_score:
                best_fallback_score = max_ry
                best_fallback_grid = test_supps
                
    # 💡 الرفض القاطع لتخطي الريأكشن، مع إرسال الرصة للعرض كخطأ واضح
    if best_fallback_grid:
        return False, best_fallback_grid, f"❌ STRICT LIMIT FAILED: Best safe grid reaction is {best_fallback_score:.2f} kN (Target: {target_rxn} kN). Try smaller spacings."
        
    return False, None, f"❌ Optimizer failed completely! All {trials_count} tested grids generated UPLIFT (Instability)."
# =========================================================
# 4. Plotting Engine (Live Preview & Diagrams Engine)
# =========================================================
def draw_base_geometry(ax, nodes, elements, supports_list, seg_sections=None, segments=None, show_seg_names=False):
    for el in elements:
        if el['type'] not in ['frame', 'truss']:
            continue
            
        n1, n2 = nodes[el['n1']], nodes[el['n2']]
        if el['type'] == 'truss':
            ax.plot([n1[0], n2[0]], [n1[1], n2[1]], color='red', linestyle='-', linewidth=0.8, zorder=1)
        else:
            if el.get('group') == 'base' and el.get('sec') == "None (Direct to Ground)":
                continue
            ax.plot([n1[0], n2[0]], [n1[1], n2[1]], color='royalblue', linestyle='-', linewidth=1.5, zorder=1)
            
    for i, sup in enumerate(supports_list):
        n, x, y, t, ang_deg = sup['node'], nodes[sup['node']][0], nodes[sup['node']][1], sup['type'], sup.get('angle', 0.0)
        ax.text(x, y - 0.4, f"J{i+1}", color='green', fontsize=7, ha='center', fontname='Arial')
        
        ang_rad = math.radians(ang_deg)
        c_a, s_a = math.cos(ang_rad), math.sin(ang_rad)
        
        def rot_pt(px, py):
            rx = x + (px - x)*c_a - (py - y)*s_a
            ry = y + (px - x)*s_a + (py - y)*c_a
            return rx, ry

        if t == 'Fixed':
            p1, p2, p3, p4 = rot_pt(x-0.1, y-0.1), rot_pt(x+0.1, y-0.1), rot_pt(x+0.1, y+0.1), rot_pt(x-0.1, y+0.1)
            ax.add_patch(Polygon([p1, p2, p3, p4], facecolor='none', edgecolor='limegreen', lw=1.0, zorder=5))
            l1, l2 = rot_pt(x-0.1, y), rot_pt(x+0.1, y)
            ax.plot([l1[0], l2[0]], [l1[1], l2[1]], color='limegreen', lw=1.0, zorder=4)
        elif t == 'Hinged':
            h, w = 0.15, 0.12
            p1, p2, p3 = rot_pt(x, y), rot_pt(x+w, y-h), rot_pt(x-w, y-h)
            ax.add_patch(Polygon([p1, p2, p3], facecolor='none', edgecolor='limegreen', lw=1.0, zorder=5))
            l1, l2 = rot_pt(x-w-0.05, y-h), rot_pt(x+w+0.05, y-h)
            ax.plot([l1[0], l2[0]], [l1[1], l2[1]], color='limegreen', lw=1.0, zorder=4)
        elif t == 'Roller':
            h, w, r = 0.15, 0.12, 0.04
            p1, p2, p3 = rot_pt(x, y), rot_pt(x+w, y-h), rot_pt(x-w, y-h)
            ax.add_patch(Polygon([p1, p2, p3], facecolor='none', edgecolor='limegreen', lw=1.0, zorder=5))
            c_pt = rot_pt(x, y-h-r)
            ax.add_patch(plt.Circle(c_pt, r, facecolor='none', edgecolor='limegreen', lw=1.0, zorder=5))
            l1, l2 = rot_pt(x-w-0.05, y-h-2*r), rot_pt(x+w+0.05, y-h-2*r)
            ax.plot([l1[0], l2[0]], [l1[1], l2[1]], color='limegreen', lw=1.0, zorder=4)

    if seg_sections and segments:
        for el in elements:
            if el['type'] == 'truss':
                n1, n2 = nodes[el['n1']], nodes[el['n2']]
                mid_x, mid_y = (n1[0]+n2[0])/2, (n1[1]+n2[1])/2
                dx, dy = n2[0]-n1[0], n2[1]-n1[1]
                rot = np.degrees(math.atan2(dy, dx))
                if rot > 90: rot -= 180
                elif rot < -90: rot += 180
                L_hyp = np.hypot(dx, dy)
                if L_hyp > 1e-4:
                    nx_s, ny_s = -dy/L_hyp, dx/L_hyp
                    st_id = el.get('strut_idx', 0) + 1
                    ax.text(mid_x + nx_s*0.1, mid_y + ny_s*0.1, f"P{st_id}: {get_short_name(el.get('sec', ''))}", color='dimgray', fontsize=6, rotation=rot, ha='center', va='center', fontname='Arial')
        
        if show_seg_names:
            for i, seg in enumerate(segments):
                mx, my, mth = eval_seg_point(seg, seg.get('L', 0)/2)
                rot_deg = math.degrees(mth)
                if rot_deg > 90: rot_deg -= 180
                elif rot_deg < -90: rot_deg += 180
                seg_name = seg.get('name', f"S{i+1}")
                ax.text(mx - math.sin(mth)*0.3, my + math.cos(mth)*0.3, seg_name, color='dimgray', fontsize=6, ha='center', va='center', rotation=rot_deg, fontname='Arial')

    if len(supports_list) > 1:
        sup_xs = []
        sup_ys = []
        for sup in supports_list:
            n_idx = sup['node']
            sup_xs.append(nodes[n_idx][0])
            sup_ys.append(nodes[n_idx][1])
            
        unique_xs = sorted(list(set([round(x, 3) for x in sup_xs])))
        if len(unique_xs) > 1:
            min_y = min(sup_ys)
            dim_y = min_y - 0.9 
            dyn_fontsize = max(4.0, min(7.0, 70.0 / len(unique_xs)))
            
            ax.plot([unique_xs[0], unique_xs[-1]], [dim_y, dim_y], color='gray', lw=0.5, zorder=1)
            
            for i in range(len(unique_xs)):
                ax.plot([unique_xs[i], unique_xs[i]], [dim_y - 0.15, dim_y + 0.15], color='gray', lw=0.5, zorder=1)
                if i < len(unique_xs) - 1:
                    dist = unique_xs[i+1] - unique_xs[i]
                    mid_x = (unique_xs[i] + unique_xs[i+1]) / 2.0
                    ax.text(mid_x, dim_y + 0.05, f"{dist:.2f}", color='dimgray', fontsize=dyn_fontsize, 
                            ha='center', va='bottom', fontname='Arial')

def draw_loads_and_geometry(ax, nodes, elements, supports_list, seg_sections, loads, segments=None, load_cat_filter=None, show_seg_names=False):
    draw_base_geometry(ax, nodes, elements, supports_list, seg_sections, segments, show_seg_names)
    scale_ld = 0.05
    
    color_map = {
        'Dead Load': ('blue', 0.15),
        'Live Load': ('red', 0.15),
        'Wind Load': ('green', 0.20)
    }
    
    for ld in loads:
        if load_cat_filter and ld.get('category') != load_cat_filter:
            continue
        if abs(ld.get('w1', 0)) < 1e-4 and abs(ld.get('w2', 0)) < 1e-4:
            continue
            
        ld_color, ld_alpha = color_map.get(ld.get('category', 'Dead Load'), ('blue', 0.15))
            
        if segments:
            i = ld.get('seg_idx', 0)
            if i >= len(segments): continue
            
            w1, w2 = ld.get('w1', 0), ld.get('w2', 0)
            num_pts = max(10, int((ld.get('end', 0) - ld.get('start', 0)) / 0.1))
            s_vals = np.linspace(ld.get('start', 0), ld.get('end', 0), num_pts)
            poly_pts, top_pts = [], []
            
            for sv in s_vals:
                px, py, th = eval_seg_point(segments[i], sv)
                w_curr = w1 + (w2 - w1) * (sv - ld.get('start', 0)) / max(1e-5, (ld.get('end', 0) - ld.get('start', 0)))
                w_val = w_curr * scale_ld
                poly_pts.append((px, py))
                
                dir_str = ld.get('dir', '')
                if 'Global Z' in dir_str or 'Global Y' in dir_str: f_vx, f_vy = 0.0, w_val
                elif 'Global X' in dir_str: f_vx, f_vy = w_val, 0.0
                else:
                    c, s = math.cos(th), math.sin(th)
                    f_vx, f_vy = -s * w_val, c * w_val
                top_pts.append((px - f_vx, py - f_vy))
                    
            poly_pts.extend(top_pts[::-1])
            if len(poly_pts) > 2:
                ax.add_patch(Polygon(poly_pts, facecolor=ld_color, edgecolor=ld_color, alpha=ld_alpha, lw=0.8, zorder=2))
                ax.add_patch(Polygon(poly_pts, facecolor='none', edgecolor=ld_color, lw=0.8, zorder=3))

def draw_live_preview(nodes, elements, supports_list, seg_sections, loads, segments=None):
    apply_plot_styles()
    fig, ax = plt.subplots(figsize=(8, 5))
    ax.set_aspect('equal', adjustable='datalim'); ax.axis('off')
    draw_loads_and_geometry(ax, nodes, elements, supports_list, seg_sections, loads, segments, show_seg_names=True)
    return safe_render_fig(fig)

def plot_sap2000_diagrams(nodes, elements, R_reactions, scales, display_nodes, supports_list, seg_sections, loads, segments=None):
    apply_plot_styles()
    figs_dict = {}
    
    # 1. Loads
    load_cats = ['Dead Load', 'Live Load', 'Wind Load']
    for cat in load_cats:
        has_load = any(ld.get('category') == cat and (abs(ld.get('w1', 0)) > 1e-4 or abs(ld.get('w2', 0)) > 1e-4) for ld in loads)
        if has_load:
            fig_ld, ax_ld = plt.subplots(figsize=(7, 4.5))
            ax_ld.set_aspect('equal', adjustable='datalim'); ax_ld.axis('off')
            draw_loads_and_geometry(ax_ld, nodes, elements, supports_list, seg_sections, loads, segments, load_cat_filter=cat, show_seg_names=False)
            figs_dict[f'Load_{cat}'] = safe_render_fig(fig_ld)
    
    # 2. Reactions
    fig_r, ax_r = plt.subplots(figsize=(7, 4.5))
    ax_r.set_aspect('equal', adjustable='datalim'); ax_r.axis('off')
    draw_base_geometry(ax_r, nodes, elements, supports_list, seg_sections, segments, show_seg_names=False)
    
    for sup in supports_list:
        n, t, ang = sup['node'], sup['type'], sup.get('angle', 0.0)
        Rx, Ry = R_reactions[3*n], R_reactions[3*n+1]
        x, y = nodes[n][0], nodes[n][1]
        
        c_a, s_a = math.cos(math.radians(ang)), math.sin(math.radians(ang))
        R_loc_x = Rx * c_a + Ry * s_a
        R_loc_y = -Rx * s_a + Ry * c_a
        
        if t == 'Roller': 
            draw_reaction_arrow(ax_r, x, y, R_loc_y, -s_a, c_a)
        else:
            draw_reaction_arrow(ax_r, x, y, R_loc_x, c_a, s_a)
            draw_reaction_arrow(ax_r, x, y, R_loc_y, -s_a, c_a)
    figs_dict['React'] = safe_render_fig(fig_r)
    
    # 3. Forces Engine
    def create_force_plot(val_key, scale, c_pos, c_neg):
        fig_f, ax_f = plt.subplots(figsize=(7, 4.5))
        ax_f.set_aspect('equal', adjustable='datalim'); ax_f.axis('off')
        draw_base_geometry(ax_f, nodes, elements, supports_list, seg_sections, segments, show_seg_names=False)
        global_texts = []
        
        def is_far(tx, ty):
            for (px, py) in global_texts:
                if math.hypot(tx-px, ty-py) < 0.35: return False
            return True

        for el in elements:
            n1, n2 = el['n1'], el['n2']
            x1, y1 = nodes[n1][0], nodes[n1][1]
            x2, y2 = nodes[n2][0], nodes[n2][1]
            c, s = el.get('c', 1), el.get('s', 0)
            xs = el.get('internal', {}).get('x', [])
            vals = el.get('internal', {}).get(val_key, [])
            if len(vals) == 0 or np.all(np.abs(vals) < 1e-6): continue
            
            plot_vals = -vals if val_key != 'N' else vals
                
            px = x1 + c * xs - s * plot_vals * scale
            py = y1 + s * xs + c * plot_vals * scale
            
            for k in range(len(px)-1):
                color = c_pos if vals[k] >= 0 else c_neg
                ax_f.plot([px[k], px[k+1]], [py[k], py[k+1]], color=color, lw=0.8)
                
            ax_f.plot([x1, px[0]], [y1, py[0]], color=c_pos if vals[0]>=0 else c_neg, lw=0.8)
            ax_f.plot([x2, px[-1]], [y2, py[-1]], color=c_pos if vals[-1]>=0 else c_neg, lw=0.8)

            def plot_val(idx):
                v_disp = vals[idx]
                if abs(v_disp) < 0.1: return
                
                tx, ty = px[idx], py[idx]
                sgn = 1 if plot_vals[idx] >= 0 else -1
                tx += -s * sgn * 0.15; ty += c * sgn * 0.15
                v_color = c_pos if vals[idx] >= 0 else c_neg
                
                if is_far(tx, ty):
                    ax_f.text(tx, ty, f"{v_disp:+.2f}", fontsize=6, color=v_color, ha='center', va='center', fontname='Arial')
                    global_texts.append((tx, ty))
                    
            if len(vals) > 0: 
                plot_val(len(vals)//2)
                max_idx = np.argmax(np.abs(vals))
                plot_val(max_idx)
                
        return safe_render_fig(fig_f)

    figs_dict['N'] = create_force_plot('N', scales['N'], 'blue', 'red')
    figs_dict['V'] = create_force_plot('V', scales['V'], 'blue', 'red')
    figs_dict['M'] = create_force_plot('M', scales['M'], 'blue', 'red')
    
    return figs_dict

def generate_chain_report(sys_data):
    doc = Document("Acrow_Template.docx") if os.path.exists("Acrow_Template.docx") else Document()
    
    def force_ltr_left(p):
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        pPr = p._element.get_or_add_pPr()
        bidi = OxmlElement('w:bidi'); bidi.set(qn('w:val'), '0'); pPr.append(bidi)
        
    def add_line(text, bold=False):
        p = doc.add_paragraph()
        force_ltr_left(p); p.paragraph_format.line_spacing = 1.5
        r = p.add_run(text)
        r.font.name = 'Arial'; r.font.size = Pt(12); r.font.bold = bold; r.font.rtl = False
        
    def add_large_diagram(doc, img_bytes, title):
        p_img = doc.add_paragraph()
        p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r_img = p_img.add_run()
        r_img.add_picture(io.BytesIO(img_bytes), width=Cm(15.0))
        
        p_title = doc.add_paragraph()
        p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r_title = p_title.add_run(title)
        r_title.font.name = 'Arial'; r_title.font.size = Pt(12); r_title.bold = True
        
        p_line = doc.add_paragraph()
        p_line.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r_line = p_line.add_run("_" * 60)
        r_line.font.size = Pt(8); r_line.font.color.rgb = RGBColor(150, 150, 150)
        doc.add_paragraph()

    p_title = doc.add_paragraph(); force_ltr_left(p_title)
    run_title = p_title.add_run("CALCULATION SHEET FOR ADVANCED SHAPES")
    run_title.font.name = 'Arial'; run_title.font.size = Pt(16); run_title.font.bold = True; run_title.font.rtl = False
    
    add_line("="*50, bold=True)
    add_line(f"1. Safety Checks:", bold=True)
    for df_row in sys_data['safety_df']:
        add_line(f"- {df_row['Component']} ({df_row['Force Type']}): {df_row['Actual']} vs {df_row['Allowable']} => {df_row['Status']}")
    
    doc.add_page_break()
    add_line("2. Analysis Diagrams:", bold=True)
    
    bufs = sys_data['img_bufs']
    for cat in ['Dead Load', 'Live Load', 'Wind Load']:
        key = f'Load_{cat}'
        if key in bufs: add_large_diagram(doc, bufs[key], f"{cat} Distribution Diagram")
            
    add_large_diagram(doc, bufs['React'], "Reactions Diagram (kN)")
    add_large_diagram(doc, bufs['N'], "Axial Force Diagram (kN)")
    add_large_diagram(doc, bufs['V'], "Shear Force Diagram (kN)")
    add_large_diagram(doc, bufs['M'], "Bending Moment Diagram (kN.m)")
    
    out = io.BytesIO(); doc.save(out)
    return out
# =========================================================
# 6. Main Streamlit UI (Smart Topology & Dynamic UX)
# =========================================================
def render_advanced_shape_module():
    st.markdown("## 🎢 The Chain Builder (Classic V3 - Ultimate Logic)")
    
    st.markdown("""
        <style>
            div[data-testid="column"]:nth-of-type(2) {
                position: sticky !important;
                top: 4rem !important;
                align-self: flex-start !important;
                z-index: 999 !important;
                height: max-content !important;
            }
        </style>
    """, unsafe_allow_html=True)

    if 'dxf_version' not in st.session_state:
        st.session_state.dxf_version = 0
        
    if 'opt_v' not in st.session_state:
        st.session_state.opt_v = 0

    vp_opts = ["Section View (XZ Axes - Vertical)", "Plan View (XY Axes - Horizontal)"]
    idx_vp = vp_opts.index(st.session_state.get("view_plane_adv", vp_opts[0])) if st.session_state.get("view_plane_adv") in vp_opts else 0
    view_plane = st.radio("📐 Structural Analysis Plane / System Projection", vp_opts, index=idx_vp, key="view_plane_adv", horizontal=True)
    
    c_upload, c_mesh = st.columns([2, 1])
    with c_upload:
        uploaded_dxf = st.file_uploader("📥 Upload DXF File (.dxf)", type=['dxf'], key="dxf_uploader")
    with c_mesh:
        st.write(""); st.write("")
        auto_mesh_size = st.number_input("Auto Frame Mesh Size (m)", min_value=0.05, max_value=5.0, value=float(st.session_state.get("auto_mesh_size_adv", 0.25)), step=0.05, key="auto_mesh_size_adv")

    if uploaded_dxf and st.button("Extract Data from DXF"):
        parsed = parse_dxf_to_data(uploaded_dxf.getvalue())
        if parsed:
            st.session_state.dxf_parsed = parsed
            st.session_state.ui_supports = parsed['supports']
            st.session_state.ui_struts = parsed['struts']
            st.session_state.ui_loads = [{'category': 'Dead Load', 'type': 'Uniform', 'dir': 'Global Z (Vertical)' if "XZ" in view_plane else 'Global Y (Vertical)', 'target_mode': 'All Segments', 'target_segs': [], 'w1': 0.0, 'w2': 0.0}]
            if 'divided_segments' in st.session_state: del st.session_state['divided_segments']
            
            st.session_state.dxf_version += 1
            st.session_state.opt_v += 1
                    
            st.success("✅ DXF Parsed! Base frames extracted perfectly.")
            st.rerun()
        else: 
            st.error("❌ Failed to parse DXF. Please check layers.")

    dxf_data = st.session_state.get('dxf_parsed', None)
    dxf_v = st.session_state.dxf_version
    opt_v = st.session_state.opt_v
    
    if 'ui_supports' not in st.session_state:
        st.session_state.ui_supports = dxf_data['supports'] if dxf_data else [{'x': 0.0, 'y': 0.0, 'type': 'Hinged', 'angle': 0.0}, {'x': 3.0, 'y': 0.0, 'type': 'Hinged', 'angle': 0.0}]
    if 'ui_struts' not in st.session_state:
        st.session_state.ui_struts = dxf_data['struts'] if dxf_data else []
    if 'ui_loads' not in st.session_state:
        st.session_state.ui_loads = [{'category': 'Dead Load', 'type': 'Uniform', 'dir': 'Global Z (Vertical)' if "XZ" in view_plane else 'Global Y (Vertical)', 'target_mode': 'All Segments', 'target_segs': [], 'w1': 0.0, 'w2': 0.0}]

    origin_x = 0.0
    origin_y = 0.0

    c_in, c_plot = st.columns([1.1, 1.9])
    
    with c_in:
        # ==========================================
        # 1. SUPPORTS
        # ==========================================
        st.markdown("### 1. Supports")
        
        base_segments = dxf_data['base_segments'] if dxf_data else [{'name': 'S1', 'L': 3.0, 'type': 'Straight Line'}]
        base_seg_names = [s.get('name', f"S{i+1}") for i, s in enumerate(base_segments)]

        for i, sup in enumerate(st.session_state.ui_supports):
            cc1, cc2, cc3, cc4, cc5 = st.columns([1.2, 1.2, 1.5, 1.2, 0.4])
            sup['x'] = cc1.number_input(f"J{i+1} X (m)", value=float(sup.get('x', 0.0)), format="%.4f", step=0.1, key=f"sx_{i}_{dxf_v}_{opt_v}")
            sup['y'] = cc2.number_input(f"J{i+1} Y (m)", value=float(sup.get('y', 0.0)), format="%.4f", step=0.1, key=f"sy_{i}_{dxf_v}_{opt_v}")
            
            type_opts = ["Hinged", "Roller", "Fixed"]
            idx_type = type_opts.index(sup.get('type', 'Hinged')) if sup.get('type') in type_opts else 0
            sup['type'] = cc3.selectbox(f"J{i+1} Type", type_opts, index=idx_type, key=f"sp_{i}_{dxf_v}_{opt_v}")
            sup['angle'] = cc4.number_input(f"J{i+1} Angle(°)", value=float(sup.get('angle', 0.0)), step=15.0, key=f"sa_{i}_{dxf_v}_{opt_v}")
            
            cc5.markdown("<br>", unsafe_allow_html=True)
            if cc5.button("❌", key=f"del_sup_{i}_{opt_v}"):
                st.session_state.ui_supports.pop(i)
                st.rerun()

        if st.button("➕ Add Support", use_container_width=True):
            st.session_state.ui_supports.append({'x': 0.0, 'y': 0.0, 'type': 'Hinged', 'angle': 0.0})
            st.rerun()

        # ==========================================
        # 2. GLOBAL & OVERRIDE SECTIONS
        # ==========================================
        st.markdown("### 2. Global & Override Sections")
        
        c_sec1, c_sec2 = st.columns([1, 1.5])
        g_sec_opts = ["Soldier U100", "Custom Section"]
        idx_gsec = g_sec_opts.index(st.session_state.get("g_sec", g_sec_opts[0])) if st.session_state.get("g_sec") in g_sec_opts else 0
        g_sec_type = c_sec1.selectbox("Global Frame Section", g_sec_opts, index=idx_gsec, key="g_sec")
        
        if g_sec_type == "Soldier U100":
            global_sec = {'name': "Soldier U100", 'E': 2100.0, 'A': 34.3 / 10000.0, 'I': 412.0 / 100000000.0, 'Mall': 13.1, 'Qall': 100.8}
            c_sec2.info("Default Acrow Soldier U100 Selected.")
        else:
            with c_sec2.expander("📝 Custom Properties", expanded=True):
                cs_A = st.number_input("Area (cm2)", value=float(st.session_state.get("cs_A_adv", 40.0)), key="cs_A_adv")
                cs_I = st.number_input("Inertia I (cm4)", value=float(st.session_state.get("cs_I_adv", 800.0)), key="cs_I_adv")
                cs_M = st.number_input("Mall (kN.m)", value=float(st.session_state.get("cs_M_adv", 20.0)), key="cs_M_adv")
                cs_Q = st.number_input("Qall (kN)", value=float(st.session_state.get("cs_Q_adv", 120.0)), key="cs_Q_adv")
                global_sec = {'name': "Custom Section", 'E': 2100.0, 'A': cs_A / 10000.0, 'I': cs_I / 100000000.0, 'Mall': cs_M, 'Qall': cs_Q}
        
        seg_sections = [global_sec.copy() for _ in range(len(base_segments))]
        
        with st.expander("🛠️ Override specific segments section", expanded=False):
            raw_override_segs = st.session_state.get("override_segs_adv", [])
            safe_override_segs = [s for s in raw_override_segs if s in base_seg_names]
            
            override_segs = st.multiselect("Select segments to override:", base_seg_names, default=safe_override_segs, key="override_segs_adv")
            if override_segs:
                o_rad_opts = ["Custom Section", "Soldier U100"]
                idx_orad = o_rad_opts.index(st.session_state.get("o_rad", o_rad_opts[0])) if st.session_state.get("o_rad") in o_rad_opts else 0
                o_sec_type = st.radio("Override Profile", o_rad_opts, index=idx_orad, key="o_rad", horizontal=True)
                if o_sec_type == "Soldier U100":
                     o_sec = {'name': "Soldier U100", 'E': 2100.0, 'A': 34.3 / 10000.0, 'I': 412.0 / 100000000.0, 'Mall': 13.1, 'Qall': 100.8}
                else:
                    o1, o2, o3, o4 = st.columns(4)
                    o_A = o1.number_input("A (cm2)", value=float(st.session_state.get("oa", 50.0)), key="oa")
                    o_I = o2.number_input("I (cm4)", value=float(st.session_state.get("oi", 1200.0)), key="oi")
                    o_M = o3.number_input("Mall (kN.m)", value=float(st.session_state.get("om", 30.0)), key="om")
                    o_Q = o4.number_input("Qall (kN)", value=float(st.session_state.get("oq", 150.0)), key="oq")
                    o_sec = {'name': "Custom Override", 'E': 2100.0, 'A': o_A / 10000.0, 'I': o_I / 100000000.0, 'Mall': o_M, 'Qall': o_Q}
                for s_name in override_segs:
                     idx = base_seg_names.index(s_name)
                     seg_sections[idx] = o_sec.copy()

        # ==========================================
        # 3. STRUTS (Push-Pulls)
        # ==========================================
        st.markdown("### 3. Struts (Push-Pulls & Ties)")
        strut_opts_base = list(STRUTS_DB.keys()) if STRUTS_DB else ["PPH 353 (1.5:3.5m)"]
        
        for i, ds in enumerate(st.session_state.ui_struts):
            with st.expander(f"📏 Strut P{i+1}", expanded=True):
                c1, c2, c3, c4, c5 = st.columns([1.5, 1.5, 1.5, 1.5, 0.4])
                ds['tx'] = c1.number_input("Top X (m)", value=float(ds.get('tx', 0.0)), format="%.3f", step=0.1, key=f"st_tx_{i}_{dxf_v}_{opt_v}")
                ds['ty'] = c2.number_input("Top Y (m)", value=float(ds.get('ty', 0.0)), format="%.3f", step=0.1, key=f"st_ty_{i}_{dxf_v}_{opt_v}")
                ds['bx'] = c3.number_input("Bot X (m)", value=float(ds.get('bx', 0.0)), format="%.3f", step=0.1, key=f"st_bx_{i}_{dxf_v}_{opt_v}")
                ds['by'] = c4.number_input("Bot Y (m)", value=float(ds.get('by', 0.0)), format="%.3f", step=0.1, key=f"st_by_{i}_{dxf_v}_{opt_v}")
                
                actual_L = math.hypot(ds['bx'] - ds['tx'], ds['by'] - ds['ty'])
                
                valid_opts = []
                for opt in strut_opts_base:
                    m = re.search(r'\((\d+\.\d+):(\d+\.\d+)m\)', opt)
                    if m and float(m.group(1)) <= actual_L <= float(m.group(2)): valid_opts.append(opt)
                if not valid_opts: valid_opts = strut_opts_base
                valid_opts.sort(key=get_strut_priority)
                
                idx_sec = valid_opts.index(ds.get('sec')) if ds.get('sec') in valid_opts else 0
                ds['sec'] = st.selectbox(f"Type (Length = {actual_L:.3f}m)", valid_opts, index=idx_sec, key=f"st_sec_{i}_{dxf_v}_{opt_v}")
                
                c5.markdown("<br>", unsafe_allow_html=True)
                if c5.button("❌", key=f"del_st_{i}_{opt_v}"):
                    st.session_state.ui_struts.pop(i)
                    st.rerun()

        if st.button("➕ Add Strut", use_container_width=True):
            st.session_state.ui_struts.append({'tx': 0.0, 'ty': 3.0, 'bx': 1.0, 'by': 0.0, 'sec': strut_opts_base[0]})
            st.rerun()

        # ==========================================
        # ✂️ زرار التقطيع الذكي (THE MAGIC BUTTON)
        # ==========================================
        st.markdown("### ✂️ Smart Topology Division")
        st.info("After placing your Supports and Struts, click below to divide the frames perfectly before applying loads.")
        if st.button("✂️ Divide Frames & Update Topology", use_container_width=True, type="primary"):
            div_segs = perform_smart_division(base_segments, st.session_state.ui_supports, st.session_state.ui_struts)
            st.session_state.divided_segments = div_segs
            st.success("✅ Frames Divided Successfully! (e.g., S1 became S1-a, S1-b). You can now apply loads.")

        working_segments = st.session_state.get('divided_segments', base_segments)
        working_seg_names = [s.get('name', f"S{i+1}") for i, s in enumerate(working_segments)]

        active_seg_sections = []
        for s in working_segments:
            m_idx = s.get('master_idx', 0)
            active_seg_sections.append(seg_sections[m_idx] if m_idx < len(seg_sections) else global_sec)

        # ==========================================
        # 🤖 AI GENERATIVE DESIGN OPTIMIZER V14
        # ==========================================
        st.markdown("### 🤖 AI Auto-Shoring Optimizer")
        with st.expander("✨ Generative Design for Bridge Supports", expanded=False):
            st.info("The AI refuses to exceed the Target Load! It will calculate total loads and use mathematical pruning to completely skip impossible combinations.")
            c_ai1, c_ai2 = st.columns(2)
            ai_target_rxn = c_ai1.number_input("Target Max Reaction per Leg (kN)", value=54.4, step=1.0)
            ai_spc_str = c_ai2.text_input("Allowed Spacings (m) [Comma Sep]", value="2.40, 2.10, 1.80, 1.50, 1.20, 0.90, 0.60")
            
            is_sym = st.checkbox("Symmetric Bridge Layout (Recommended)", value=True)
            
            if st.button("✨ Run AI Optimizer & Generate Supports", type="primary", use_container_width=True):
                with st.spinner("🧠 AI is executing Physics-Aware Deep Search..."):
                    succ, res, msg = run_auto_optimizer(
                        base_segments, working_segments, active_seg_sections, 
                        st.session_state.ui_struts, st.session_state.ui_loads, 
                        ai_target_rxn, ai_spc_str, view_plane, auto_mesh_size, dxf_v, is_sym
                    )
                    
                    if succ: # 💡 لو نجح، بيطبق الرصة وينظف الذاكرة
                        st.session_state.ui_supports = res
                        st.session_state.opt_v += 1 
                        new_opt_v = st.session_state.opt_v
                        
                        for idx, sp in enumerate(res):
                            st.session_state[f"sx_{idx}_{dxf_v}_{new_opt_v}"] = sp['x']
                            st.session_state[f"sy_{idx}_{dxf_v}_{new_opt_v}"] = sp['y']
                            st.session_state[f"sp_{idx}_{dxf_v}_{new_opt_v}"] = sp['type']
                            st.session_state[f"sa_{idx}_{dxf_v}_{new_opt_v}"] = sp['angle']
                            
                        for key in list(st.session_state.keys()):
                            if key.startswith("sx_") or key.startswith("sy_") or key.startswith("sp_") or key.startswith("sa_"):
                                parts = key.split('_')
                                if len(parts) >= 4 and int(parts[-1]) != new_opt_v:
                                    del st.session_state[key]
                                    
                        st.success(msg)
                        st.rerun()
                    else:
                        # 💡 لو فشل، بيطبعلك الإيرور ويرفض يوقع الرصة الغلط!
                        st.error(msg)
                        if res: 
                            st.info("⚠️ A fallback grid was calculated but REJECTED because it violated Target Reaction.")

    with c_plot:
        st.markdown("<h3 style='text-align: center; border-bottom: 2px solid #ddd; padding-bottom: 10px; font-family: Arial; color: #1e3d59;'>Live Geometry & Loads Preview</h3>", unsafe_allow_html=True)
        preview_spot = st.empty()

        # ==========================================
        # 5. LOADS (الوراثة الذكية للأحمال)
        # ==========================================
        st.markdown("### 5. Applied Loads")
        cc_d, cc_l, cc_w = st.columns(3)
        fac_d = cc_d.number_input("Dead Factor", value=float(st.session_state.get("cmb_d", 1.00)), step=0.1, format="%.2f", key="cmb_d")
        fac_l = cc_l.number_input("Live Factor", value=float(st.session_state.get("cmb_l", 1.00)), step=0.1, format="%.2f", key="cmb_l")
        fac_w = cc_w.number_input("Wind Factor", value=float(st.session_state.get("cmb_w", 1.00)), step=0.1, format="%.2f", key="cmb_w")
        combo_factors = {'Dead Load': fac_d, 'Live Load': fac_l, 'Wind Load': fac_w}
        
        combined_loads = []
        
        if "XY" in view_plane:
            dir_options = ["Global X (Horizontal)", "Global Y (Vertical)", "Local Y (Perpendicular)"]
        else:
            dir_options = ["Global X (Horizontal)", "Global Z (Vertical)", "Local Z (Perpendicular)"]
        
        for i, ld in enumerate(st.session_state.ui_loads):
            with st.expander(f"📥 Load Item {i+1}", expanded=(i==0)):
                col_l1, col_l2, col_l3, col_l4 = st.columns([1.5, 1.5, 1.5, 0.5])
                
                cat_opts = ["Dead Load", "Live Load", "Wind Load"]
                ld['category'] = col_l1.selectbox("Category", cat_opts, index=cat_opts.index(ld.get('category', 'Dead Load')), key=f"ld_cat_{i}_{dxf_v}")
                
                type_opts = ["Uniform", "Trapezoidal", "Point Load"]
                ld['type'] = col_l2.selectbox("Type", type_opts, index=type_opts.index(ld.get('type', 'Uniform')), key=f"ld_t_{i}_{dxf_v}")
                
                if ld.get('dir') not in dir_options:
                    ld['dir'] = dir_options[1]
                idx_dir = dir_options.index(ld.get('dir'))
                ld['dir'] = col_l3.selectbox("Direction", dir_options, index=idx_dir, key=f"ld_d_{i}_{dxf_v}")
                
                col_l4.markdown("<br>", unsafe_allow_html=True)
                if col_l4.button("❌", key=f"del_ld_{i}"):
                    st.session_state.ui_loads.pop(i)
                    st.rerun()
                
                target_mode_opts = ["Single Segment", "Multiple Segments", "All Segments"]
                idx_tm = target_mode_opts.index(ld.get('target_mode', 'All Segments'))
                target_mode = st.radio("Apply Load To:", target_mode_opts, index=idx_tm, key=f"ld_mode_{i}", horizontal=True)
                ld['target_mode'] = target_mode
                
                target_base_indices = []
                
                if target_mode == "Single Segment":
                    idx_schoice = base_seg_names.index(st.session_state.get(f"ld_single_{i}", base_seg_names[0])) if st.session_state.get(f"ld_single_{i}") in base_seg_names else 0
                    s_choice = st.selectbox("Select Base Segment", base_seg_names, index=idx_schoice, key=f"ld_single_{i}")
                    target_base_indices.append(base_seg_names.index(s_choice))
                elif target_mode == "Multiple Segments":
                    raw_def_segs = st.session_state.get(f"ld_multi_{i}", ld.get('target_segs', []))
                    safe_def_segs = [s for s in raw_def_segs if s in base_seg_names]
                    selected_segs = st.multiselect("Select Base Segments", base_seg_names, default=safe_def_segs, key=f"ld_multi_{i}")
                    ld['target_segs'] = selected_segs
                    target_base_indices = [base_seg_names.index(s) for s in selected_segs]
                else:
                    target_base_indices = list(range(len(base_segments)))
                
                sc1, sc2 = st.columns(2)
                ld['w1'] = sc1.number_input("Value W1", value=float(ld.get('w1', 0.0)), format="%.3f", key=f"ld_w1_{i}_{dxf_v}")
                ld['w2'] = sc2.number_input("Value W2", value=float(ld.get('w2', 0.0)), format="%.3f", key=f"ld_w2_{i}_{dxf_v}") if ld['type'] == "Trapezoidal" else ld['w1']
                
                target_working_indices = [w_idx for w_idx, w_seg in enumerate(working_segments) if w_seg.get('master_idx', 0) in target_base_indices]
                
                for s_idx_num in target_working_indices:
                    max_s = float(working_segments[s_idx_num].get('L', 0.0))
                    combined_loads.append({
                        'seg_idx': s_idx_num, 'category': ld['category'], 'type': ld['type'], 'dir': ld['dir'], 
                        'start': 0.0, 'end': max_s if ld['type'] != 'Point Load' else 0.0, 
                        'w1': ld['w1'] * combo_factors[ld['category']], 'w2': ld['w2'] * combo_factors[ld['category']]
                    })

        if st.button("➕ Add Load", use_container_width=True):
            st.session_state.ui_loads.append({'category': 'Dead Load', 'type': 'Uniform', 'dir': dir_options[1], 'target_mode': 'All Segments', 'target_segs': [], 'w1': 0.0, 'w2': 0.0})
            st.rerun()

        nodes_base, elements_base, nodal_loads_base, display_nodes_base, supports_list_base = build_chain_mesh(
            working_segments, active_seg_sections, combined_loads, st.session_state.ui_struts, st.session_state.ui_supports, base_segments, mesh_size=auto_mesh_size
        )
        live_img = draw_live_preview(nodes_base, elements_base, supports_list_base, active_seg_sections, combined_loads, working_segments)
        preview_spot.image(live_img, use_container_width=True)
        
        st.write("")
        if st.button("🚀 Run Advanced FEA & Generate Diagrams", type="primary", use_container_width=True):
            if 'divided_segments' not in st.session_state:
                st.warning("⚠️ Warning: You haven't clicked 'Divide Frames' yet. The analysis will run on base frames.")
            with st.spinner("Solving Finite Element Matrix..."):
                U_full, R_full, net_load_full = solve_fea_engine(nodes_base, elements_base, nodal_loads_base, supports_list_base) 
                
                total_rxn = sum([R_full[3*sup['node']+1] for sup in supports_list_base])
                
                st.session_state.adv_fea_data = {
                    'U': U_full, 'R': R_full, 'nodes': nodes_base, 'elements': elements_base, 'display_nodes': display_nodes_base,
                    'supports_list': supports_list_base, 'seg_sections': active_seg_sections, 'loads_data': combined_loads, 
                    'segments': working_segments, 'total_load': net_load_full, 'total_rxn': total_rxn
                }
                st.session_state.adv_solved = True
            st.success("✅ Analysis Complete! Scroll down for accurate diagrams.")

    if getattr(st.session_state, 'adv_solved', False):
        st.markdown("---")
        
        fea_data = st.session_state.adv_fea_data
        
        # 💡 شاشة التأكد من الاتزان
        st.markdown("### 📊 Equilibrium & Safety Summary")
        st.info(f"⚖️ **Physics Engine Check:** Total Applied Vertical Load = **{fea_data['total_load']:.2f} kN** ➔ Sum of Vertical Reactions = **{fea_data['total_rxn']:.2f} kN**")
        
        with st.expander("⚙️ Diagram Scale Controls", expanded=False):
            c_s1, c_s2, c_s3 = st.columns(3)
            sc_n = c_s1.slider("Axial Scale", 0.001, 0.100, float(st.session_state.get("adv_sc_n", 0.015)), step=0.001, key="adv_sc_n")
            sc_v = c_s2.slider("Shear Scale", 0.001, 0.100, float(st.session_state.get("adv_sc_v", 0.015)), step=0.001, key="adv_sc_v")
            sc_m = c_s3.slider("Moment Scale", 0.001, 0.100, float(st.session_state.get("adv_sc_m", 0.015)), step=0.001, key="adv_sc_m")
            
        img_bufs = plot_sap2000_diagrams(
            fea_data['nodes'], fea_data['elements'], fea_data['R'], 
            {'N': sc_n, 'V': sc_v, 'M': sc_m}, 
            fea_data['display_nodes'], fea_data['supports_list'], fea_data['seg_sections'], 
            loads=fea_data['loads_data'], segments=fea_data.get('segments')
        )

        titles = {
            'Load_Dead Load': "Dead Load Distribution",
            'Load_Live Load': "Live Load Distribution",
            'Load_Wind Load': "Wind Load Distribution",
            'React': "Reactions (kN)", 'N': "Axial Force (kN)", 'V': "Shear Force (kN)", 'M': "Bending Moment (kN.m)"
        }
        
        load_keys = [k for k in ['Load_Dead Load', 'Load_Live Load', 'Load_Wind Load'] if k in img_bufs]
        if load_keys:
            st.markdown("#### 📥 Applied Loads Diagrams")
            cols_ld = st.columns(len(load_keys))
            for idx, lk in enumerate(load_keys):
                cols_ld[idx].image(img_bufs[lk], use_container_width=True)
                cols_ld[idx].markdown(f"<p align='center' style='font-weight:bold; border-bottom:1px solid #ccc;'>{titles[lk]}</p>", unsafe_allow_html=True)
                
        st.markdown("#### ⚙️ Internal Forces & Reactions")
        c_p1, c_p2 = st.columns(2)
        c_p1.image(img_bufs['React'], use_container_width=True); c_p1.markdown(f"<p align='center' style='font-weight:bold; border-bottom:1px solid #ccc;'>{titles['React']}</p>", unsafe_allow_html=True)
        c_p2.image(img_bufs['N'], use_container_width=True); c_p2.markdown(f"<p align='center' style='font-weight:bold; border-bottom:1px solid #ccc;'>{titles['N']}</p>", unsafe_allow_html=True)
        
        c_p3, c_p4 = st.columns(2)
        c_p3.image(img_bufs['V'], use_container_width=True); c_p3.markdown(f"<p align='center' style='font-weight:bold; border-bottom:1px solid #ccc;'>{titles['V']}</p>", unsafe_allow_html=True)
        c_p4.image(img_bufs['M'], use_container_width=True); c_p4.markdown(f"<p align='center' style='font-weight:bold; border-bottom:1px solid #ccc;'>{titles['M']}</p>", unsafe_allow_html=True)
        
        safety_data = []
        for i, sec in enumerate(fea_data['seg_sections']):
            max_m, max_v = 0.0, 0.0
            for el in fea_data['elements']:
                if el.get('group') == 'segment' and el.get('seg_idx') == i:
                    max_m = max(max_m, np.max(np.abs(el.get('internal', {}).get('M', [0]))))
                    max_v = max(max_v, np.max(np.abs(el.get('internal', {}).get('V', [0]))))
            s_status = "SAFE" if max_m <= sec['Mall'] and max_v <= sec['Qall'] else "UNSAFE"
            safety_data.append({
                "Component": f"{fea_data['segments'][i].get('name', f'S{i+1}')} ({get_short_name(sec['name'])})", "Force Type": "Bending & Shear", 
                "Actual": f"M={max_m:.2f}, V={max_v:.2f}", "Allowable": f"M={sec['Mall']:.2f}, V={sec['Qall']:.2f}", "Status": s_status
            })
            
        st.table(pd.DataFrame(safety_data))
        fea_data['safety_df'] = safety_data; fea_data['img_bufs'] = img_bufs
        
        doc_out = generate_chain_report(fea_data)
        st.download_button("⬇️ Download Calculation Sheet (Word)", data=doc_out.getvalue(), file_name="Advanced_Shape_Calculation_Sheet.docx")

if __name__ == "__main__":
    render_advanced_shape_module()